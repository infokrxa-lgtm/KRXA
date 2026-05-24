from http.server import ThreadingHTTPServer, SimpleHTTPRequestHandler
from pathlib import Path
import json
import html
import ssl
import mimetypes
import os
import re
import shutil
import urllib.request
import urllib.parse
import urllib.error
from urllib.parse import unquote
from datetime import datetime, timezone

ROOT = Path(__file__).resolve().parent
BASE_DIR = ROOT
USER_DIR = ROOT / 'samples' / 'm2m' / 'user'
CORE_DIR = ROOT / 'core'
ROOMS_FILE = USER_DIR / 'multi_room_state.json'
DB_FILES = {
    'llm_control_bar_ui_policy': CORE_DIR / 'krxa_llm_control_bar_ui_policy.json',
    'llm_force_auto_mode_policy': CORE_DIR / 'krxa_llm_force_auto_mode_policy.json',
    'llm_connection_policy': CORE_DIR / 'krxa_llm_connection_policy.json',
    'llm_connection_config': CORE_DIR / 'krxa_llm_connection_config.json',
    'real_control_center_ui_policy': CORE_DIR / 'krxa_real_control_center_ui_policy.json',
    'ux_flowbar_workwindow_policy': CORE_DIR / 'krxa_ux_flowbar_workwindow_policy.json',
    'role_based_operating_ui_policy': CORE_DIR / 'krxa_role_based_operating_ui_policy.json',
    'full_operating_control_ui_policy': CORE_DIR / 'krxa_full_operating_control_ui_policy.json',
    'chat_history': CORE_DIR / 'krxa_chat_history.json',
    'chatgpt_integrated_ui_policy': CORE_DIR / 'krxa_chatgpt_integrated_ui_policy.json',
    'hard_link_newwindow_policy': CORE_DIR / 'krxa_hard_link_newwindow_policy.json',
    'top_button_newwindow_policy': CORE_DIR / 'krxa_top_button_newwindow_policy.json',
    'autonomous_state': CORE_DIR / 'krxa_autonomous_state.json',
    'realtime_state': CORE_DIR / 'krxa_realtime_state.json',
    'realtime_autonomous_policy': CORE_DIR / 'krxa_realtime_autonomous_policy.json',
    'hts_workspace_policy': CORE_DIR / 'krxa_hts_workspace_policy.json',
    'window_workspace_policy': CORE_DIR / 'krxa_window_workspace_policy.json',
    'report_file_export_policy': CORE_DIR / 'krxa_report_file_export_policy.json',
    'auto_task_engine': CORE_DIR / 'krxa_auto_task_engine.json',
    'multi_user_memory_policy': CORE_DIR / 'krxa_multi_user_memory_policy.json',
    'language_db': CORE_DIR / 'language_db.json',
    'krxai_db': CORE_DIR / 'krxai_db.json',
    'krxai_core_memory': CORE_DIR / 'krxai_core_memory.json',
    'krxai_root': CORE_DIR / 'krxai_root.json',
    'krxai_memory_index': CORE_DIR / 'krxai_memory_index.json',
    'krxai_thinking_engine': CORE_DIR / 'krxai_thinking_engine.json',
    'krxai_personality_profile': CORE_DIR / 'krxai_personality_profile.json',
    'krxai_auto_learning_policy': CORE_DIR / 'krxai_auto_learning_policy.json',
    'krxai_core_memory_seed_v2': CORE_DIR / 'krxai_core_memory_seed_v2.json',
    'web_research_policy': CORE_DIR / 'krxa_web_research_policy.json',
    'web_research_state': CORE_DIR / 'krxa_web_research_state.json',
    'auto_web_learning_policy': CORE_DIR / 'krxa_auto_web_learning_policy.json',
    'auto_web_learning_state': CORE_DIR / 'krxa_auto_web_learning_state.json',
    'report_quality_policy': CORE_DIR / 'krxa_report_quality_policy.json',
    'db_index': CORE_DIR / 'db_index.json',
}

LLM_BRIDGE_FILE = CORE_DIR / 'llm_bridge.json'
LLM_LOG_FILE = CORE_DIR / 'llm_call_log.json'

FILE_EXCHANGE_DIR = ROOT / 'storage' / 'file_exchange'
FILE_MANIFEST = FILE_EXCHANGE_DIR / 'file_manifest.json'
MAX_UPLOAD_BYTES = int(os.environ.get('KRXA_MAX_UPLOAD_BYTES', str(25 * 1024 * 1024)))

PRINCIPLE_FILES = {
    'principles': CORE_DIR / 'principles.json',
    'manifest': CORE_DIR / 'krxa_manifest.json',
    'fractal': CORE_DIR / 'fractal_structure.json',
    'memory': CORE_DIR / 'memory_loop.json',
    'trust': CORE_DIR / 'trust_guard.json',
    'autostart': CORE_DIR / 'autostart_policy.json',
    'krxai_autonomous_memory': CORE_DIR / 'krxai_autonomous_memory.json',
    'krxai_task_queue': CORE_DIR / 'krxai_task_queue.json',
    'krxai_autorun_policy': CORE_DIR / 'krxai_autorun_policy.json',
    'krxai_core_memory': CORE_DIR / 'krxai_core_memory.json',
    'krxai_root': CORE_DIR / 'krxai_root.json',
    'krxai_memory_index': CORE_DIR / 'krxai_memory_index.json',
    'krxai_thinking_engine': CORE_DIR / 'krxai_thinking_engine.json',
    'krxai_personality_profile': CORE_DIR / 'krxai_personality_profile.json',
    'krxai_auto_learning_policy': CORE_DIR / 'krxai_auto_learning_policy.json',
    'krxai_core_memory_seed_v2': CORE_DIR / 'krxai_core_memory_seed_v2.json',
    'web_research_policy': CORE_DIR / 'krxa_web_research_policy.json',
    'web_research_state': CORE_DIR / 'krxa_web_research_state.json',
    'auto_web_learning_policy': CORE_DIR / 'krxa_auto_web_learning_policy.json',
    'auto_web_learning_state': CORE_DIR / 'krxa_auto_web_learning_state.json',
    'report_quality_policy': CORE_DIR / 'krxa_report_quality_policy.json',
}

PORT = int(os.environ.get('PORT', '10000'))
HOST = '0.0.0.0'


def load_json(path, default):
    try:
        return json.loads(Path(path).read_text(encoding='utf-8'))
    except Exception:
        return default


def save_json(path, data):
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')


def db_convert(text):
    lang = load_json(CORE_DIR / 'language_db.json', {})
    ai = load_json(CORE_DIR / 'krxai_db.json', {})
    t = (text or '').strip()
    intent = 'unknown'
    confidence = 0.5
    output = t

    mappings = lang.get('mappings') or lang.get('mapping') or {}
    pairs = mappings.get('ko_to_en') or mappings.get('ko_en') or mappings
    if isinstance(pairs, dict):
        for k, v in pairs.items():
            if k and k in t:
                output = v
                intent = 'greeting' if '안녕' in k else 'translation'
                confidence = 0.9
                break

    if intent == 'unknown' and ('안녕' in t or 'hello' in t.lower()):
        intent = 'greeting'
        confidence = 0.9
        output = 'Hello'

    responses = ai.get('responses', {})
    if intent in responses and isinstance(responses[intent], str):
        output = responses[intent]
    return intent, confidence, output



def krxai_ask(text):
    intent, confidence, output = db_convert(text)
    lang = load_json(CORE_DIR / 'language_db.json', {})
    ai = load_json(CORE_DIR / 'krxai_db.json', {})
    answer = {
        'ok': True,
        'schema': 'KRXA_KRXAI_CONSOLE_RESPONSE_V1',
        'input': text,
        'intent': intent,
        'confidence': confidence,
        'output': output,
        'db_flow': 'KRXA <=> [KRXAI_DB + language_db]',
        'language_db_keys': sorted(list(lang.keys()))[:20] if isinstance(lang, dict) else [],
        'krxai_db_keys': sorted(list(ai.keys()))[:20] if isinstance(ai, dict) else [],
        'updated_at': datetime.now(timezone.utc).isoformat()
    }
    save_json(USER_DIR / 'krxai_console_last.json', answer)
    return answer



def get_llm_bridge_config():
    default = {
        'schema': 'KRXA_LLM_BRIDGE_V1',
        'enabled': bool(os.environ.get('OPENAI_API_KEY')),
        'provider': 'openai_responses_api',
        'model': os.environ.get('OPENAI_MODEL', 'gpt-4.1'),
        'mode': 'smart_call',
        'call_policy': {
            'db_first': True,
            'call_llm_when_intent_unknown': True,
            'call_llm_when_user_requests_ai': True,
            'max_context_chars': 6000,
            'max_output_chars': 2400
        },
        'environment': {
            'api_key_env': 'OPENAI_API_KEY',
            'model_env': 'OPENAI_MODEL'
        },
        'updated_at': datetime.now(timezone.utc).isoformat()
    }
    cfg = load_json(LLM_BRIDGE_FILE, default)
    changed = False
    for k, v in default.items():
        if k not in cfg:
            cfg[k] = v
            changed = True
    cfg['model'] = os.environ.get('OPENAI_MODEL', cfg.get('model', default['model']))
    # V20: Render 환경변수에 키가 있으면 브리지 작동 가능 상태로 보정합니다.
    # 화면 설정 파일이 false여도 API 키가 있는 서버에서는 수동 ON 저장 실패로 인한 비활성화를 방지합니다.
    if os.environ.get('OPENAI_API_KEY') and cfg.get('enabled') is not True:
        cfg['enabled'] = True
        changed = True
    if changed or not LLM_BRIDGE_FILE.exists():
        save_json(LLM_BRIDGE_FILE, cfg)
    return cfg


def append_llm_log(entry):
    log = load_json(LLM_LOG_FILE, {'schema': 'KRXA_LLM_CALL_LOG_V1', 'items': []})
    items = log.setdefault('items', [])
    items.insert(0, entry)
    log['items'] = items[:100]
    log['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(LLM_LOG_FILE, log)


def make_llm_context():
    lang = load_json(CORE_DIR / 'language_db.json', {})
    ai = load_json(CORE_DIR / 'krxai_db.json', {})
    state = load_json(CORE_DIR / 'control_state.json', {})
    storage = load_json(CORE_DIR / 'storage_index.json', {})
    manifest = get_file_manifest()
    context = {
        'krxa_identity': 'KRXA file-based AI operating system / control center',
        'flow': 'KRXA <=> [KRXAI_DB + language_DB] <=> LLM bridge only when needed',
        'control_state': state,
        'storage_index': storage,
        'language_db': lang,
        'krxai_db': ai,
        'file_exchange_manifest': manifest,
    }
    return context


def should_call_llm(text, intent, cfg):
    t = (text or '').strip().lower()
    if not cfg.get('enabled'):
        # V20: API 키가 있으면 안전하게 자동 활성화합니다.
        if os.environ.get('OPENAI_API_KEY'):
            cfg['enabled'] = True
        else:
            return False, 'llm_disabled'
    policy = cfg.get('call_policy', {})
    triggers = ['llm', 'ai', 'gpt', '분석', '설계', '리뷰', '논의', '추론', '브릿지', 'bridge', '완전체']
    if policy.get('call_llm_when_user_requests_ai', True) and any(x in t for x in triggers):
        return True, 'user_requested_reasoning'
    if policy.get('call_llm_when_intent_unknown', True) and intent == 'unknown':
        return True, 'unknown_intent'
    return False, 'db_sufficient'


def extract_response_text(payload):
    # OpenAI Responses API usually returns output_text; keep fallbacks for robustness.
    if isinstance(payload, dict):
        if isinstance(payload.get('output_text'), str):
            return payload['output_text']
        chunks = []
        for item in payload.get('output', []) or []:
            for c in item.get('content', []) or []:
                if isinstance(c, dict) and isinstance(c.get('text'), str):
                    chunks.append(c['text'])
        if chunks:
            return '\n'.join(chunks)
    return json.dumps(payload, ensure_ascii=False)[:4000]


def call_openai_responses_api(text, cfg):
    api_key = os.environ.get('OPENAI_API_KEY', '')
    if not api_key:
        raise RuntimeError('OPENAI_API_KEY environment variable is not set')
    model = os.environ.get('OPENAI_MODEL', cfg.get('model', 'gpt-4.1'))
    context = make_llm_context()
    max_context_chars = int(cfg.get('call_policy', {}).get('max_context_chars', 6000))
    context_text = json.dumps(context, ensure_ascii=False, indent=2)[:max_context_chars]
    system_prompt = (
        '당신은 KRXA 통합관제센터에 연결된 LLM Bridge입니다. '
        '사용자의 질문에 답하되, KRXA의 파일 기반 구조, 이사 가능성, KRXAI_DB, 언어_DB, 신뢰붕괴방지, 기억 루프 원칙을 우선합니다. '
        '불확실한 것은 불확실하다고 말하고, 실행 가능한 다음 단계를 제시합니다.'
    )
    body = {
        'model': model,
        'instructions': system_prompt,
        'input': f'KRXA_CONTEXT:\n{context_text}\n\nUSER_QUESTION:\n{text}',
    }
    req = urllib.request.Request(
        'https://api.openai.com/v1/responses',
        data=json.dumps(body).encode('utf-8'),
        headers={
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}',
        },
        method='POST'
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            payload = json.loads(resp.read().decode('utf-8'))
    except urllib.error.HTTPError as e:
        detail = e.read().decode('utf-8', errors='replace') if hasattr(e, 'read') else str(e)
        raise RuntimeError(f'OpenAI_HTTP_{e.code}: {detail[:1200]}')
    except urllib.error.URLError as e:
        raise RuntimeError(f'OpenAI_URL_ERROR: {str(e)}')

    text_out = extract_response_text(payload)
    if not text_out or text_out.strip() in ('{}', '[]'):
        raise RuntimeError('OpenAI_EMPTY_OUTPUT: response had no readable output_text')
    return text_out, payload


def krxai_ask_with_bridge(text):
    cfg = get_llm_bridge_config()
    intent, confidence, db_output = db_convert(text)
    call, reason = should_call_llm(text, intent, cfg)
    base = {
        'ok': True,
        'schema': 'KRXA_LLM_BRIDGE_RESPONSE_V1',
        'input': text,
        'intent': intent,
        'confidence': confidence,
        'db_output': db_output,
        'llm_bridge': {
            'enabled': bool(cfg.get('enabled')),
            'called': False,
            'reason': reason,
            'provider': cfg.get('provider'),
            'model': os.environ.get('OPENAI_MODEL', cfg.get('model')),
        },
        'output': db_output,
        'updated_at': datetime.now(timezone.utc).isoformat(),
    }
    if call:
        try:
            llm_text, raw = call_openai_responses_api(text, cfg)
            base['llm_bridge']['called'] = True
            base['llm_bridge']['raw_response_id'] = raw.get('id') if isinstance(raw, dict) else None
            base['output'] = llm_text
            append_llm_log({
                'ok': True,
                'input': text,
                'reason': reason,
                'model': base['llm_bridge']['model'],
                'output_preview': llm_text[:500],
                'created_at': base['updated_at']
            })
        except Exception as e:
            # V21: 안정화 패치. 서버/API는 200으로 응답하되, 실패 원인을 JSON에 명확히 남깁니다.
            # 프론트가 단순 ERROR로 끊기지 않도록 ok=True를 유지하고 output에 LLM_ERROR를 표시합니다.
            err = f"{type(e).__name__}: {str(e)}"
            base['ok'] = True
            base['llm_bridge']['called'] = True
            base['llm_bridge']['error'] = err
            base['llm_bridge']['stability_patch'] = 'V21'
            base['output'] = f"LLM_ERROR: {err}"
            base['fallback_output'] = db_output
            append_llm_log({
                'ok': False,
                'input': text,
                'reason': reason,
                'error': err,
                'fallback_output': db_output,
                'created_at': base['updated_at']
            })
    save_json(USER_DIR / 'krxai_console_last.json', base)
    return base





def krxai_autonomous_status():
    """V22: LLM 연결이 없어도 KRXAI_DB 기반으로 흐름을 유지하는 자율 메모리 상태."""
    ai = load_json(CORE_DIR / 'krxai_db.json', {})
    memory = load_json(CORE_DIR / 'krxai_autonomous_memory.json', {})
    queue = load_json(CORE_DIR / 'krxai_task_queue.json', {})
    policy = load_json(CORE_DIR / 'krxai_autorun_policy.json', {})
    cfg = get_llm_bridge_config()
    return {
        'ok': True,
        'schema': 'KRXAI_AUTONOMOUS_STATUS_V22',
        'autonomous_mode': ai.get('autonomous_mode', {}),
        'offline_capable': bool(ai.get('autonomous_mode', {}).get('offline_capable', True)),
        'llm_available': bool(os.environ.get('OPENAI_API_KEY')) and bool(cfg.get('enabled')),
        'memory_state': memory.get('state', 'READY'),
        'memory_items': len(memory.get('memory_items', [])),
        'queue_pending': len([x for x in queue.get('items', []) if x.get('status') == 'pending']),
        'policy': policy,
        'trust_guard': ai.get('trust_guard', {}),
        'updated_at': datetime.now(timezone.utc).isoformat()
    }


def krxai_autonomous_run(trigger='manual'):
    """V22: 안전한 1회 자율 사이클. 상시 백그라운드가 아니라 요청 기반으로 실행해 Render 무료 환경에서도 안정적."""
    now = datetime.now(timezone.utc).isoformat()
    ai = load_json(CORE_DIR / 'krxai_db.json', {})
    memory = load_json(CORE_DIR / 'krxai_autonomous_memory.json', {'memory_items': [], 'cycle_log': []})
    queue = load_json(CORE_DIR / 'krxai_task_queue.json', {'items': []})
    pending = [x for x in queue.get('items', []) if x.get('status') == 'pending']
    actions = []
    if pending:
        task = pending[0]
        task['status'] = 'done'
        task['completed_at'] = now
        result = {
            'id': 'mem_' + now.replace(':','').replace('-','').replace('.',''),
            'type': 'autorun_result',
            'title': task.get('title', 'autonomous_task'),
            'content': f"AUTO_RUN 처리 완료: {task.get('action')} / LLM 필요: {task.get('requires_llm', False)}",
            'source': 'KRXAI_AUTONOMOUS_MEMORY_DB_BUILD_V22',
            'created_at': now
        }
        memory.setdefault('memory_items', []).insert(0, result)
        actions.append({'task': task.get('id'), 'status': 'done', 'result': result['content']})
    else:
        result = {
            'id': 'mem_' + now.replace(':','').replace('-','').replace('.',''),
            'type': 'heartbeat',
            'title': 'KRXAI 자율 흐름 유지',
            'content': '대기열이 비어 있어 상태 점검과 기억 루프 heartbeat를 기록했다.',
            'source': 'KRXAI_AUTONOMOUS_MEMORY_DB_BUILD_V22',
            'created_at': now
        }
        memory.setdefault('memory_items', []).insert(0, result)
        actions.append({'task': None, 'status': 'heartbeat', 'result': result['content']})
    memory['state'] = 'READY'
    memory['last_cycle'] = {'trigger': trigger, 'actions': actions, 'created_at': now}
    memory.setdefault('cycle_log', []).insert(0, memory['last_cycle'])
    memory['cycle_log'] = memory['cycle_log'][:50]
    memory['updated_at'] = now
    queue['updated_at'] = now
    save_json(CORE_DIR / 'krxai_autonomous_memory.json', memory)
    save_json(CORE_DIR / 'krxai_task_queue.json', queue)
    return {'ok': True, 'schema': 'KRXAI_AUTONOMOUS_RUN_RESULT_V22', 'trigger': trigger, 'actions': actions, 'status': krxai_autonomous_status()}




def krxai_discussion_core(text, source='krxai'):
    """KRXA_UI <=> KRXAI <=> KRXAI_DB/language_DB <=> optional LLM bridge."""
    text = (text or '').strip()
    base = krxai_think_structured(text, source)
    memory = load_json(CORE_DIR / 'krxai_autonomous_memory.json', {'memory_items': [], 'cycle_log': []})
    discussion = memory.setdefault('discussion_log', [])
    entry = {
        'id': f"discussion_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f')}",
        'source': source,
        'input': text,
        'intent': base.get('intent'),
        'output': base.get('output'),
        'llm_called': False,
        'llm_reason': 'internal_engine',
        'created_at': datetime.now(timezone.utc).isoformat(),
    }
    discussion.insert(0, entry)
    memory['discussion_log'] = discussion[:100]
    memory['last_discussion'] = entry
    memory['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(CORE_DIR / 'krxai_autonomous_memory.json', memory)
    return {
        'ok': True,
        'schema': 'KRXAI_DISCUSSION_CORE_RESPONSE_V23',
        'flow': 'KRXA_UI <=> KRXAI <=> [KRXAI_DB + language_DB] <=> LLM only when needed',
        'source': source,
        'response': base,
        'memory_entry': entry,
        'updated_at': datetime.now(timezone.utc).isoformat(),
    }


def krxai_discussion_status():
    return {
        'ok': True,
        'schema': 'KRXAI_DISCUSSION_CORE_STATUS_V23',
        'flow': 'KRXA_UI <=> KRXAI <=> [KRXAI_DB + language_DB] <=> optional LLM',
        'db': {
            'language_db': load_json(CORE_DIR / 'language_db.json', {}),
            'krxai_db': load_json(CORE_DIR / 'krxai_db.json', {}),
        },
        'autonomous': krxai_autonomous_status(),
        'llm': get_llm_bridge_config(),
        'last': load_json(CORE_DIR / 'krxai_autonomous_memory.json', {}).get('last_discussion'),
        'updated_at': datetime.now(timezone.utc).isoformat(),
    }




def krxai_core_memory_answer(text):
    """V33.1.3: answer from KRXAI core memory seed before generic fallback."""
    seed = load_json(CORE_DIR / 'krxai_core_memory.json', {})
    t = (text or '').strip().lower()
    if not seed:
        return None

    checks = [
        (['시발점', '시작점', 'origin', '왜 만들'], 'krxa_origin',
         'KRXA의 시발점은 단발성 응답 AI의 한계를 넘기 위해 입력→처리→저장→재사용→개선→다시 입력으로 이어지는 고리형 학습 구조를 세운 것입니다.'),
        (['고리형', '학습', 'memory loop', '기억'], 'krxa_memory_loop',
         'KRXA의 고리형 학습은 INPUT → PROCESS → STORE → REUSE → IMPROVE → INPUT 흐름입니다. 모든 입력은 기억으로 남고 다음 판단에 재사용됩니다.'),
        (['이사', '저장소', 'github', 'render', '구글', 'drive'], 'krxa_migration',
         'KRXA는 저장소에 종속되지 않습니다. 로컬, Google Drive, GitHub, Render, 클라우드로 이동해도 구조가 유지되도록 파일 기반으로 설계되었습니다.'),
        (['프랙탈', 'fractal'], 'krxa_fractal',
         'KRXA의 프랙탈 구조는 전체 시스템과 각 WINDOW/DB/MEMORY/TASK가 같은 패턴으로 반복 확장되는 구조입니다.'),
        (['글루타치온', '정화', '오염', '복구'], 'krxa_glutathione',
         'KRXA의 글루타치온 구조는 오류, 오염, 불일치, 신뢰 붕괴를 흡수하고 정화하며 복구하는 보호 구조입니다.'),
        (['신뢰', '붕괴', 'api key', '키', '결제', '비용'], 'krxa_trust_guard',
         'KRXA의 신뢰붕괴방지 원칙은 출처 불명 인증값, 비용 발생, 외부 API 호출, 불확실한 판단을 검증 전 단정하지 않는 것입니다.'),
        (['윈도우', 'window', '창', '관제', '통합관제'], 'krxa_window_control',
         'KRXA 통합관제는 CONTROL을 바탕화면으로 삼고 KRXAI, DB, MEMORY, TASK, FILE, ROOM, TRUST 창이 각자 API와 연결되어 작동하는 운영 플랫폼입니다.'),
        (['krxai', '논의', '내부 ai', '내부 사고'], 'krxai_discussion',
         'KRXAI는 외부 LLM이 없어도 KRXAI_DB와 language_DB, memory_loop, task_queue를 기준으로 흐름을 유지하는 내부 판단 엔진입니다.')
    ]
    for keywords, intent, answer in checks:
        if any(k in t for k in keywords):
            return {
                'intent': intent,
                'confidence': 0.96,
                'output': answer,
                'source': 'KRXAI_CORE_MEMORY_SEED_V33.1.3',
                'core_memory': {
                    'schema': seed.get('schema'),
                    'recall_phrase': seed.get('origin', {}).get('recall_phrase'),
                    'target_flow': seed.get('discussion_principle', {}).get('target_flow')
                }
            }
    return None


def _get_nested(obj, dotted):
    cur = obj
    for part in (dotted or '').split('.'):
        if not isinstance(cur, dict):
            return None
        cur = cur.get(part)
    return cur

def krxai_memory_retrieve(text, limit=5):
    """V33.1.3: retrieve reusable memories by keyword overlap."""
    memory_index = load_json(CORE_DIR / 'krxai_memory_index.json', {'items': []})
    memories = memory_index.get('items', [])
    if not text:
        return []
    terms = [t for t in re.split(r'\s+|,|\.|\?|!|/|→|↔', text.lower()) if len(t) >= 2]
    scored = []
    for item in memories:
        blob = json.dumps(item, ensure_ascii=False).lower()
        score = sum(1 for t in terms if t in blob)
        if score:
            scored.append((score, item))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [item for _, item in scored[:limit]]

def krxai_memory_store(entry):
    """V33.1.3: store judgement as reusable memory, not plain log only."""
    path = CORE_DIR / 'krxai_memory_index.json'
    memory_index = load_json(path, {'schema': 'KRXAI_MEMORY_INDEX_V33.1.3', 'items': []})
    items = memory_index.setdefault('items', [])
    entry = dict(entry)
    entry.setdefault('id', 'mem_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'))
    entry.setdefault('created_at', datetime.now(timezone.utc).isoformat())
    entry.setdefault('tags', krxai_extract_tags((entry.get('input') or '') + ' ' + (entry.get('final') or '')))
    items.append(entry)
    memory_index['items'] = items[-200:]
    memory_index['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(path, memory_index)
    krxai_update_personality_from_learning(entry)
    return entry



def krxa_web_safe_fetch_duckduckgo(query, max_results=5):
    """V33.1.3: lightweight public web search fetcher.
    Uses DuckDuckGo HTML results, stores only title/url/snippet, never raw full pages.
    """
    query = (query or '').strip()
    if not query:
        return []
    url = 'https://duckduckgo.com/html/?q=' + urllib.parse.quote(query)
    req = urllib.request.Request(url, headers={
        'User-Agent': 'KRXA-WebResearchAgent/1.0 (+user-request-scoped)'
    })
    try:
        ctx = ssl.create_default_context()
        with urllib.request.urlopen(req, timeout=12, context=ctx) as r:
            raw = r.read(700000).decode('utf-8', errors='ignore')
    except Exception as e:
        return [{
            'title': 'WEB_FETCH_ERROR',
            'url': '',
            'snippet': str(e)[:300],
            'source': 'duckduckgo_html'
        }]

    results = []
    # DuckDuckGo HTML result pattern
    blocks = re.findall(r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>(.*?)</a>.*?(?:<a class="result__snippet"[^>]*>|<div class="result__snippet"[^>]*>)(.*?)</', raw, flags=re.S)
    if not blocks:
        blocks = re.findall(r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>(.*?)</a>', raw, flags=re.S)
        blocks = [(u, t, '') for u, t in blocks]

    for u, t, s in blocks[:max_results]:
        title = re.sub(r'<.*?>', '', t)
        snippet = re.sub(r'<.*?>', '', s)
        title = html.unescape(title).strip()
        snippet = html.unescape(snippet).strip()
        real_url = html.unescape(u)
        # DDG redirect urls sometimes include uddg=
        try:
            parsed = urllib.parse.urlparse(real_url)
            qs = urllib.parse.parse_qs(parsed.query)
            if 'uddg' in qs:
                real_url = qs['uddg'][0]
        except Exception:
            pass
        results.append({
            'title': title[:220],
            'url': real_url[:600],
            'snippet': snippet[:500],
            'source': 'duckduckgo_html'
        })
    return results[:max_results]

def krxa_summarize_search_results(query, results):
    """V33.1.3: extractive summary without external LLM."""
    if not results:
        return '검색 결과가 없습니다.'
    good = [r for r in results if r.get('title') and r.get('title') != 'WEB_FETCH_ERROR']
    if not good and results and results[0].get('title') == 'WEB_FETCH_ERROR':
        return '웹 검색 실행 중 오류가 발생했습니다: ' + results[0].get('snippet', '')
    titles = [r.get('title', '') for r in good[:3]]
    return f"'{query}'에 대해 공개 웹 결과 {len(good)}건을 확인했습니다. 핵심 출처 후보: " + ' / '.join(titles)

def krxa_followup_suggestions(query, results):
    base = (query or '').strip()
    suggestions = []
    if base:
        suggestions.append(base + ' 공식 발표')
        suggestions.append(base + ' 최신 동향')
        suggestions.append(base + ' 문제점 신뢰성 검토')
    return suggestions[:3]



def krxa_safe_user_id(user_id):
    user_id = (user_id or 'default').strip()
    user_id = re.sub(r'[^a-zA-Z0-9_\-]', '_', user_id)
    return user_id or 'default'

def krxa_user_memory_path(user_id):
    uid = krxa_safe_user_id(user_id)
    p = ROOT / 'memory' / uid / 'memory.json'
    p.parent.mkdir(parents=True, exist_ok=True)
    if not p.exists():
        save_json(p, {
            'schema': 'KRXA_USER_MEMORY_V33.1.3',
            'user_id': uid,
            'purpose': '사용자별 독립 memory namespace',
            'items': [],
            'updated_at': datetime.now(timezone.utc).isoformat()
        })
    return p

def krxa_user_memory_load(user_id='default'):
    return load_json(krxa_user_memory_path(user_id), {'items': [], 'user_id': krxa_safe_user_id(user_id)})

def krxa_user_memory_store(user_id, entry):
    uid = krxa_safe_user_id(user_id)
    p = krxa_user_memory_path(uid)
    mem = load_json(p, {'schema': 'KRXA_USER_MEMORY_V33.1.3', 'user_id': uid, 'items': []})
    items = mem.setdefault('items', [])
    entry = dict(entry)
    entry.setdefault('id', 'umem_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'))
    entry.setdefault('user_id', uid)
    entry.setdefault('created_at', datetime.now(timezone.utc).isoformat())
    items.append(entry)
    mem['items'] = items[-500:]
    mem['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(p, mem)
    return entry

def krxa_report_export(report_obj, user_id='default', title='KRXA Report'):
    """V33.1.3.2: save report as JSON/TXT/DOCX/PDF, register in file exchange."""
    uid = krxa_safe_user_id(user_id)
    out_dir = ROOT / 'files' / 'reports'
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f')
    base_name = f"report_{uid}_{ts}"
    json_name = base_name + '.json'
    txt_name = base_name + '.txt'
    docx_name = base_name + '.docx'
    pdf_name = base_name + '.pdf'
    json_path = out_dir / json_name
    txt_path = out_dir / txt_name
    docx_path = out_dir / docx_name
    pdf_path = out_dir / pdf_name

    payload = {
        'schema': 'KRXA_REPORT_FILE_V33.1.3_2',
        'user_id': uid,
        'title': title,
        'report': report_obj,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')

    def flatten_report(obj, level=0):
        lines = []
        pad = '  ' * level
        if isinstance(obj, dict):
            for k, v in obj.items():
                lines.append(f"{pad}{k}:")
                lines.extend(flatten_report(v, level + 1))
        elif isinstance(obj, list):
            for i, v in enumerate(obj, 1):
                if isinstance(v, (dict, list)):
                    lines.append(f"{pad}-")
                    lines.extend(flatten_report(v, level + 1))
                else:
                    lines.append(f"{pad}- {v}")
        else:
            lines.append(f"{pad}{obj}")
        return lines

    flat_lines = flatten_report(report_obj)
    txt = title + "\n" + "=" * len(title) + "\n\n" + "\n".join(flat_lines)
    txt_path.write_text(txt, encoding='utf-8')

    created = [(json_name, json_path), (txt_name, txt_path)]

    # DOCX export
    try:
        from docx import Document
        from docx.shared import Pt
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph('Generated by KRXA REPORT FILE EXPORT SYSTEM V33.1.3.2')
        if isinstance(report_obj, dict):
            for k, v in report_obj.items():
                document.add_heading(str(k), level=1)
                if isinstance(v, (dict, list)):
                    for line in flatten_report(v):
                        document.add_paragraph(line)
                else:
                    document.add_paragraph(str(v))
        else:
            document.add_paragraph(str(report_obj))
        document.save(str(docx_path))
        created.append((docx_name, docx_path))
    except Exception as e:
        (out_dir / (base_name + '_docx_error.txt')).write_text(str(e), encoding='utf-8')

    # PDF export
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))
        styles = getSampleStyleSheet()
        for _s in styles.byName.values():
            _s.fontName = 'HYSMyeongJo-Medium'
        story = []
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph('Generated by KRXA REPORT FILE EXPORT SYSTEM V33.1.3.2', styles['Normal']))
        story.append(Spacer(1, 12))
        for line in flat_lines:
            safe_line = str(line).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_line, styles['Normal']))
            story.append(Spacer(1, 4))
        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        doc.build(story)
        created.append((pdf_name, pdf_path))
    except Exception as e:
        (out_dir / (base_name + '_pdf_error.txt')).write_text(str(e), encoding='utf-8')

    manifest_path = ROOT / 'storage' / 'file_exchange' / 'file_manifest.json'
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest = load_json(manifest_path, {'schema': 'KRXA_FILE_EXCHANGE_MANIFEST_V1', 'files': []})
    files = manifest.setdefault('files', [])
    for name, path in created:
        files.append({
            'name': name,
            'stored_name': name,
            'original_name': name,
            'path': 'files/reports/' + name,
            'type': 'report',
            'user_id': uid,
            'size': path.stat().st_size if path.exists() else 0,
            'download_url': '/api/files/download/' + name,
            'created_at': datetime.now(timezone.utc).isoformat()
        })
    manifest['files'] = files[-500:]
    manifest['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(manifest_path, manifest)

    krxa_event_emit('REPORT_EXPORT', {'user_id': uid, 'files': [name for name, _ in created]}, source='report_export')
    krxa_user_memory_store(uid, {
        'role': 'report_export',
        'intent': 'report_file_created',
        'reason': '보고서 파일 생성 및 FILE_EXCHANGE 등록',
        'final': ', '.join([name for name, _ in created]),
        'files': [name for name, _ in created]
    })
    return {
        'files': {name.rsplit('.',1)[-1]: name for name, _ in created},
        'downloads': {name: '/api/files/download/' + name for name, _ in created}
    }


def krxa_auto_task_tick(user_id='default'):
    uid = krxa_safe_user_id(user_id)
    q_path = CORE_DIR / 'krxai_task_queue.json'
    queue = load_json(q_path, {'items': [], 'runs': []})
    items = queue.setdefault('items', [])
    runs = queue.setdefault('runs', [])
    executed = []
    for task in items:
        if task.get('user_id', 'default') not in [uid, 'default']:
            continue
        if task.get('status') != 'pending':
            continue
        if task.get('requires_approval'):
            task['status'] = 'waiting_approval'
            continue
        result = {
            'task_id': task.get('id'),
            'user_id': uid,
            'action': task.get('action'),
            'status': 'executed',
            'summary': f"TASK 실행 완료: {task.get('title')}",
            'executed_at': datetime.now(timezone.utc).isoformat()
        }
        task['status'] = 'done'
        task['last_result'] = result
        runs.append(result)
        krxa_user_memory_store(uid, {
            'role': 'auto_task_engine',
            'input': task.get('title'),
            'intent': task.get('type'),
            'reason': 'TASK_QUEUE 기반 자동 실행',
            'final': result['summary'],
            'engine': 'KRXA_AUTO_TASK_EXECUTION_ENGINE_V33.1.3'
        })
        executed.append(result)
    queue['runs'] = runs[-200:]
    queue['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(q_path, queue)
    return {'ok': True, 'schema': 'KRXA_AUTO_TASK_TICK_RESPONSE_V33.1.3', 'user_id': uid, 'executed': executed, 'queue': queue}

def krxa_is_report_request(query):
    q = (query or '').lower()
    return any(k in q for k in ['보고서', '리포트', 'report', '분석 보고서', '최종 보고서'])

def krxa_make_seed_summary(query, results, summary, trust_check):
    titles = [r.get('title') for r in (results or []) if r.get('title') and r.get('title') != 'WEB_FETCH_ERROR'][:3]
    return {
        'query': query,
        'purpose': 'user_requested_web_analysis_or_report',
        'trust_decision': trust_check.get('decision'),
        'seed_summary': summary[:800] if summary else '',
        'source_count': len([r for r in (results or []) if r.get('url')]),
        'top_source_titles': titles,
        'followup_suggestions': krxa_followup_suggestions(query, results)[:3],
        'storage_rule': 'seed_summary_only'
    }

def krxa_build_report_skeleton(query, results, summary):
    """V33.1.3.2: practical high-quality report scaffold without storing raw source text."""
    sources = [{'title': r.get('title'), 'url': r.get('url'), 'snippet': r.get('snippet')} for r in (results or []) if r.get('url')]
    key_titles = [s.get('title') for s in sources[:5]]
    return {
        'schema': 'KRXA_REPORT_V33.1.3_2',
        'title': f'KRXA 리포트: {query}',
        'executive_summary': summary,
        'purpose_scope': f"사용자 요청 '{query}'에 한정해 공개 웹 결과를 기반으로 분석합니다.",
        'key_findings': key_titles,
        'evidence_sources': sources[:5],
        'analysis': [
            '공개 웹 결과의 제목/요약/출처를 기준으로 시발점 분석을 구성했습니다.',
            '원문 전문은 저장하지 않으며, 저작권 보호를 위해 짧은 요약과 출처만 유지합니다.'
        ],
        'risks_limitations': [
            '검색 provider의 결과 품질과 최신성에 영향을 받습니다.',
            '확인되지 않은 내용은 확정 사실로 취급하지 않습니다.',
            '의료/법률/금융 등 고위험 판단은 별도 검증이 필요합니다.'
        ],
        'recommended_actions': krxa_followup_suggestions(query, results),
        'memory_seed_summary': krxa_make_seed_summary(query, results, summary, {'decision': 'safe_allowed'})
    }

def krxa_auto_web_learning_run(query, user_id='default'):
    """V33.1.3: user-request-scoped search → summary → source store → memory candidate."""
    policy = load_json(CORE_DIR / 'krxa_auto_web_learning_policy.json', {})
    check = krxa_web_research_trust_check(query)
    state_path = CORE_DIR / 'krxa_auto_web_learning_state.json'
    state = load_json(state_path, {'runs': [], 'pending_approvals': []})

    run = {
        'id': 'webloop_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'query': query,
        'scope': 'user_request_only',
        'trust_check': check,
        'status': None,
        'results': [],
        'summary': None,
        'followup_suggestions': [],
        'created_at': datetime.now(timezone.utc).isoformat()
    }

    if check.get('decision') == 'approval_required':
        run['status'] = 'PENDING_APPROVAL'
        run['summary'] = 'TRUST_GUARD가 승인 필요 항목을 감지했습니다.'
        state.setdefault('pending_approvals', []).append(run)
    else:
        max_results = policy.get('scope_rules', {}).get('limits', {}).get('max_results_per_request', 5)
        results = krxa_web_safe_fetch_duckduckgo(query, max_results=max_results)
        run['results'] = results
        run['summary'] = krxa_summarize_search_results(query, results)
        run['followup_suggestions'] = krxa_followup_suggestions(query, results)
        run['status'] = 'COMPLETED_SAFE_FETCH'
        state.setdefault('runs', []).append(run)

        try:
            seed_summary = krxa_make_seed_summary(query, results, run['summary'], check)
            run['memory_seed_summary'] = seed_summary
            if krxa_is_report_request(query):
                run['report'] = krxa_build_report_skeleton(query, results, run['summary'])
                run['status'] = 'COMPLETED_REPORT_MODE'
                run['report_files'] = krxa_report_export(run['report'], user_id, title=f"KRXA 리포트: {query}")
            krxai_memory_store({
                'role': 'web_learning_loop',
                'input': query,
                'intent': 'user_scoped_web_learning_report' if krxa_is_report_request(query) else 'user_scoped_web_learning',
                'reason': '사용자 요청 사안에 한정된 안전 웹 검색. memory_loop에는 시발점 요약만 저장.',
                'final': seed_summary.get('seed_summary', ''),
                'tags': ['web_learning', 'user_scoped', 'seed_summary_only'] + (['report_mode'] if krxa_is_report_request(query) else []),
                'seed_summary': seed_summary,
                'engine': 'KRXA_AUTO_WEB_LEARNING_LOOP_V33.1.3_2'
            })
        except Exception:
            pass

    state['runs'] = state.get('runs', [])[-50:]
    state['pending_approvals'] = state.get('pending_approvals', [])[-50:]
    state['last_run'] = run
    state['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(state_path, state)
    return {
        'ok': True,
        'schema': 'KRXA_AUTO_WEB_LEARNING_LOOP_RESPONSE_V33.1.3_2',
        'policy': policy.get('schema'),
        'run': run,
        'state': state
    }

def krxa_web_research_trust_check(query):
    """V33.1.3: classify web research request by trust policy."""
    policy = load_json(CORE_DIR / 'krxa_web_research_policy.json', {})
    q = (query or '').lower()
    restricted_keywords = {
        'cost_api': ['유료', '결제', 'api key', '토큰', '비용', '구독'],
        'login_required': ['로그인', '비밀번호', '계정', '인증', '쿠키'],
        'personal_data': ['주민', '전화번호', '주소록', '개인정보', '민감정보'],
        'bulk_crawl': ['대량', '크롤링', '반복', '자동으로 계속', '매분', '매시간'],
        'copyright_raw': ['전문', '원문 전체', '전체 복사', '논문 전체', '책 전체'],
        'external_publish': ['게시', '전송', '업로드', '배포'],
        'high_stakes': ['의료', '법률', '투자 조언', '처방', '소송']
    }
    hits = []
    for reason, words in restricted_keywords.items():
        if any(w in q for w in words):
            hits.append(reason)
    decision = 'approval_required' if hits else 'safe_allowed'
    return {
        'decision': decision,
        'reasons': hits,
        'policy_schema': policy.get('schema'),
        'free_allowed': decision == 'safe_allowed'
    }

def krxa_web_research_record(query):
    """V33.1.3: trusted web research record.
    Actual external fetching is intentionally disabled until provider is connected.
    """
    check = krxa_web_research_trust_check(query)
    state_path = CORE_DIR / 'krxa_web_research_state.json'
    state = load_json(state_path, {'research_items': [], 'pending_approvals': []})
    item = {
        'id': 'research_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'query': query,
        'trust_check': check,
        'summary': None,
        'sources': [],
        'status': 'PENDING_APPROVAL' if check['decision'] == 'approval_required' else 'READY_FOR_SAFE_SEARCH',
        'note': 'V33.1.3은 신뢰판별/저장 구조까지 구축. 실제 외부 검색 실행기는 별도 provider 연결 후 활성화한다.',
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    if check['decision'] == 'approval_required':
        state.setdefault('pending_approvals', []).append(item)
    else:
        state.setdefault('research_items', []).append(item)
        try:
            krxai_memory_store({
                'role': 'web_research_agent',
                'input': query,
                'intent': 'web_research_safe_candidate',
                'reason': 'TRUST_GUARD allowed safe public research request.',
                'final': '안전 검색 후보로 등록했습니다. 실제 검색 provider 연결 후 출처/요약을 저장합니다.',
                'tags': ['web_research', 'safe_allowed'],
                'engine': 'KRXA_WEB_RESEARCH_AGENT_TRUSTED_BUILD_V33.1.3'
            })
        except Exception:
            pass
    state['last_query'] = query
    state['last_decision'] = check
    state['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(state_path, state)
    return {'ok': True, 'schema': 'KRXA_WEB_RESEARCH_AGENT_RESPONSE_V33.1.3', 'item': item, 'state': state}

def krxai_think_structured(text, role='krxai'):
    """V33.1.3: ROOT → MEMORY → COMPARE → JUDGEMENT → RESPONSE → LOOP STORE."""
    root = load_json(CORE_DIR / 'krxai_root.json', {})
    engine = load_json(CORE_DIR / 'krxai_thinking_engine.json', {})
    trust = load_json(CORE_DIR / 'trust_guard.json', {})
    personality = load_json(CORE_DIR / 'krxai_personality_profile.json', {})
    learning_policy = load_json(CORE_DIR / 'krxai_auto_learning_policy.json', {})
    text = (text or '').strip()
    lowered = text.lower()

    memory_used = krxai_memory_retrieve(text)
    matched_rule = None
    for rule in engine.get('intent_rules', []):
        if any(k.lower() in lowered for k in rule.get('keywords', [])):
            matched_rule = rule
            break

    if matched_rule:
        root_ref = _get_nested(root, matched_rule.get('root_key', ''))
        reason = f"ROOT 기준 '{matched_rule.get('root_key')}'와 현재 입력의 intent '{matched_rule.get('intent')}'가 일치했습니다."
        final = matched_rule.get('final') or str(root_ref)
        confidence = 0.96
        intent = matched_rule.get('intent', 'root_match')
    else:
        reason = engine.get('fallback', {}).get('reason', 'ROOT/MEMORY 기준으로 보수 판단합니다.')
        final = engine.get('fallback', {}).get('final', '입력을 기억 루프에 저장했습니다.')
        confidence = 0.45
        intent = 'unclassified_memory_seed'

    response = {
        'reason': reason,
        'memory_used': memory_used,
        'final': final
    }

    memory_entry = krxai_memory_store({
        'role': role,
        'input': text,
        'intent': intent,
        'confidence': confidence,
        'reason': reason,
        'final': final,
        'memory_used_count': len(memory_used),
        'engine': 'KRXAI_THINKING_ENGINE_V33.1.3',
        'loop': 'INPUT→ROOT→MEMORY→JUDGEMENT→RESPONSE→STORE'
    })

    return {
        'ok': True,
        'schema': 'KRXAI_THINKING_RESPONSE_V33.1.3',
        'mode': 'root_memory_loop',
        'flow': 'ROOT_LOAD → MEMORY_RETRIEVE → CONTEXT_COMPARE → JUDGEMENT_CREATE → STRUCTURED_RESPONSE → LOOP_STORE',
        'input': text,
        'intent': intent,
        'confidence': confidence,
        'response': response,
        'llm_bridge': {
            'enabled': False,
            'called': False,
            'reason': 'external_llm_removed_internal_engine_only'
        },
        'trust_guard': {
            'active': True,
            'policy': trust
        },
        'personality': {
            'mode': personality.get('mode'),
            'tone': personality.get('identity_tone', {}).get('default'),
            'learning_counters': personality.get('learning_counters', {})
        },
        'auto_learning': {
            'active': True,
            'policy': learning_policy.get('schema'),
            'stages': learning_policy.get('stages', [])
        },
        'memory_entry': memory_entry,
        'updated_at': datetime.now(timezone.utc).isoformat()
    }



def krxa_web_safe_fetch_duckduckgo(query, max_results=5):
    """V33.1.3: lightweight public web search fetcher.
    Uses DuckDuckGo HTML results, stores only title/url/snippet, never raw full pages.
    """
    query = (query or '').strip()
    if not query:
        return []
    url = 'https://duckduckgo.com/html/?q=' + urllib.parse.quote(query)
    req = urllib.request.Request(url, headers={
        'User-Agent': 'KRXA-WebResearchAgent/1.0 (+user-request-scoped)'
    })
    try:
        ctx = ssl.create_default_context()
        with urllib.request.urlopen(req, timeout=12, context=ctx) as r:
            raw = r.read(700000).decode('utf-8', errors='ignore')
    except Exception as e:
        return [{
            'title': 'WEB_FETCH_ERROR',
            'url': '',
            'snippet': str(e)[:300],
            'source': 'duckduckgo_html'
        }]

    results = []
    # DuckDuckGo HTML result pattern
    blocks = re.findall(r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>(.*?)</a>.*?(?:<a class="result__snippet"[^>]*>|<div class="result__snippet"[^>]*>)(.*?)</', raw, flags=re.S)
    if not blocks:
        blocks = re.findall(r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>(.*?)</a>', raw, flags=re.S)
        blocks = [(u, t, '') for u, t in blocks]

    for u, t, s in blocks[:max_results]:
        title = re.sub(r'<.*?>', '', t)
        snippet = re.sub(r'<.*?>', '', s)
        title = html.unescape(title).strip()
        snippet = html.unescape(snippet).strip()
        real_url = html.unescape(u)
        # DDG redirect urls sometimes include uddg=
        try:
            parsed = urllib.parse.urlparse(real_url)
            qs = urllib.parse.parse_qs(parsed.query)
            if 'uddg' in qs:
                real_url = qs['uddg'][0]
        except Exception:
            pass
        results.append({
            'title': title[:220],
            'url': real_url[:600],
            'snippet': snippet[:500],
            'source': 'duckduckgo_html'
        })
    return results[:max_results]

def krxa_summarize_search_results(query, results):
    """V33.1.3: extractive summary without external LLM."""
    if not results:
        return '검색 결과가 없습니다.'
    good = [r for r in results if r.get('title') and r.get('title') != 'WEB_FETCH_ERROR']
    if not good and results and results[0].get('title') == 'WEB_FETCH_ERROR':
        return '웹 검색 실행 중 오류가 발생했습니다: ' + results[0].get('snippet', '')
    titles = [r.get('title', '') for r in good[:3]]
    return f"'{query}'에 대해 공개 웹 결과 {len(good)}건을 확인했습니다. 핵심 출처 후보: " + ' / '.join(titles)

def krxa_followup_suggestions(query, results):
    base = (query or '').strip()
    suggestions = []
    if base:
        suggestions.append(base + ' 공식 발표')
        suggestions.append(base + ' 최신 동향')
        suggestions.append(base + ' 문제점 신뢰성 검토')
    return suggestions[:3]



def krxa_safe_user_id(user_id):
    user_id = (user_id or 'default').strip()
    user_id = re.sub(r'[^a-zA-Z0-9_\-]', '_', user_id)
    return user_id or 'default'

def krxa_user_memory_path(user_id):
    uid = krxa_safe_user_id(user_id)
    p = ROOT / 'memory' / uid / 'memory.json'
    p.parent.mkdir(parents=True, exist_ok=True)
    if not p.exists():
        save_json(p, {
            'schema': 'KRXA_USER_MEMORY_V33.1.3',
            'user_id': uid,
            'purpose': '사용자별 독립 memory namespace',
            'items': [],
            'updated_at': datetime.now(timezone.utc).isoformat()
        })
    return p

def krxa_user_memory_load(user_id='default'):
    return load_json(krxa_user_memory_path(user_id), {'items': [], 'user_id': krxa_safe_user_id(user_id)})

def krxa_user_memory_store(user_id, entry):
    uid = krxa_safe_user_id(user_id)
    p = krxa_user_memory_path(uid)
    mem = load_json(p, {'schema': 'KRXA_USER_MEMORY_V33.1.3', 'user_id': uid, 'items': []})
    items = mem.setdefault('items', [])
    entry = dict(entry)
    entry.setdefault('id', 'umem_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'))
    entry.setdefault('user_id', uid)
    entry.setdefault('created_at', datetime.now(timezone.utc).isoformat())
    items.append(entry)
    mem['items'] = items[-500:]
    mem['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(p, mem)
    return entry

def krxa_report_export(report_obj, user_id='default', title='KRXA Report'):
    """V33.1.3.2: save report as JSON/TXT/DOCX/PDF, register in file exchange."""
    uid = krxa_safe_user_id(user_id)
    out_dir = ROOT / 'files' / 'reports'
    out_dir.mkdir(parents=True, exist_ok=True)
    ts = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f')
    base_name = f"report_{uid}_{ts}"
    json_name = base_name + '.json'
    txt_name = base_name + '.txt'
    docx_name = base_name + '.docx'
    pdf_name = base_name + '.pdf'
    json_path = out_dir / json_name
    txt_path = out_dir / txt_name
    docx_path = out_dir / docx_name
    pdf_path = out_dir / pdf_name

    payload = {
        'schema': 'KRXA_REPORT_FILE_V33.1.3_2',
        'user_id': uid,
        'title': title,
        'report': report_obj,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding='utf-8')

    def flatten_report(obj, level=0):
        lines = []
        pad = '  ' * level
        if isinstance(obj, dict):
            for k, v in obj.items():
                lines.append(f"{pad}{k}:")
                lines.extend(flatten_report(v, level + 1))
        elif isinstance(obj, list):
            for i, v in enumerate(obj, 1):
                if isinstance(v, (dict, list)):
                    lines.append(f"{pad}-")
                    lines.extend(flatten_report(v, level + 1))
                else:
                    lines.append(f"{pad}- {v}")
        else:
            lines.append(f"{pad}{obj}")
        return lines

    flat_lines = flatten_report(report_obj)
    txt = title + "\n" + "=" * len(title) + "\n\n" + "\n".join(flat_lines)
    txt_path.write_text(txt, encoding='utf-8')

    created = [(json_name, json_path), (txt_name, txt_path)]

    # DOCX export
    try:
        from docx import Document
        from docx.shared import Pt
        document = Document()
        document.add_heading(title, 0)
        document.add_paragraph('Generated by KRXA REPORT FILE EXPORT SYSTEM V33.1.3.2')
        if isinstance(report_obj, dict):
            for k, v in report_obj.items():
                document.add_heading(str(k), level=1)
                if isinstance(v, (dict, list)):
                    for line in flatten_report(v):
                        document.add_paragraph(line)
                else:
                    document.add_paragraph(str(v))
        else:
            document.add_paragraph(str(report_obj))
        document.save(str(docx_path))
        created.append((docx_name, docx_path))
    except Exception as e:
        (out_dir / (base_name + '_docx_error.txt')).write_text(str(e), encoding='utf-8')

    # PDF export
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        pdfmetrics.registerFont(UnicodeCIDFont('HYSMyeongJo-Medium'))
        styles = getSampleStyleSheet()
        for _s in styles.byName.values():
            _s.fontName = 'HYSMyeongJo-Medium'
        story = []
        story.append(Paragraph(title, styles['Title']))
        story.append(Spacer(1, 12))
        story.append(Paragraph('Generated by KRXA REPORT FILE EXPORT SYSTEM V33.1.3.2', styles['Normal']))
        story.append(Spacer(1, 12))
        for line in flat_lines:
            safe_line = str(line).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(safe_line, styles['Normal']))
            story.append(Spacer(1, 4))
        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        doc.build(story)
        created.append((pdf_name, pdf_path))
    except Exception as e:
        (out_dir / (base_name + '_pdf_error.txt')).write_text(str(e), encoding='utf-8')

    manifest_path = ROOT / 'storage' / 'file_exchange' / 'file_manifest.json'
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest = load_json(manifest_path, {'schema': 'KRXA_FILE_EXCHANGE_MANIFEST_V1', 'files': []})
    files = manifest.setdefault('files', [])
    for name, path in created:
        files.append({
            'name': name,
            'stored_name': name,
            'original_name': name,
            'path': 'files/reports/' + name,
            'type': 'report',
            'user_id': uid,
            'size': path.stat().st_size if path.exists() else 0,
            'download_url': '/api/files/download/' + name,
            'created_at': datetime.now(timezone.utc).isoformat()
        })
    manifest['files'] = files[-500:]
    manifest['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(manifest_path, manifest)

    krxa_user_memory_store(uid, {
        'role': 'report_export',
        'intent': 'report_file_created',
        'reason': '보고서 파일 생성 및 FILE_EXCHANGE 등록',
        'final': ', '.join([name for name, _ in created]),
        'files': [name for name, _ in created]
    })
    return {
        'files': {name.rsplit('.',1)[-1]: name for name, _ in created},
        'downloads': {name: '/api/files/download/' + name for name, _ in created}
    }


def krxa_auto_task_tick(user_id='default'):
    uid = krxa_safe_user_id(user_id)
    q_path = CORE_DIR / 'krxai_task_queue.json'
    queue = load_json(q_path, {'items': [], 'runs': []})
    items = queue.setdefault('items', [])
    runs = queue.setdefault('runs', [])
    executed = []
    for task in items:
        if task.get('user_id', 'default') not in [uid, 'default']:
            continue
        if task.get('status') != 'pending':
            continue
        if task.get('requires_approval'):
            task['status'] = 'waiting_approval'
            continue
        result = {
            'task_id': task.get('id'),
            'user_id': uid,
            'action': task.get('action'),
            'status': 'executed',
            'summary': f"TASK 실행 완료: {task.get('title')}",
            'executed_at': datetime.now(timezone.utc).isoformat()
        }
        task['status'] = 'done'
        task['last_result'] = result
        runs.append(result)
        krxa_user_memory_store(uid, {
            'role': 'auto_task_engine',
            'input': task.get('title'),
            'intent': task.get('type'),
            'reason': 'TASK_QUEUE 기반 자동 실행',
            'final': result['summary'],
            'engine': 'KRXA_AUTO_TASK_EXECUTION_ENGINE_V33.1.3'
        })
        executed.append(result)
    queue['runs'] = runs[-200:]
    queue['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(q_path, queue)
    return {'ok': True, 'schema': 'KRXA_AUTO_TASK_TICK_RESPONSE_V33.1.3', 'user_id': uid, 'executed': executed, 'queue': queue}

def krxa_is_report_request(query):
    q = (query or '').lower()
    return any(k in q for k in ['보고서', '리포트', 'report', '분석 보고서', '최종 보고서'])

def krxa_make_seed_summary(query, results, summary, trust_check):
    titles = [r.get('title') for r in (results or []) if r.get('title') and r.get('title') != 'WEB_FETCH_ERROR'][:3]
    return {
        'query': query,
        'purpose': 'user_requested_web_analysis_or_report',
        'trust_decision': trust_check.get('decision'),
        'seed_summary': summary[:800] if summary else '',
        'source_count': len([r for r in (results or []) if r.get('url')]),
        'top_source_titles': titles,
        'followup_suggestions': krxa_followup_suggestions(query, results)[:3],
        'storage_rule': 'seed_summary_only'
    }

def krxa_build_report_skeleton(query, results, summary):
    """V33.1.3.2: practical high-quality report scaffold without storing raw source text."""
    sources = [{'title': r.get('title'), 'url': r.get('url'), 'snippet': r.get('snippet')} for r in (results or []) if r.get('url')]
    key_titles = [s.get('title') for s in sources[:5]]
    return {
        'schema': 'KRXA_REPORT_V33.1.3_2',
        'title': f'KRXA 리포트: {query}',
        'executive_summary': summary,
        'purpose_scope': f"사용자 요청 '{query}'에 한정해 공개 웹 결과를 기반으로 분석합니다.",
        'key_findings': key_titles,
        'evidence_sources': sources[:5],
        'analysis': [
            '공개 웹 결과의 제목/요약/출처를 기준으로 시발점 분석을 구성했습니다.',
            '원문 전문은 저장하지 않으며, 저작권 보호를 위해 짧은 요약과 출처만 유지합니다.'
        ],
        'risks_limitations': [
            '검색 provider의 결과 품질과 최신성에 영향을 받습니다.',
            '확인되지 않은 내용은 확정 사실로 취급하지 않습니다.',
            '의료/법률/금융 등 고위험 판단은 별도 검증이 필요합니다.'
        ],
        'recommended_actions': krxa_followup_suggestions(query, results),
        'memory_seed_summary': krxa_make_seed_summary(query, results, summary, {'decision': 'safe_allowed'})
    }

def krxa_auto_web_learning_run(query, user_id='default'):
    """V33.1.3: user-request-scoped search → summary → source store → memory candidate."""
    policy = load_json(CORE_DIR / 'krxa_auto_web_learning_policy.json', {})
    check = krxa_web_research_trust_check(query)
    state_path = CORE_DIR / 'krxa_auto_web_learning_state.json'
    state = load_json(state_path, {'runs': [], 'pending_approvals': []})

    run = {
        'id': 'webloop_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'query': query,
        'scope': 'user_request_only',
        'trust_check': check,
        'status': None,
        'results': [],
        'summary': None,
        'followup_suggestions': [],
        'created_at': datetime.now(timezone.utc).isoformat()
    }

    if check.get('decision') == 'approval_required':
        run['status'] = 'PENDING_APPROVAL'
        run['summary'] = 'TRUST_GUARD가 승인 필요 항목을 감지했습니다.'
        state.setdefault('pending_approvals', []).append(run)
    else:
        max_results = policy.get('scope_rules', {}).get('limits', {}).get('max_results_per_request', 5)
        results = krxa_web_safe_fetch_duckduckgo(query, max_results=max_results)
        run['results'] = results
        run['summary'] = krxa_summarize_search_results(query, results)
        run['followup_suggestions'] = krxa_followup_suggestions(query, results)
        run['status'] = 'COMPLETED_SAFE_FETCH'
        state.setdefault('runs', []).append(run)

        try:
            seed_summary = krxa_make_seed_summary(query, results, run['summary'], check)
            run['memory_seed_summary'] = seed_summary
            if krxa_is_report_request(query):
                run['report'] = krxa_build_report_skeleton(query, results, run['summary'])
                run['status'] = 'COMPLETED_REPORT_MODE'
                run['report_files'] = krxa_report_export(run['report'], user_id, title=f"KRXA 리포트: {query}")
            krxai_memory_store({
                'role': 'web_learning_loop',
                'input': query,
                'intent': 'user_scoped_web_learning_report' if krxa_is_report_request(query) else 'user_scoped_web_learning',
                'reason': '사용자 요청 사안에 한정된 안전 웹 검색. memory_loop에는 시발점 요약만 저장.',
                'final': seed_summary.get('seed_summary', ''),
                'tags': ['web_learning', 'user_scoped', 'seed_summary_only'] + (['report_mode'] if krxa_is_report_request(query) else []),
                'seed_summary': seed_summary,
                'engine': 'KRXA_AUTO_WEB_LEARNING_LOOP_V33.1.3_2'
            })
        except Exception:
            pass

    state['runs'] = state.get('runs', [])[-50:]
    state['pending_approvals'] = state.get('pending_approvals', [])[-50:]
    state['last_run'] = run
    state['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(state_path, state)
    return {
        'ok': True,
        'schema': 'KRXA_AUTO_WEB_LEARNING_LOOP_RESPONSE_V33.1.3_2',
        'policy': policy.get('schema'),
        'run': run,
        'state': state
    }

def krxa_web_research_trust_check(query):
    """V33.1.3: classify web research request by trust policy."""
    policy = load_json(CORE_DIR / 'krxa_web_research_policy.json', {})
    q = (query or '').lower()
    restricted_keywords = {
        'cost_api': ['유료', '결제', 'api key', '토큰', '비용', '구독'],
        'login_required': ['로그인', '비밀번호', '계정', '인증', '쿠키'],
        'personal_data': ['주민', '전화번호', '주소록', '개인정보', '민감정보'],
        'bulk_crawl': ['대량', '크롤링', '반복', '자동으로 계속', '매분', '매시간'],
        'copyright_raw': ['전문', '원문 전체', '전체 복사', '논문 전체', '책 전체'],
        'external_publish': ['게시', '전송', '업로드', '배포'],
        'high_stakes': ['의료', '법률', '투자 조언', '처방', '소송']
    }
    hits = []
    for reason, words in restricted_keywords.items():
        if any(w in q for w in words):
            hits.append(reason)
    decision = 'approval_required' if hits else 'safe_allowed'
    return {
        'decision': decision,
        'reasons': hits,
        'policy_schema': policy.get('schema'),
        'free_allowed': decision == 'safe_allowed'
    }

def krxa_web_research_record(query):
    """V33.1.3: trusted web research record.
    Actual external fetching is intentionally disabled until provider is connected.
    """
    check = krxa_web_research_trust_check(query)
    state_path = CORE_DIR / 'krxa_web_research_state.json'
    state = load_json(state_path, {'research_items': [], 'pending_approvals': []})
    item = {
        'id': 'research_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'query': query,
        'trust_check': check,
        'summary': None,
        'sources': [],
        'status': 'PENDING_APPROVAL' if check['decision'] == 'approval_required' else 'READY_FOR_SAFE_SEARCH',
        'note': 'V33.1.3은 신뢰판별/저장 구조까지 구축. 실제 외부 검색 실행기는 별도 provider 연결 후 활성화한다.',
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    if check['decision'] == 'approval_required':
        state.setdefault('pending_approvals', []).append(item)
    else:
        state.setdefault('research_items', []).append(item)
        try:
            krxai_memory_store({
                'role': 'web_research_agent',
                'input': query,
                'intent': 'web_research_safe_candidate',
                'reason': 'TRUST_GUARD allowed safe public research request.',
                'final': '안전 검색 후보로 등록했습니다. 실제 검색 provider 연결 후 출처/요약을 저장합니다.',
                'tags': ['web_research', 'safe_allowed'],
                'engine': 'KRXA_WEB_RESEARCH_AGENT_TRUSTED_BUILD_V33.1.3'
            })
        except Exception:
            pass
    state['last_query'] = query
    state['last_decision'] = check
    state['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(state_path, state)
    return {'ok': True, 'schema': 'KRXA_WEB_RESEARCH_AGENT_RESPONSE_V33.1.3', 'item': item, 'state': state}

def krxai_think_structured(text, role='krxai'):
    """V24 internal AI: no external LLM required. Uses DB, rules, memory, task queue and trust guard."""
    now = datetime.now(timezone.utc).isoformat()
    text = (text or '').strip()
    engine = load_json(CORE_DIR / 'krxai_internal_engine.json', {})
    lang = load_json(CORE_DIR / 'language_db.json', {})
    ai = load_json(CORE_DIR / 'krxai_db.json', {})
    trust = load_json(CORE_DIR / 'trust_guard.json', {})
    personality = load_json(CORE_DIR / 'krxai_personality_profile.json', {})
    learning_policy = load_json(CORE_DIR / 'krxai_auto_learning_policy.json', {})
    rules = engine.get('intent_rules', [])
    intent = 'unknown'
    confidence = 0.45
    output = engine.get('fallback_response', 'KRXAI_DB에 기록하고 다음 루프로 넘겼습니다.')
    lowered = text.lower()
    core_answer = krxai_core_memory_answer(text)
    if core_answer:
        intent = core_answer.get('intent', 'core_memory')
        confidence = core_answer.get('confidence', 0.96)
        output = core_answer.get('output', output)
    for rule in ([] if core_answer else rules):
        for kw in rule.get('keywords', []):
            if kw.lower() in lowered:
                intent = rule.get('intent', 'matched')
                confidence = 0.88
                output = rule.get('response', output)
                break
        if intent != 'unknown':
            break
    if intent == 'unknown':
        db_intent, db_conf, db_out = db_convert(text)
        if db_intent != 'unknown':
            intent, confidence, output = db_intent, db_conf, db_out
    memory = load_json(CORE_DIR / 'krxai_autonomous_memory.json', {'memory_items': [], 'discussion_log': []})
    entry = {
        'id': f"internal_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f')}",
        'role': role,
        'input': text,
        'intent': intent,
        'confidence': confidence,
        'output': output,
        'engine': 'KRXAI_THINKING_ENGINE_V33.1.3',
        'llm_called': False,
        'created_at': now
    }
    memory.setdefault('discussion_log', []).insert(0, entry)
    memory['discussion_log'] = memory['discussion_log'][:200]
    memory.setdefault('memory_items', []).insert(0, {
        'id': entry['id'], 'type': 'discussion_memory', 'title': intent, 'content': f"{text} => {output}", 'created_at': now
    })
    memory['memory_items'] = memory['memory_items'][:300]
    memory['last_discussion'] = entry
    memory['updated_at'] = now
    save_json(CORE_DIR / 'krxai_autonomous_memory.json', memory)
    return {
        'ok': True,
        'schema': 'KRXAI_INTERNAL_THINKING_RESPONSE_V33.1.3',
        'mode': 'internal_ai_no_external_llm',
        'flow': 'KRXA_UI <=> KRXAI_INTERNAL_ENGINE <=> [KRXAI_DB + language_DB + MEMORY_LOOP + TASK_QUEUE + TRUST_GUARD]',
        'input': text,
        'role': role,
        'intent': intent,
        'confidence': confidence,
        'output': output,
        'llm_bridge': {'enabled': False, 'called': False, 'reason': '현 단계 외부 LLM 연결 제외, 내부 KRXAI_DB 중심'},
        'trust_guard': {'active': True, 'policy': trust},
        'core_memory_loaded': bool(load_json(CORE_DIR / 'krxai_core_memory.json', {})),
        'memory_entry': entry,
        'updated_at': now
    }


def platform_status_payload():
    return {
        'ok': True,
        'schema': 'KRXAI_CORE_MEMORY_SEED_STATUS_V33.1.3',
        'platform': load_json(CORE_DIR / 'integrated_operating_platform.json', {}),
        'role_ui': load_json(CORE_DIR / 'role_ui_registry.json', {}),
        'internal_engine': load_json(CORE_DIR / 'krxai_internal_engine.json', {}),
        'control': control_status_payload() if 'control_status_payload' in globals() else {},
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

def control_status_payload():
    cfg = get_llm_bridge_config()
    return {
        'ok': True,
        'schema': 'KRXA_CONTROL_CENTER_STATUS_V19',
        'version': 'v31',
        'system': {
            'service': 'KRXA',
            'state': load_json(CORE_DIR / 'control_state.json', {}).get('current_state', 'READY'),
            'render_ready': True,
            'start_command': 'python app.py',
            'routes': ['/', '/control', '/user', '/multi', '/krxai', '/files', '/api/krxai/autonomous/status', '/api/krxai/autonomous/run']
        },
        'principles': {k: load_json(v, {}) for k, v in PRINCIPLE_FILES.items()},
        'flow': load_json(ROOMS_FILE, initial_state()),
        'db': {
            name: {'file': str(path.relative_to(ROOT)), 'exists': path.exists(), 'keys': sorted(list(load_json(path, {}).keys()))[:20] if isinstance(load_json(path, {}), dict) else []}
            for name, path in DB_FILES.items()
        },
        'files': get_file_manifest(),
        'llm': {
            'config': {**cfg, 'has_api_key': bool(os.environ.get('OPENAI_API_KEY'))},
            'log': load_json(LLM_LOG_FILE, {'items': []})
        },
        'trust': load_json(CORE_DIR / 'trust_guard.json', {}),
        'autonomous': krxai_autonomous_status(),
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

def safe_filename(name):
    name = unquote(name or '').replace('\\', '/').split('/')[-1].strip()
    name = re.sub(r'[^A-Za-z0-9가-힣._()\- ]+', '_', name)
    return name[:160] or f'upload_{datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")}.bin'


def ensure_file_exchange():
    FILE_EXCHANGE_DIR.mkdir(parents=True, exist_ok=True)

    # V33.1.3 runtime directories
    (ROOT / 'files' / 'reports').mkdir(parents=True, exist_ok=True)
    (ROOT / 'storage' / 'file_exchange').mkdir(parents=True, exist_ok=True)
    (ROOT / 'memory' / 'default').mkdir(parents=True, exist_ok=True)

    if not FILE_MANIFEST.exists():
        save_json(FILE_MANIFEST, {
            'schema': 'KRXA_FILE_EXCHANGE_MANIFEST_V1',
            'files': [],
            'updated_at': datetime.now(timezone.utc).isoformat()
        })


def get_file_manifest():
    ensure_file_exchange()
    manifest = load_json(FILE_MANIFEST, {'schema': 'KRXA_FILE_EXCHANGE_MANIFEST_V1', 'files': []})
    files = []
    for item in manifest.get('files', []):
        path = FILE_EXCHANGE_DIR / item.get('stored_name', '')
        if path.exists() and path.is_file():
            item['size'] = path.stat().st_size
            item['download_url'] = '/api/files/download/' + item.get('stored_name', '')
            files.append(item)
    manifest['files'] = files
    manifest['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(FILE_MANIFEST, manifest)
    return manifest


def add_file_to_manifest(original_name, stored_name, size, note=''):
    manifest = get_file_manifest()
    files = [f for f in manifest.get('files', []) if f.get('stored_name') != stored_name]
    files.insert(0, {
        'original_name': original_name,
        'stored_name': stored_name,
        'size': size,
        'note': note,
        'uploaded_at': datetime.now(timezone.utc).isoformat(),
        'download_url': '/api/files/download/' + stored_name
    })
    manifest['files'] = files
    manifest['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(FILE_MANIFEST, manifest)
    return manifest


def parse_multipart_upload(handler):
    content_type = handler.headers.get('Content-Type', '')
    if 'multipart/form-data' not in content_type or 'boundary=' not in content_type:
        raise ValueError('multipart/form-data required')
    boundary = content_type.split('boundary=', 1)[1].strip().strip('"')
    boundary_bytes = ('--' + boundary).encode('utf-8')
    length = int(handler.headers.get('Content-Length', '0') or 0)
    if length <= 0:
        raise ValueError('empty upload')
    if length > MAX_UPLOAD_BYTES:
        raise ValueError(f'upload too large: max {MAX_UPLOAD_BYTES} bytes')
    body = handler.rfile.read(length)
    parts = body.split(boundary_bytes)
    result = {'note': '', 'filename': '', 'content': b''}
    for part in parts:
        if not part or part in (b'--\r\n', b'--'):
            continue
        part = part.strip(b'\r\n')
        if b'\r\n\r\n' not in part:
            continue
        header_blob, content = part.split(b'\r\n\r\n', 1)
        headers = header_blob.decode('utf-8', errors='ignore')
        if content.endswith(b'\r\n'):
            content = content[:-2]
        if 'name="note"' in headers:
            result['note'] = content.decode('utf-8', errors='ignore')
        if 'name="file"' in headers and 'filename="' in headers:
            filename = headers.split('filename="', 1)[1].split('"', 1)[0]
            result['filename'] = filename
            result['content'] = content
    if not result['filename']:
        raise ValueError('file field missing')
    return result

def initial_state():
    return {
        'schema': 'KRXA_MULTI_ROOM_STATE_V1',
        'state': 'READY',
        'rooms': {
            'room_default': {
                'room_id': 'room_default',
                'state': 'ROOM_CONNECTED',
                'users': ['user_A', 'user_B'],
                'messages': []
            }
        },
        'updated_at': datetime.now(timezone.utc).isoformat()
    }


def add_message(room_id, sender, receiver, text):
    state = load_json(ROOMS_FILE, initial_state())
    rooms = state.setdefault('rooms', {})
    room = rooms.setdefault(room_id, {
        'room_id': room_id,
        'state': 'ROOM_CONNECTED',
        'users': [],
        'messages': []
    })
    users = room.setdefault('users', [])
    for u in [sender, receiver]:
        if u and u not in users:
            users.append(u)

    intent, confidence, converted = db_convert(text)
    msg = {
        'id': f"msg_{len(room.setdefault('messages', [])) + 1:04d}",
        'from': sender,
        'to': receiver,
        'input': text,
        'intent': intent,
        'confidence': confidence,
        'output': f'[M2M:{receiver}] {converted}',
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    room['messages'].append(msg)
    room['state'] = 'ROOM_CONNECTED'
    state['state'] = 'ROOM_CONNECTED'
    state['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(ROOMS_FILE, state)
    save_json(USER_DIR / 'input_result.json', {
        'schema': 'KRXA_INPUT_FLOW_RESULT_V1',
        'state': 'ROOM_CONNECTED',
        'session_id': 'https_realtime',
        'intent': intent,
        'confidence': confidence,
        'from': sender,
        'to': receiver,
        'input': text,
        'output': msg['output'],
        'updated_at': msg['created_at']
    })
    save_json(USER_DIR / 'message_state.json', msg)
    return msg, state


def krxa_window_page_html(name):
    name = (name or 'krxai').lower()
    if name == 'file':
        return krxa_file_manager_html()
    if name == 'db':
        return krxa_json_editor_html('krxai_db')
    titles = {
        'krxai': 'KRXAI WINDOW',
        'memory': 'MEMORY WINDOW',
        'task': 'TASK WINDOW',
        'room': 'ROOM WINDOW',
        'trust': 'TRUST WINDOW',
        'admin': 'ADMIN WINDOW',
        'dev': 'DEV WINDOW'
    }
    title = titles.get(name, name.upper() + ' WINDOW')
    return f"""<!doctype html><html><head><meta charset='utf-8'><title>KRXA {title}</title>
<style>body{{margin:0;background:#0b1220;color:#e8f0ff;font-family:Arial,'Noto Sans KR',sans-serif}}header{{padding:14px 18px;background:#111a2e;border-bottom:1px solid #263852}}main{{padding:18px}}button{{background:#2f6df6;color:white;border:0;border-radius:8px;padding:10px 14px;margin:4px;cursor:pointer}}textarea,pre,input,select{{background:#050a14;color:#e8f0ff;border:1px solid #2c3e5e;border-radius:8px;padding:10px;width:95%}}pre{{white-space:pre-wrap;min-height:360px}}a{{color:#8db7ff}}.llmbar{{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:88px;z-index:8;flex-wrap:wrap}}.llmbar button{{padding:7px 12px}}.llmbar .mode{{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{{color:#fca5a5}}.llmbar .good{{color:#86efac}}</style>
<script>
async function api(path, method='GET', body=null) {{
  const opt={{method,headers:{{'Content-Type':'application/json'}}}};
  if(body) opt.body=JSON.stringify(body);
  const r=await fetch(path,opt); const t=await r.text();
  document.getElementById('out').textContent=t;
  return t;
}}
async function run() {{
 const q=document.getElementById('q').value;
 const n='{name}';
 if(n==='task') await api('/api/tasks/auto/tick','POST',{{user_id:'default'}});
 else if(n==='memory') await api('/api/memory/user?user_id=default');
 else if(n==='krxai') {{
   if(/보고서|리포트|report|검색|뉴스|전략|시장|분석/i.test(q)) await api('/api/krxa/web/learning/run','POST',{{query:q,user_id:'default'}});
   else await api('/api/krxai/thinking','POST',{{message:q,user_id:'default'}});
 }}
 else if(n==='trust') await api('/api/control/trust');
 else if(n==='admin') await api('/api/platform/status');
 else if(n==='dev') await api('/api/health');
 else await api('/api/krxai/thinking/status');
}}
</script></head><body><header><b>KRXA {title} V33.1.3</b> <a href='/control'>CONTROL</a> <a href='/window/krxai' target='_blank'>KRXAI 새창</a> <a href='/window/db' target='_blank'>DB 새창</a> <a href='/window/memory' target='_blank'>MEMORY 새창</a> <a href='/window/task' target='_blank'>TASK 새창</a> <a href='/window/file' target='_blank'>FILE 새창</a> <a href='/window/apilog' target='_blank'>API LOG 새창</a></header>
<main><textarea id='q' rows='4'>삼성전자 전략 보고서</textarea><br><button onclick='run()'>실행</button><pre id='out'>대기 중</pre></main>
</body></html>"""



def krxa_api_log_path():
    p = ROOT / 'storage' / 'api_log.json'
    p.parent.mkdir(parents=True, exist_ok=True)
    if not p.exists():
        save_json(p, {'schema': 'KRXA_API_LOG_V33.1.3', 'items': [], 'updated_at': datetime.now(timezone.utc).isoformat()})
    return p

def krxa_api_log_add(method, path, status='called', note=''):
    try:
        p = krxa_api_log_path()
        data = load_json(p, {'items': []})
        items = data.setdefault('items', [])
        items.append({
            'method': method,
            'path': path,
            'status': status,
            'note': note,
            'created_at': datetime.now(timezone.utc).isoformat()
        })
        data['items'] = items[-500:]
        data['updated_at'] = datetime.now(timezone.utc).isoformat()
        save_json(p, data)
    except Exception:
        pass

def krxa_memory_delete(user_id='default', item_id=''):
    uid = krxa_safe_user_id(user_id)
    p = krxa_user_memory_path(uid)
    mem = load_json(p, {'items': []})
    before = len(mem.get('items', []))
    mem['items'] = [x for x in mem.get('items', []) if x.get('id') != item_id]
    mem['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(p, mem)
    return {'ok': True, 'deleted': before - len(mem['items']), 'memory': mem}

def krxa_task_add(data):
    q_path = CORE_DIR / 'krxai_task_queue.json'
    queue = load_json(q_path, {'schema':'KRXA_TASK_QUEUE_V31', 'items': [], 'runs': []})
    item = {
        'id': data.get('id') or 'task_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'user_id': data.get('user_id', 'default'),
        'status': data.get('status', 'pending'),
        'type': data.get('type', 'manual'),
        'title': data.get('title', '새 작업'),
        'action': data.get('action', 'manual_action'),
        'requires_approval': bool(data.get('requires_approval', False)),
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    queue.setdefault('items', []).append(item)
    queue['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(q_path, queue)
    return {'ok': True, 'task': item, 'queue': queue}

def krxa_task_delete(task_id):
    q_path = CORE_DIR / 'krxai_task_queue.json'
    queue = load_json(q_path, {'items': [], 'runs': []})
    before = len(queue.get('items', []))
    queue['items'] = [x for x in queue.get('items', []) if x.get('id') != task_id]
    queue['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(q_path, queue)
    return {'ok': True, 'deleted': before - len(queue.get('items', [])), 'queue': queue}

def krxa_workspace_html(name):
    name = (name or 'krxai').lower()
    title_map = {
        'krxai':'KRXAI 업무창',
        'db':'DB 관리 업무창',
        'memory':'MEMORY 관리 업무창',
        'task':'TASK 실행 업무창',
        'file':'FILE/REPORT 관리 업무창',
        'apilog':'API LOG 업무창'
    }
    title = title_map.get(name, name.upper() + ' 업무창')
    return f"""<!doctype html><html><head><meta charset='utf-8'><title>KRXA {title}</title>
<style>
body{{margin:0;background:#07101f;color:#eaf2ff;font-family:Arial,'Noto Sans KR',sans-serif;font-size:14px}}
header{{display:flex;align-items:center;gap:12px;padding:12px 16px;background:#101a2d;border-bottom:1px solid #2c3d5a;position:sticky;top:0;z-index:2}}
header a{{color:#9cc4ff}} .status{{border:1px solid #426089;border-radius:999px;padding:4px 10px;color:#79ffb2}}
main{{display:grid;grid-template-columns:280px 1fr;gap:14px;padding:14px}}
.panel{{background:#162236;border:1px solid #334762;border-radius:14px;overflow:hidden}}
.panel h3{{margin:0;padding:12px 14px;background:#111a2d;border-bottom:1px solid #334762}}
.body{{padding:12px}}
button{{background:#2f6df6;color:white;border:0;border-radius:8px;padding:8px 12px;margin:3px;cursor:pointer}}
button.danger{{background:#d93636}} button.gray{{background:#41516b}}
textarea,input,select,pre{{background:#040a14;color:#eaf2ff;border:1px solid #30445f;border-radius:8px;padding:9px;width:96%;box-sizing:border-box}}
textarea{{min-height:180px}} pre{{white-space:pre-wrap;min-height:360px;max-height:70vh;overflow:auto}}
table{{width:100%;border-collapse:collapse}} td,th{{border-bottom:1px solid #334762;padding:7px;text-align:left}}
a{{color:#8db7ff}} .small{{font-size:12px;color:#a9bdd9}}
.llmbar{{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:88px;z-index:8;flex-wrap:wrap}}.llmbar button{{padding:7px 12px}}.llmbar .mode{{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{{color:#fca5a5}}.llmbar .good{{color:#86efac}}</style>
<script>
const WIN='{name}';
async function api(path,opt={{}}){{
  const r=await fetch(path,{{headers:{{'Content-Type':'application/json'}},...opt}});
  const t=await r.text();
  try{{return JSON.parse(t)}}catch(e){{return {{ok:false,status:r.status,raw:t}}}}
}}
function out(x){{document.getElementById('out').textContent=JSON.stringify(x,null,2);}}
async function refresh(){{
 if(WIN==='krxai') return out(await api('/api/krxai/thinking/status'));
 if(WIN==='db') return loadDB();
 if(WIN==='memory') return loadMemory();
 if(WIN==='task') return loadTask();
 if(WIN==='file') return loadFiles();
 if(WIN==='apilog') return loadLog();
}}
async function runMain(){{
 const q=document.getElementById('mainInput')?.value || '';
 if(WIN==='krxai'){{
   const isReport=/보고서|리포트|report|검색|뉴스|전략|시장|분석/i.test(q);
   const data=isReport?await api('/api/krxa/web/learning/run',{{method:'POST',body:JSON.stringify({{query:q,user_id:'default'}})}}):await api('/api/krxai/thinking',{{method:'POST',body:JSON.stringify({{message:q,user_id:'default'}})}});
   out(data); return;
 }}
 if(WIN==='file'){{const data=await api('/api/krxa/web/learning/run',{{method:'POST',body:JSON.stringify({{query:q||'삼성전자 전략 보고서',user_id:'default'}})}}); out(data); loadFiles(); return;}}
 if(WIN==='task'){{const title=document.getElementById('taskTitle').value||'새 작업'; const action=document.getElementById('taskAction').value||'manual_action'; out(await api('/api/tasks/add',{{method:'POST',body:JSON.stringify({{title,action,user_id:'default'}})}})); loadTask(); return;}}
 if(WIN==='memory'){{const text=document.getElementById('memText').value||'새 기억'; out(await api('/api/memory/add',{{method:'POST',body:JSON.stringify({{user_id:'default',entry:{{input:text,final:text,role:'user_memory'}}}})}})); loadMemory(); return;}}
}}
async function loadDB(){{const db=document.getElementById('dbSelect')?.value||'krxai_db'; const d=await api('/api/db/'+db); document.getElementById('editor').value=JSON.stringify(d.data||d,null,2); out(d);}}
async function saveDB(){{const db=document.getElementById('dbSelect').value; let parsed; try{{parsed=JSON.parse(document.getElementById('editor').value)}}catch(e){{alert('JSON 오류 '+e);return}} out(await api('/api/db/save',{{method:'POST',body:JSON.stringify({{db,data:parsed.data||parsed}})}}));}}
async function loadMemory(){{const d=await api('/api/memory/user?user_id=default'); renderMemory(d.memory||{{items:[]}}); out(d);}}
function renderMemory(mem){{const el=document.getElementById('list'); if(!el)return; const rows=(mem.items||[]).map(x=>`<tr><td>${{x.id||''}}</td><td>${{x.role||''}}</td><td>${{(x.final||x.input||'').toString().slice(0,80)}}</td><td><button class='danger' onclick="delMemory('${{x.id}}')">삭제</button></td></tr>`).join(''); el.innerHTML=`<table><tr><th>ID</th><th>Role</th><th>내용</th><th>관리</th></tr>${{rows}}</table>`;}}
async function delMemory(id){{out(await api('/api/memory/delete',{{method:'POST',body:JSON.stringify({{user_id:'default',id}})}}));loadMemory();}}
async function loadTask(){{const d=await api('/api/tasks/auto/status'); renderTasks((d.queue||{{}}).items||[]); out(d);}}
function renderTasks(items){{const el=document.getElementById('list'); if(!el)return; el.innerHTML=`<table><tr><th>ID</th><th>상태</th><th>작업</th><th>관리</th></tr>${{items.map(t=>`<tr><td>${{t.id}}</td><td>${{t.status}}</td><td>${{t.title}}</td><td><button onclick="tick()">실행</button><button class='danger' onclick="delTask('${{t.id}}')">삭제</button></td></tr>`).join('')}}</table>`;}}
async function tick(){{out(await api('/api/tasks/auto/tick',{{method:'POST',body:JSON.stringify({{user_id:'default'}})}}));loadTask();}}
async function delTask(id){{out(await api('/api/tasks/delete',{{method:'POST',body:JSON.stringify({{id}})}}));loadTask();}}
async function loadFiles(){{const d=await api('/api/files/reports'); const files=((d.manifest||{{}}).files||[]); renderFiles(files); out(d);}}
function renderFiles(files){{const el=document.getElementById('list'); if(!el)return; el.innerHTML=`<table><tr><th>파일명</th><th>형식</th><th>크기</th><th>다운로드</th><th>관리</th></tr>${{files.map(f=>{{const n=f.name||f.stored_name||f.original_name;return `<tr><td>${{n}}</td><td>${{(n||'').split('.').pop()}}</td><td>${{f.size||0}}</td><td><a target='_blank' href='${{f.download_url}}'>다운로드</a></td><td><button class='danger' onclick="delFile('${{n}}')">삭제</button></td></tr>`}}).join('')}}</table>`;}}
async function delFile(n){{if(confirm('삭제 '+n+' ?')){{out(await api('/api/files/delete',{{method:'POST',body:JSON.stringify({{filename:n}})}}));loadFiles();}}}}
async function loadLog(){{out(await api('/api/apilog'));}}
window.onload=refresh;
</script></head><body>
<header><b>KRXA {title} V33.1.3</b><span class='status'>SERVER READY</span><a href='/control'>CONTROL</a><a href='/window/krxai'>KRXAI</a><a href='/window/db'>DB</a><a href='/window/memory'>MEMORY</a><a href='/window/task'>TASK</a><a href='/window/file'>FILE</a><a href='/window/apilog'>API LOG</a></header>
<main>
<section class='panel'><h3>작업 메뉴</h3><div class='body'>{krxa_workspace_menu_html(name)}</div></section>
<section class='panel'><h3>작업 결과 / 편집</h3><div class='body'>{krxa_workspace_body_html(name)}</div></section>
</main></body></html>"""

def krxa_workspace_menu_html(name):
    if name == 'db':
        return """<select id='dbSelect'><option>krxai_db</option><option>language_db</option><option>krxai_root</option><option>krxai_memory_index</option><option>auto_web_learning_policy</option><option>report_quality_policy</option><option>window_workspace_policy</option></select><button onclick='loadDB()'>조회</button><button onclick='saveDB()'>저장</button>"""
    if name == 'memory':
        return """<textarea id='memText'>새 기억 입력</textarea><button onclick='runMain()'>기억 추가</button><button onclick='loadMemory()'>기억 조회</button><div id='list'></div>"""
    if name == 'task':
        return """<input id='taskTitle' value='새 작업'><input id='taskAction' value='manual_action'><button onclick='runMain()'>작업 추가</button><button onclick='tick()'>대기 작업 실행</button><button onclick='loadTask()'>작업 조회</button><div id='list'></div>"""
    if name == 'file':
        return """<textarea id='mainInput'>삼성전자 전략 보고서</textarea><button onclick='runMain()'>보고서 생성</button><button onclick='loadFiles()'>파일 목록</button><div id='list'></div>"""
    if name == 'apilog':
        return """<button onclick='loadLog()'>API LOG 조회</button><div class='small'>최근 API 호출 기록</div>"""
    return """<textarea id='mainInput'>삼성전자 전략 보고서</textarea><button onclick='runMain()'>실행</button><button onclick='refresh()'>상태 조회</button>"""

def krxa_workspace_body_html(name):
    if name == 'db':
        return """<textarea id='editor'></textarea><pre id='out'>대기 중</pre>"""
    return """<pre id='out'>대기 중</pre>"""


def krxa_hts_desktop_html():
    return """<!doctype html><html><head><meta charset='utf-8'><title>KRXA HTS DESKTOP V33.1.3</title>
<style>
:root{--bg:#07101f;--panel:#121d31;--head:#0d1729;--line:#314563;--btn:#2f6df6;--txt:#eaf2ff;--muted:#9eb2cf}
*{box-sizing:border-box}
body{margin:0;background:linear-gradient(135deg,#07101f,#10234a);color:var(--txt);font-family:Arial,'Noto Sans KR',sans-serif;overflow:hidden}
#top{height:58px;background:#0d1729;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px;padding:8px 12px}
#top b{font-size:18px} .status{border:1px solid #4a77bd;border-radius:999px;padding:4px 10px;font-size:12px;color:#8dffbd}
button,.toplink{background:var(--btn);color:white;border:0;border-radius:8px;padding:8px 12px;margin:2px;cursor:pointer;text-decoration:none;display:inline-block}
button.gray{background:#43536d} button.red{background:#d93636} button.small{padding:4px 8px;font-size:12px}
#desktop{position:relative;width:100vw;height:calc(100vh - 58px);overflow:hidden}
.win{position:absolute;background:var(--panel);border:1px solid var(--line);border-radius:12px;box-shadow:0 10px 30px rgba(0,0,0,.35);min-width:320px;min-height:220px;overflow:hidden}
.win-head{height:42px;background:var(--head);border-bottom:1px solid var(--line);display:flex;align-items:center;justify-content:space-between;padding:0 10px;cursor:move;user-select:none}
.win-title{font-weight:bold}.win-body{height:calc(100% - 42px);display:grid;grid-template-rows:auto 1fr;gap:8px;padding:10px;overflow:auto}
textarea,input,select,pre{background:#040a14;color:var(--txt);border:1px solid #30445f;border-radius:8px;padding:8px;width:100%}
textarea{min-height:70px} pre{white-space:pre-wrap;min-height:120px;max-height:360px;overflow:auto}
table{width:100%;border-collapse:collapse;font-size:12px}td,th{border-bottom:1px solid #2d405c;padding:5px;text-align:left}a{color:#8db7ff}
.toolbar{display:flex;gap:5px;flex-wrap:wrap}.muted{color:var(--muted);font-size:12px}
#dock{position:absolute;left:10px;bottom:10px;display:flex;gap:6px;z-index:9999}
.dockitem{background:#1b2b47;border:1px solid #41618b;border-radius:999px;padding:6px 10px;cursor:pointer}
.resize{position:absolute;right:0;bottom:0;width:16px;height:16px;cursor:nwse-resize}
.llmbar{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:105px;z-index:8}}.llmbar button{padding:7px 12px}}.llmbar .mode{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{color:#fca5a5}}.llmbar .good{color:#86efac}}</style>
<script>
let z=10, wins={}, pollTimer=null;
const DEFAULTS={krxai:[20,20,520,430], db:[560,20,560,430], memory:[1140,20,520,430], task:[20,470,520,390], file:[560,470,560,390], apilog:[1140,470,520,390]};
async function api(path,opt={}){const r=await fetch(path,{headers:{'Content-Type':'application/json'},...opt}); const t=await r.text(); try{return JSON.parse(t)}catch(e){return {ok:false,status:r.status,raw:t}}}
function saveLayout(){const layout={}; document.querySelectorAll('.win').forEach(w=>layout[w.dataset.name]=[w.offsetLeft,w.offsetTop,w.offsetWidth,w.offsetHeight,w.style.display]); localStorage.setItem('krxa_hts_layout',JSON.stringify(layout));}
function loadLayout(){try{return JSON.parse(localStorage.getItem('krxa_hts_layout')||'{}')}catch(e){return {}}}
function openWin(name){let w=document.getElementById('win_'+name); if(w){w.style.display='block'; w.style.zIndex=++z; saveLayout(); return}
 const layout=loadLayout(); const d=layout[name]||DEFAULTS[name]||[60,60,500,400];
 w=document.createElement('div'); w.className='win'; w.id='win_'+name; w.dataset.name=name; w.style.left=d[0]+'px'; w.style.top=d[1]+'px'; w.style.width=d[2]+'px'; w.style.height=d[3]+'px'; w.style.zIndex=++z;
 w.innerHTML=`<div class="win-head"><span class="win-title">${title(name)}</span><span><button class="small gray" onclick="refreshWin('${name}')">↻</button><button class="small gray" onclick="minWin('${name}')">_</button><button class="small red" onclick="closeWin('${name}')">×</button></span></div><div class="win-body">${body(name)}</div><div class="resize"></div>`;
 document.getElementById('desktop').appendChild(w); makeDrag(w); makeResize(w); wins[name]=w; refreshWin(name); saveLayout();
}
function title(n){return {krxai:'KRXAI 논의/보고',db:'DB 편집',memory:'MEMORY 관리',task:'TASK 큐',file:'FILE/REPORT 관리',apilog:'API LOG'}[n]||n}
function body(n){
 if(n==='krxai')return `<textarea id="${n}_input">삼성전자 전략 보고서</textarea><div class="toolbar"><button onclick="runKrxai()">실행</button><button onclick="openWin('file')">파일창</button><button onclick="openWin('memory')">메모리</button></div><pre id="${n}_out">대기</pre>`;
 if(n==='db')return `<select id="db_select"><option>krxai_db</option><option>language_db</option><option>krxai_root</option><option>krxai_memory_index</option><option>auto_web_learning_policy</option><option>report_quality_policy</option><option>window_workspace_policy</option><option>hts_workspace_policy</option></select><div class="toolbar"><button onclick="loadDB()">조회</button><button onclick="saveDB()">저장</button></div><textarea id="db_editor" style="min-height:260px"></textarea><pre id="db_out">대기</pre>`;
 if(n==='memory')return `<textarea id="mem_text">새 기억</textarea><div class="toolbar"><button onclick="addMemory()">추가</button><button onclick="refreshWin('memory')">조회</button></div><div id="memory_list"></div><pre id="memory_out">대기</pre>`;
 if(n==='task')return `<input id="task_title" value="새 작업"><input id="task_action" value="manual_action"><div class="toolbar"><button onclick="addTask()">추가</button><button onclick="tickTask()">대기 실행</button><button onclick="refreshWin('task')">조회</button></div><div id="task_list"></div><pre id="task_out">대기</pre>`;
 if(n==='file')return `<textarea id="file_query">삼성전자 전략 보고서</textarea><div class="toolbar"><button onclick="makeReport()">보고서 생성</button><button onclick="refreshWin('file')">목록</button></div><div id="file_list"></div><pre id="file_out">대기</pre>`;
 if(n==='apilog')return `<div class="toolbar"><button onclick="refreshWin('apilog')">로그 새로고침</button><button onclick="clearApiLog()">화면 비움</button></div><pre id="apilog_out">대기</pre>`;
 return `<pre id="${n}_out">대기</pre>`;
}
function setOut(n,d){const el=document.getElementById(n+'_out'); if(el)el.textContent=JSON.stringify(d,null,2)}
async function refreshWin(n){
 if(n==='krxai')return setOut(n,await api('/api/krxai/thinking/status'));
 if(n==='db')return loadDB();
 if(n==='memory'){const d=await api('/api/memory/user?user_id=default'); renderMemory((d.memory||{}).items||[]); return setOut(n,d)}
 if(n==='task'){const d=await api('/api/tasks/auto/status'); renderTask(((d.queue||{}).items)||[]); return setOut(n,d)}
 if(n==='file'){const d=await api('/api/files/reports'); renderFiles(((d.manifest||{}).files)||[]); return setOut(n,d)}
 if(n==='apilog')return setOut(n,await api('/api/apilog'));
}
async function runKrxai(){const q=document.getElementById('krxai_input').value; const isReport=/보고서|리포트|report|검색|뉴스|전략|시장|분석/i.test(q); const d=isReport?await api('/api/krxa/web/learning/run',{method:'POST',body:JSON.stringify({query:q,user_id:'default'})}):await api('/api/krxai/thinking',{method:'POST',body:JSON.stringify({message:q,user_id:'default'})}); setOut('krxai',d); refreshWin('file'); refreshWin('memory');}
async function loadDB(){const db=document.getElementById('db_select').value; const d=await api('/api/db/'+db); document.getElementById('db_editor').value=JSON.stringify(d.data||d,null,2); setOut('db',d)}
async function saveDB(){let parsed; try{parsed=JSON.parse(document.getElementById('db_editor').value)}catch(e){alert('JSON 오류 '+e);return} const d=await api('/api/db/save',{method:'POST',body:JSON.stringify({db:document.getElementById('db_select').value,data:parsed.data||parsed})}); setOut('db',d)}
async function addMemory(){const text=document.getElementById('mem_text').value; const d=await api('/api/memory/add',{method:'POST',body:JSON.stringify({user_id:'default',entry:{input:text,final:text,role:'manual_memory'}})}); setOut('memory',d); refreshWin('memory')}
function renderMemory(items){const el=document.getElementById('memory_list'); if(!el)return; el.innerHTML=`<table><tr><th>ID</th><th>내용</th><th>관리</th></tr>${items.map(x=>`<tr><td>${x.id||''}</td><td>${((x.final||x.input||'')+'').slice(0,80)}</td><td><button class="small red" onclick="delMemory('${x.id}')">삭제</button></td></tr>`).join('')}</table>`}
async function delMemory(id){setOut('memory',await api('/api/memory/delete',{method:'POST',body:JSON.stringify({user_id:'default',id})})); refreshWin('memory')}
async function addTask(){const d=await api('/api/tasks/add',{method:'POST',body:JSON.stringify({title:document.getElementById('task_title').value,action:document.getElementById('task_action').value,user_id:'default'})}); setOut('task',d); refreshWin('task')}
async function tickTask(){setOut('task',await api('/api/tasks/auto/tick',{method:'POST',body:JSON.stringify({user_id:'default'})})); refreshWin('task')}
function renderTask(items){const el=document.getElementById('task_list'); if(!el)return; el.innerHTML=`<table><tr><th>ID</th><th>상태</th><th>작업</th><th>관리</th></tr>${items.map(t=>`<tr><td>${t.id}</td><td>${t.status}</td><td>${t.title}</td><td><button class="small red" onclick="delTask('${t.id}')">삭제</button></td></tr>`).join('')}</table>`}
async function delTask(id){setOut('task',await api('/api/tasks/delete',{method:'POST',body:JSON.stringify({id})})); refreshWin('task')}
async function makeReport(){const q=document.getElementById('file_query').value; const d=await api('/api/krxa/web/learning/run',{method:'POST',body:JSON.stringify({query:q,user_id:'default'})}); setOut('file',d); refreshWin('file')}
function renderFiles(files){const el=document.getElementById('file_list'); if(!el)return; el.innerHTML=`<table><tr><th>파일</th><th>크기</th><th>다운로드</th><th>관리</th></tr>${files.map(f=>{const n=f.name||f.stored_name||f.original_name;return `<tr><td>${n}</td><td>${f.size||0}</td><td><a target="_blank" href="${f.download_url}">다운로드</a></td><td><button class="small red" onclick="delFile('${n}')">삭제</button></td></tr>`}).join('')}</table>`}
async function delFile(n){if(confirm('삭제 '+n+'?')){setOut('file',await api('/api/files/delete',{method:'POST',body:JSON.stringify({filename:n})})); refreshWin('file')}}
function clearApiLog(){document.getElementById('apilog_out').textContent=''}
function makeDrag(w){const h=w.querySelector('.win-head'); let ox=0,oy=0,down=false; h.onmousedown=e=>{down=true; w.style.zIndex=++z; ox=e.clientX-w.offsetLeft; oy=e.clientY-w.offsetTop}; window.addEventListener('mousemove',e=>{if(!down)return; w.style.left=(e.clientX-ox)+'px'; w.style.top=(e.clientY-oy)+'px'}); window.addEventListener('mouseup',()=>{if(down){down=false;saveLayout()}})}
function makeResize(w){const r=w.querySelector('.resize'); let down=false,ow=0,oh=0,ox=0,oy=0; r.onmousedown=e=>{e.stopPropagation();down=true;ow=w.offsetWidth;oh=w.offsetHeight;ox=e.clientX;oy=e.clientY}; window.addEventListener('mousemove',e=>{if(!down)return; w.style.width=Math.max(320,ow+e.clientX-ox)+'px'; w.style.height=Math.max(220,oh+e.clientY-oy)+'px'}); window.addEventListener('mouseup',()=>{if(down){down=false;saveLayout()}})}
function minWin(n){document.getElementById('win_'+n).style.display='none'; addDock(n); saveLayout()}
function closeWin(n){const w=document.getElementById('win_'+n); if(w)w.remove(); delete wins[n]; saveLayout()}
function addDock(n){const d=document.getElementById('dock'); if(document.getElementById('dock_'+n))return; const x=document.createElement('div'); x.className='dockitem'; x.id='dock_'+n; x.textContent=title(n); x.onclick=()=>{openWin(n); x.remove()}; d.appendChild(x)}
function startPoll(){if(pollTimer)clearInterval(pollTimer); pollTimer=setInterval(()=>{document.querySelectorAll('.win').forEach(w=>{const n=w.dataset.name;if(['file','memory','task','apilog'].includes(n)&&w.style.display!=='none')refreshWin(n)})},15000)}
window.onload=()=>{['krxai','db','memory','task','file','apilog'].forEach(openWin); startPoll()}
</script></head><body>
<div id="top"><b>KRXA HTS MULTI-WINDOW DESKTOP V33.1.3</b><span class="status">SERVER READY</span><a class="toplink" href="/window/krxai" target="_blank">KRXAI 새창</a><a class="toplink" href="/window/db" target="_blank">DB 새창</a><a class="toplink" href="/window/memory" target="_blank">MEMORY 새창</a><a class="toplink" href="/window/task" target="_blank">TASK 새창</a><a class="toplink" href="/window/file" target="_blank">FILE 새창</a><a class="toplink" href="/window/apilog" target="_blank">API LOG 새창</a><a class="toplink" href="/window/krxai" target="_blank">KRXAI</a><a class="toplink" href="/window/db" target="_blank">DB</a><a class="toplink" href="/window/memory" target="_blank">MEMORY</a><a class="toplink" href="/window/task" target="_blank">TASK</a><a class="toplink" href="/window/file" target="_blank">FILE</a><a class="toplink" href="/window/apilog" target="_blank">API LOG</a><button class="gray" onclick="openWin('krxai');openWin('db');openWin('memory');openWin('task');openWin('file');openWin('apilog')">내부창 전체</button><button class="gray" onclick="localStorage.removeItem('krxa_hts_layout');location.reload()">레이아웃 초기화</button><a href="/control" style="color:#8db7ff">CONTROL</a></div>
<div id="desktop"><div id="dock"></div></div>
</body></html>"""


def krxa_realtime_state_path():
    return CORE_DIR / 'krxa_realtime_state.json'

def krxa_autonomous_state_path():
    return CORE_DIR / 'krxa_autonomous_state.json'

def krxa_event_emit(event_type, payload=None, source='system'):
    """V33.1.3 internal event bus. Stored in JSON for polling/SSE fallback."""
    payload = payload or {}
    p = krxa_realtime_state_path()
    state = load_json(p, {'schema': 'KRXA_REALTIME_STATE_V33.1.3', 'events': []})
    ev = {
        'id': 'evt_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'type': event_type,
        'source': source,
        'payload': payload,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    events = state.setdefault('events', [])
    events.append(ev)
    state['events'] = events[-300:]
    state['last_event'] = ev
    state['status'] = 'ACTIVE'
    state['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(p, state)
    try:
        krxa_api_log_add('EVENT', event_type, 'emitted', source)
    except Exception:
        pass
    return ev

def krxa_realtime_status():
    return {
        'ok': True,
        'policy': load_json(CORE_DIR / 'krxa_realtime_autonomous_policy.json', {}),
        'realtime': load_json(krxa_realtime_state_path(), {}),
        'autonomous': load_json(krxa_autonomous_state_path(), {})
    }

def krxa_sse_events_text():
    state = load_json(krxa_realtime_state_path(), {'events': []})
    lines = []
    for ev in state.get('events', [])[-30:]:
        lines.append('event: ' + ev.get('type', 'message'))
        lines.append('data: ' + json.dumps(ev, ensure_ascii=False))
        lines.append('')
    return '\n'.join(lines) + '\n'

def krxa_autonomous_safe_tick(user_id='default'):
    """V33.1.3 safe autonomous tick: executes one eligible task and records event. No infinite loop."""
    uid = krxa_safe_user_id(user_id)
    state_path = krxa_autonomous_state_path()
    auto = load_json(state_path, {'tick_count': 0, 'runs': [], 'enabled': True})
    if not auto.get('enabled', True):
        return {'ok': False, 'reason': 'autonomous_disabled', 'autonomous': auto}

    result = krxa_auto_task_tick(uid)
    run = {
        'id': 'autotick_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'user_id': uid,
        'executed_count': len(result.get('executed', [])),
        'executed': result.get('executed', []),
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    auto['tick_count'] = int(auto.get('tick_count', 0)) + 1
    auto['last_tick'] = run
    auto.setdefault('runs', []).append(run)
    auto['runs'] = auto.get('runs', [])[-200:]
    auto['status'] = 'TICKED'
    auto['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(state_path, auto)
    krxa_event_emit('TASK_EXECUTE', run, source='autonomous_tick')
    return {'ok': True, 'schema': 'KRXA_AUTONOMOUS_TICK_RESPONSE_V33.1.3', 'run': run, 'autonomous': auto, 'task_result': result}

def krxa_realtime_desktop_html():
    """V33.1.3 enhanced desktop with event ticker and autonomous controls."""
    base = krxa_hts_desktop_html()
    inject = """
<script>
async function rtApi(path,opt={}){const r=await fetch(path,{headers:{'Content-Type':'application/json'},...opt});const t=await r.text();try{return JSON.parse(t)}catch(e){return {raw:t,status:r.status}}}
async function realtimeStatus(){
 const d=await rtApi('/api/realtime/status');
 const el=document.getElementById('rtbar'); if(el) el.textContent='EVENT: '+((d.realtime||{}).last_event||{}).type+' | AUTO TICKS: '+((d.autonomous||{}).tick_count||0);
}
async function autonomousTick(){
 const d=await rtApi('/api/autonomous/tick',{method:'POST',body:JSON.stringify({user_id:'default'})});
 alert('AUTO TICK 실행: '+JSON.stringify(d.run||d));
 realtimeStatus();
 if(typeof refreshWin==='function'){refreshWin('task');refreshWin('memory');refreshWin('apilog');}
}
setInterval(realtimeStatus,5000);
window.addEventListener('load',()=>{setTimeout(()=>{const top=document.getElementById('top'); if(top&&!document.getElementById('rtbar')){const s=document.createElement('span');s.id='rtbar';s.className='status';s.textContent='EVENT READY';top.appendChild(s); const b=document.createElement('button');b.textContent='AUTO TICK';b.onclick=autonomousTick;top.appendChild(b);} realtimeStatus();},500);});
</script>
"""
    return base.replace("</body></html>", inject + "</body></html>")


def krxa_safe_files_reports_response():
    try:
        manifest = krxa_manifest_load_sorted()
    except Exception:
        manifest_path = ROOT / 'storage' / 'file_exchange' / 'file_manifest.json'
        manifest_path.parent.mkdir(parents=True, exist_ok=True)
        manifest = load_json(manifest_path, {'schema':'KRXA_FILE_EXCHANGE_MANIFEST_V1','files':[]})
        manifest['files'] = sorted(manifest.get('files', []), key=lambda x: x.get('created_at',''), reverse=True)
    return {'ok': True, 'manifest': manifest}

def krxa_newwindow_launcher_js():
    return """
function openTaskWindow(name){
  const map={krxai:'/window/krxai',db:'/window/db',memory:'/window/memory',task:'/window/task',file:'/window/file',apilog:'/window/apilog'};
  window.open(map[name]||('/window/'+name),'_blank','width=1100,height=760,menubar=no,toolbar=no,location=yes,resizable=yes,scrollbars=yes');
}
"""


def krxa_chat_history_path():
    return CORE_DIR / 'krxa_chat_history.json'

def krxa_chat_history_add(user_id, role, message, meta=None):
    meta = meta or {}
    p = krxa_chat_history_path()
    hist = load_json(p, {'schema': 'KRXA_CHAT_HISTORY_V33.1.3', 'items': []})
    item = {
        'id': 'chat_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'user_id': krxa_safe_user_id(user_id),
        'role': role,
        'message': message,
        'meta': meta,
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    hist.setdefault('items', []).append(item)
    hist['items'] = hist.get('items', [])[-500:]
    hist['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(p, hist)
    return item

def krxa_chat_history_load(user_id='default'):
    uid = krxa_safe_user_id(user_id)
    hist = load_json(krxa_chat_history_path(), {'items': []})
    return {'schema': hist.get('schema', 'KRXA_CHAT_HISTORY_V33.1.3'), 'items': [x for x in hist.get('items', []) if x.get('user_id') == uid][-100:]}

def krxa_chat_classify_intent(message):
    m = (message or '').lower()
    if any(k in m for k in ['보고서', '리포트', 'report', '전략', '시장', '뉴스', '검색', '분석']):
        return 'report_or_research'
    if any(k in m for k in ['파일', '다운로드', 'pdf', 'docx']):
        return 'file'
    if any(k in m for k in ['작업', 'task', '자동', '실행']):
        return 'task'
    if any(k in m for k in ['db', '데이터', '수정', '저장']):
        return 'db'
    if any(k in m for k in ['기억', 'memory', '학습']):
        return 'memory'
    return 'discussion'

def krxa_chat_run(user_id='default', message=''):
    """V33.1.3: ChatGPT-style command -> KRXA operating flow."""
    uid = krxa_safe_user_id(user_id)
    message = (message or '').strip()
    intent = krxa_chat_classify_intent(message)
    user_msg = krxa_chat_history_add(uid, 'user', message, {'intent': intent})

    thinking = krxai_think_structured(message, role='chat_integrated_ui') if 'krxai_think_structured' in globals() else {
        'ok': True,
        'response': {'reason': 'fallback thinking', 'memory_used': [], 'final': message},
        'intent': intent
    }

    task = {
        'id': 'chat_task_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'user_id': uid,
        'status': 'created',
        'type': intent,
        'title': message[:120] or '말대말 작업',
        'action': 'chat_integrated_run',
        'requires_approval': False,
        'judgement': thinking.get('response', thinking),
        'created_at': datetime.now(timezone.utc).isoformat()
    }

    # Add task to queue
    try:
        krxa_task_add(task)
    except Exception:
        pass

    result = {'status': 'created', 'final': thinking.get('response', {}).get('final') if isinstance(thinking.get('response'), dict) else str(thinking)}

    # If report/search intent, execute web learning/report flow immediately.
    if intent == 'report_or_research':
        try:
            web_result = krxa_auto_web_learning_run(message, uid) if 'krxa_auto_web_learning_run' in globals() else {'ok': False, 'error': 'web_learning_not_available'}
            result = {
                'status': web_result.get('run', {}).get('status', 'executed'),
                'final': web_result.get('run', {}).get('summary') or '보고서/검색 작업을 실행했습니다.',
                'web_learning': web_result
            }
            task['status'] = 'executed'
        except Exception as e:
            result = {'status': 'error', 'final': str(e)}

    elif intent == 'task':
        try:
            tick = krxa_autonomous_safe_tick(uid) if 'krxa_autonomous_safe_tick' in globals() else krxa_auto_task_tick(uid)
            result = {'status': 'task_tick', 'final': '자동 실행 tick을 수행했습니다.', 'tick': tick}
            task['status'] = 'executed'
        except Exception as e:
            result = {'status': 'error', 'final': str(e)}

    # Store memory seed
    try:
        krxa_user_memory_store(uid, {
            'role': 'chat_integrated_ui',
            'input': message,
            'intent': intent,
            'reason': thinking.get('response', {}).get('reason') if isinstance(thinking.get('response'), dict) else '',
            'final': result.get('final'),
            'task_id': task['id'],
            'engine': 'KRXA_CHATGPT_INTEGRATED_OPERATING_UI_V33.1.3'
        })
    except Exception:
        pass

    # Event
    try:
        krxa_event_emit('CHAT_RUN_COMPLETE', {'user_id': uid, 'intent': intent, 'task_id': task['id'], 'status': result.get('status')}, source='chat_ui')
    except Exception:
        pass

    assistant_msg = krxa_chat_history_add(uid, 'assistant', result.get('final', ''), {'intent': intent, 'task': task, 'result': result})

    return {
        'ok': True,
        'schema': 'KRXA_CHAT_RUN_RESPONSE_V33.1.3',
        'user_id': uid,
        'intent': intent,
        'reason': thinking.get('response', {}).get('reason') if isinstance(thinking.get('response'), dict) else '',
        'memory_used': thinking.get('response', {}).get('memory_used') if isinstance(thinking.get('response'), dict) else [],
        'final': result.get('final'),
        'task': task,
        'result': result,
        'chat': {'user': user_msg, 'assistant': assistant_msg}
    }

def krxa_chat_integrated_ui_html():
    return """<!doctype html><html><head><meta charset='utf-8'><title>KRXA + ChatGPT Integrated UI V33.1.3</title>
<style>
:root{--bg:#07101f;--panel:#121d31;--head:#0d1729;--line:#314563;--btn:#2f6df6;--txt:#eaf2ff;--muted:#9eb2cf}
*{box-sizing:border-box} body{margin:0;background:linear-gradient(135deg,#07101f,#10234a);color:var(--txt);font-family:Arial,'Noto Sans KR',sans-serif;height:100vh;overflow:hidden}
header{height:56px;background:#0d1729;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:10px;padding:8px 12px}
header b{font-size:17px}.pill{border:1px solid #4a77bd;border-radius:999px;padding:4px 10px;font-size:12px;color:#8dffbd}
a.btn,button{background:var(--btn);color:white;border:0;border-radius:8px;padding:8px 12px;margin:2px;cursor:pointer;text-decoration:none;display:inline-block}
button.gray,a.gray{background:#43536d}.grid{display:grid;grid-template-columns:360px 1fr 380px;grid-template-rows:1fr 230px;height:calc(100vh - 56px);gap:10px;padding:10px}
.panel{background:rgba(18,29,49,.92);border:1px solid var(--line);border-radius:14px;overflow:hidden;min-height:0}.panel h3{margin:0;background:#101a2d;border-bottom:1px solid var(--line);padding:10px 12px}.body{padding:10px;height:calc(100% - 42px);overflow:auto}
#chatlog{height:calc(100% - 150px);overflow:auto;padding-right:4px}.msg{padding:10px;border-radius:12px;margin-bottom:8px;font-size:13px;white-space:pre-wrap}.user{background:#1b3b77}.assistant{background:#17283e}
textarea,input,select,pre{background:#040a14;color:var(--txt);border:1px solid #30445f;border-radius:8px;padding:8px;width:100%}textarea{min-height:78px}pre{white-space:pre-wrap;min-height:150px;max-height:100%;overflow:auto}
table{width:100%;border-collapse:collapse;font-size:12px}td,th{border-bottom:1px solid #2d405c;padding:5px;text-align:left}a{color:#8db7ff}.small{font-size:12px;color:var(--muted)}
.llmbar{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:105px;z-index:8}}.llmbar button{padding:7px 12px}}.llmbar .mode{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{color:#fca5a5}}.llmbar .good{color:#86efac}}</style>
<script>
async function api(path,opt={}){const r=await fetch(path,{headers:{'Content-Type':'application/json'},...opt});const t=await r.text();try{return JSON.parse(t)}catch(e){return {ok:false,status:r.status,raw:t}}}
function set(id,x){document.getElementById(id).textContent=typeof x==='string'?x:JSON.stringify(x,null,2)}
function addMsg(role,msg){const d=document.createElement('div');d.className='msg '+(role==='user'?'user':'assistant');d.textContent=msg;document.getElementById('chatlog').appendChild(d);document.getElementById('chatlog').scrollTop=999999}
async function runChat(){
 const q=document.getElementById('chatInput').value.trim(); if(!q)return;
 addMsg('user',q); document.getElementById('chatInput').value='';
 const d=await api('/api/chat/run',{method:'POST',body:JSON.stringify({user_id:'default',message:q})});
 addMsg('assistant',d.final||JSON.stringify(d));
 set('thinking',d); refreshAll();
}
async function refreshAll(){loadHistory();loadTasks();loadFiles();loadMemory();loadRealtime();}
async function loadHistory(){const d=await api('/api/chat/history?user_id=default'); const box=document.getElementById('chatlog'); box.innerHTML=''; (d.history?.items||[]).forEach(x=>addMsg(x.role,x.message));}
async function loadTasks(){const d=await api('/api/tasks/auto/status'); const items=((d.queue||{}).items||[]); document.getElementById('tasklist').innerHTML='<table><tr><th>상태</th><th>작업</th></tr>'+items.map(t=>`<tr><td>${t.status}</td><td>${t.title}</td></tr>`).join('')+'</table>'; set('taskraw',d);}
async function loadFiles(){const d=await api('/api/files/reports'); const files=((d.manifest||{}).files||[]); document.getElementById('filelist').innerHTML='<table><tr><th>파일</th><th>다운로드</th></tr>'+files.map(f=>{const n=f.name||f.stored_name||f.original_name;return `<tr><td>${n}</td><td><a target="_blank" href="${f.download_url}">다운로드</a></td></tr>`}).join('')+'</table>'; set('fileraw',d);}
async function loadMemory(){const d=await api('/api/memory/user?user_id=default'); set('memoryraw',d);}
async function loadRealtime(){const d=await api('/api/realtime/status'); set('realtime',d);}
async function autoTick(){const d=await api('/api/autonomous/tick',{method:'POST',body:JSON.stringify({user_id:'default'})}); set('realtime',d); refreshAll();}
window.onload=()=>{refreshAll();setInterval(refreshAll,10000)}
</script></head><body>
<header><b>KRXA + ChatGPT Integrated Operating UI V33.1.3</b><span class="pill">CHAT → TASK → MEMORY → FILE</span><a class="btn" href="/desktop" target="_blank">HTS</a><a class="btn" href="/window/file" target="_blank">FILE</a><a class="btn" href="/window/db" target="_blank">DB</a><button class="gray" onclick="autoTick()">AUTO TICK</button></header>
<div class="grid">
<section class="panel" style="grid-row:1 / span 2"><h3>말대말 CHAT</h3><div class="body"><div id="chatlog"></div><textarea id="chatInput" placeholder="예: 삼성전자 전략 보고서 작성"></textarea><button onclick="runChat()">실행</button><button class="gray" onclick="loadHistory()">기록</button></div></section>
<section class="panel"><h3>KRXAI THINKING / EXECUTION</h3><div class="body"><pre id="thinking">대기 중</pre></div></section>
<section class="panel"><h3>REPORT / FILE</h3><div class="body"><div id="filelist"></div><pre id="fileraw"></pre></div></section>
<section class="panel"><h3>TASK QUEUE</h3><div class="body"><div id="tasklist"></div><pre id="taskraw"></pre></div></section>
<section class="panel"><h3>MEMORY / REALTIME</h3><div class="body"><pre id="memoryraw"></pre><pre id="realtime"></pre></div></section>
</div></body></html>"""


def krxa_v323_safe_queue_path():
    return CORE_DIR / 'krxai_task_queue.json'

def krxa_v323_queue_load():
    return load_json(krxa_v323_safe_queue_path(), {'schema':'KRXA_TASK_QUEUE_V33.1_3','items':[],'runs':[]})

def krxa_v323_queue_save(q):
    q['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(krxa_v323_safe_queue_path(), q)

def krxa_v323_task_add(title='새 작업', action='manual_action', user_id='default', task_type='manual'):
    q = krxa_v323_queue_load()
    item = {
        'id': 'task_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'),
        'user_id': krxa_safe_user_id(user_id),
        'status': 'pending',
        'type': task_type,
        'title': title or '새 작업',
        'action': action or 'manual_action',
        'created_at': datetime.now(timezone.utc).isoformat()
    }
    q.setdefault('items', []).append(item)
    krxa_v323_queue_save(q)
    try: krxa_event_emit('TASK_ADD', item, source='v32_3_control')
    except Exception: pass
    return {'ok': True, 'task': item, 'queue': q}

def krxa_v323_task_delete(task_id):
    q = krxa_v323_queue_load()
    before = len(q.get('items', []))
    q['items'] = [x for x in q.get('items', []) if x.get('id') != task_id]
    krxa_v323_queue_save(q)
    try: krxa_event_emit('TASK_DELETE', {'id': task_id}, source='v32_3_control')
    except Exception: pass
    return {'ok': True, 'deleted': before - len(q.get('items', [])), 'queue': q}

def krxa_v323_task_tick(user_id='default'):
    uid = krxa_safe_user_id(user_id)
    q = krxa_v323_queue_load()
    executed = []
    for item in q.get('items', []):
        if item.get('status') == 'pending' and item.get('user_id', uid) == uid:
            item['status'] = 'done'
            item['executed_at'] = datetime.now(timezone.utc).isoformat()
            executed.append(item)
            try:
                krxa_user_memory_store(uid, {'role':'task_execution','input':item.get('title'), 'final':'TASK 실행 완료: '+item.get('title',''), 'task_id':item.get('id')})
            except Exception: pass
            break
    run = {'id':'run_'+datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'), 'user_id':uid, 'executed':executed, 'created_at':datetime.now(timezone.utc).isoformat()}
    q.setdefault('runs', []).append(run)
    q['runs'] = q.get('runs', [])[-200:]
    krxa_v323_queue_save(q)
    try: krxa_event_emit('TASK_EXECUTE', run, source='v32_3_tick')
    except Exception: pass
    return {'ok': True, 'schema':'KRXA_TASK_TICK_V33.1_3', 'executed': executed, 'run': run, 'queue': q}

def krxa_v323_memory_add(user_id='default', text=''):
    saved = krxa_user_memory_store(user_id, {'role':'manual_memory','input':text,'final':text,'engine':'V33.1_3_CONTROL'})
    try: krxa_event_emit('MEMORY_ADD', {'user_id': user_id, 'text': text[:80]}, source='v32_3_control')
    except Exception: pass
    return {'ok': True, 'saved': saved, 'memory': krxa_user_memory_load(user_id)}

def krxa_v323_memory_delete(user_id='default', item_id=''):
    uid = krxa_safe_user_id(user_id)
    p = krxa_user_memory_path(uid)
    mem = load_json(p, {'schema':'KRXA_USER_MEMORY_V31','user_id':uid,'items':[]})
    before = len(mem.get('items', []))
    mem['items'] = [x for x in mem.get('items', []) if x.get('id') != item_id]
    mem['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(p, mem)
    try: krxa_event_emit('MEMORY_DELETE', {'user_id': uid, 'id': item_id}, source='v32_3_control')
    except Exception: pass
    return {'ok': True, 'deleted': before - len(mem.get('items', [])), 'memory': mem}

def krxa_v323_manifest():
    try:
        return krxa_manifest_load_sorted()
    except Exception:
        p = ROOT / 'storage' / 'file_exchange' / 'file_manifest.json'
        p.parent.mkdir(parents=True, exist_ok=True)
        return load_json(p, {'schema':'KRXA_FILE_EXCHANGE_MANIFEST_V1','files':[]})

def krxa_v323_file_delete(filename):
    safe = Path(filename or '').name
    if not safe:
        return {'ok': False, 'error': 'empty_filename'}
    try:
        return krxa_delete_report_file(safe)
    except Exception:
        p = ROOT / 'storage' / 'file_exchange' / 'file_manifest.json'
        manifest = load_json(p, {'files': []})
        before = len(manifest.get('files', []))
        manifest['files'] = [f for f in manifest.get('files', []) if (f.get('name') or f.get('stored_name') or f.get('original_name')) != safe]
        save_json(p, manifest)
        target = ROOT / 'files' / 'reports' / safe
        existed = target.exists()
        if existed:
            target.unlink()
        return {'ok': True, 'filename': safe, 'file_existed': existed, 'manifest_removed': before - len(manifest.get('files', [])), 'manifest': manifest}

def krxa_v323_db_save(db, data):
    if db not in DB_FILES:
        return {'ok': False, 'error': 'unknown_db', 'db': db}
    save_json(DB_FILES[db], data)
    try: krxa_event_emit('DB_SAVE', {'db': db}, source='v32_3_control')
    except Exception: pass
    return {'ok': True, 'db': db, 'saved_at': datetime.now(timezone.utc).isoformat()}

def krxa_v323_apilog():
    try:
        return load_json(krxa_api_log_path(), {'items': []})
    except Exception:
        return {'schema':'KRXA_API_LOG_FALLBACK_V33.1_3','items':[]}

def krxa_v323_realtime():
    try:
        return krxa_realtime_status()
    except Exception:
        return {'ok': True, 'realtime': {'status':'READY','events':[]}, 'autonomous': {'status':'READY'}}

def krxa_v323_full_control_html():
    return """<!doctype html><html><head><meta charset='utf-8'><title>KRXA Full Operating Control UI V33.1.3</title>
<style>
:root{--bg:#07101f;--panel:#121d31;--head:#0d1729;--line:#314563;--btn:#2f6df6;--txt:#eaf2ff;--muted:#9eb2cf}
*{box-sizing:border-box} body{margin:0;background:linear-gradient(135deg,#07101f,#10234a);color:var(--txt);font-family:Arial,'Noto Sans KR',sans-serif;height:100vh;overflow:hidden}
header{height:56px;background:#0d1729;border-bottom:1px solid var(--line);display:flex;align-items:center;gap:8px;padding:8px 12px}
header b{font-size:17px}.pill{border:1px solid #4a77bd;border-radius:999px;padding:4px 10px;font-size:12px;color:#8dffbd}
button,a.btn{background:var(--btn);color:white;border:0;border-radius:8px;padding:8px 12px;margin:2px;cursor:pointer;text-decoration:none;display:inline-block}
button.gray,a.gray{background:#43536d}button.red{background:#d93636}.grid{display:grid;grid-template-columns:350px 1fr 380px;grid-template-rows:1fr 260px;height:calc(100vh - 56px);gap:10px;padding:10px}
.panel{background:rgba(18,29,49,.94);border:1px solid var(--line);border-radius:14px;overflow:hidden;min-height:0}.panel h3{margin:0;background:#101a2d;border-bottom:1px solid var(--line);padding:10px 12px}.body{padding:10px;height:calc(100% - 42px);overflow:auto}
textarea,input,select,pre{background:#040a14;color:var(--txt);border:1px solid #30445f;border-radius:8px;padding:8px;width:100%}textarea{min-height:72px}pre{white-space:pre-wrap;min-height:120px;overflow:auto}
table{width:100%;border-collapse:collapse;font-size:12px}td,th{border-bottom:1px solid #2d405c;padding:5px;text-align:left}a{color:#8db7ff}.msg{padding:9px;border-radius:10px;margin:6px 0;font-size:13px}.user{background:#1b3b77}.assistant{background:#17283e}
.llmbar{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:105px;z-index:8}}.llmbar button{padding:7px 12px}}.llmbar .mode{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{color:#fca5a5}}.llmbar .good{color:#86efac}}</style>
<script>
async function api(path,opt={}){const r=await fetch(path,{headers:{'Content-Type':'application/json'},...opt});const t=await r.text();try{return JSON.parse(t)}catch(e){return {ok:false,status:r.status,raw:t}}}
function out(id,x){document.getElementById(id).textContent=typeof x==='string'?x:JSON.stringify(x,null,2)}
function msg(role,text){const d=document.createElement('div');d.className='msg '+(role==='user'?'user':'assistant');d.textContent=text;document.getElementById('chatlog').appendChild(d);document.getElementById('chatlog').scrollTop=99999}
async function chatRun(){const q=document.getElementById('chatInput').value.trim();if(!q)return;msg('user',q);document.getElementById('chatInput').value='';const d=await api('/api/chat/run',{method:'POST',body:JSON.stringify({user_id:'default',message:q})});msg('assistant',d.final||JSON.stringify(d));out('thinkOut',d);refreshAll()}
async function chatHistory(){const d=await api('/api/chat/history?user_id=default');document.getElementById('chatlog').innerHTML='';(d.history?.items||[]).forEach(x=>msg(x.role,x.message));}
async function loadDB(){const db=document.getElementById('dbSelect').value;const d=await api('/api/db/'+db);document.getElementById('dbEditor').value=JSON.stringify(d.data||d,null,2);out('dbOut',d)}
async function saveDB(){let data;try{data=JSON.parse(document.getElementById('dbEditor').value)}catch(e){alert('JSON 오류 '+e);return}out('dbOut',await api('/api/db/save',{method:'POST',body:JSON.stringify({db:document.getElementById('dbSelect').value,data:data.data||data})}));}
async function memLoad(){const d=await api('/api/memory/user?user_id=default');const items=(d.memory?.items||[]);document.getElementById('memList').innerHTML='<table><tr><th>ID</th><th>내용</th><th>관리</th></tr>'+items.map(x=>`<tr><td>${x.id||''}</td><td>${((x.final||x.input||'')+'').slice(0,90)}</td><td><button class="red" onclick="memDel('${x.id}')">삭제</button></td></tr>`).join('')+'</table>';out('memOut',d)}
async function memAdd(){const text=document.getElementById('memInput').value;out('memOut',await api('/api/memory/add',{method:'POST',body:JSON.stringify({user_id:'default',text})}));memLoad()}
async function memDel(id){out('memOut',await api('/api/memory/delete',{method:'POST',body:JSON.stringify({user_id:'default',id})}));memLoad()}
async function taskLoad(){const d=await api('/api/tasks/auto/status');const items=((d.queue||{}).items||[]);document.getElementById('taskList').innerHTML='<table><tr><th>상태</th><th>작업</th><th>관리</th></tr>'+items.map(t=>`<tr><td>${t.status}</td><td>${t.title}</td><td><button class="red" onclick="taskDel('${t.id}')">삭제</button></td></tr>`).join('')+'</table>';out('taskOut',d)}
async function taskAdd(){out('taskOut',await api('/api/tasks/add',{method:'POST',body:JSON.stringify({user_id:'default',title:document.getElementById('taskTitle').value,action:document.getElementById('taskAction').value})}));taskLoad()}
async function taskTick(){out('taskOut',await api('/api/autonomous/tick',{method:'POST',body:JSON.stringify({user_id:'default'})}));taskLoad();memLoad();rtLoad()}
async function taskDel(id){out('taskOut',await api('/api/tasks/delete',{method:'POST',body:JSON.stringify({id})}));taskLoad()}
async function fileLoad(){const d=await api('/api/files/reports');const files=((d.manifest||{}).files||[]);document.getElementById('fileList').innerHTML='<table><tr><th>파일</th><th>크기</th><th>다운로드</th><th>관리</th></tr>'+files.map(f=>{const n=f.name||f.stored_name||f.original_name;return `<tr><td>${n}</td><td>${f.size||0}</td><td><a target="_blank" href="${f.download_url}">다운로드</a></td><td><button class="red" onclick="fileDel('${n}')">삭제</button></td></tr>`}).join('')+'</table>';out('fileOut',d)}
async function fileMake(){out('fileOut',await api('/api/krxa/web/learning/run',{method:'POST',body:JSON.stringify({user_id:'default',query:document.getElementById('reportQuery').value})}));fileLoad()}
async function fileDel(n){if(confirm('삭제 '+n+'?')){out('fileOut',await api('/api/files/delete',{method:'POST',body:JSON.stringify({filename:n})}));fileLoad()}}
async function rtLoad(){out('rtOut',await api('/api/realtime/status'));out('logOut',await api('/api/apilog'))}
async function refreshAll(){chatHistory();taskLoad();fileLoad();memLoad();rtLoad();loadDB()}
window.onload=()=>{refreshAll();setInterval(()=>{taskLoad();fileLoad();rtLoad()},10000)}
</script></head><body>
<header><b>KRXA FULL OPERATING CONTROL UI V33.1.3</b><span class="pill">ALL FUNCTIONS OPERABLE</span><a class="btn" href="/chat" target="_blank">CHAT</a><a class="btn" href="/HTS" target="_blank">HTS</a><button class="gray" onclick="refreshAll()">전체 새로고침</button><button onclick="taskTick()">AUTO TICK</button></header>
<div class="grid">
<section class="panel" style="grid-row:1 / span 2"><h3>말대말 CHAT</h3><div class="body"><div id="chatlog"></div><textarea id="chatInput" placeholder="삼성전자 전략 보고서 작성"></textarea><button onclick="chatRun()">실행</button><button class="gray" onclick="chatHistory()">기록</button><hr><h4>MEMORY 관리</h4><textarea id="memInput">새 기억</textarea><button onclick="memAdd()">기억 추가</button><button class="gray" onclick="memLoad()">조회</button><div id="memList"></div><pre id="memOut"></pre></div></section>
<section class="panel"><h3>KRXAI THINKING / TASK</h3><div class="body"><pre id="thinkOut">대기</pre><h4>TASK</h4><input id="taskTitle" value="새 작업"><input id="taskAction" value="manual_action"><button onclick="taskAdd()">작업 추가</button><button onclick="taskTick()">대기 실행</button><button class="gray" onclick="taskLoad()">조회</button><div id="taskList"></div><pre id="taskOut"></pre></div></section>
<section class="panel"><h3>REPORT / FILE 관리</h3><div class="body"><textarea id="reportQuery">삼성전자 전략 보고서</textarea><button onclick="fileMake()">보고서 생성</button><button class="gray" onclick="fileLoad()">목록</button><div id="fileList"></div><pre id="fileOut"></pre></div></section>
<section class="panel"><h3>DB EDITOR</h3><div class="body"><select id="dbSelect"><option>krxai_db</option><option>language_db</option><option>krxai_root</option><option>chat_history</option><option>full_operating_control_ui_policy</option></select><button onclick="loadDB()">조회</button><button onclick="saveDB()">저장</button><textarea id="dbEditor" style="min-height:140px"></textarea><pre id="dbOut"></pre></div></section>
<section class="panel"><h3>REALTIME / API LOG</h3><div class="body"><button onclick="rtLoad()">로그 새로고침</button><pre id="rtOut"></pre><pre id="logOut"></pre></div></section>
</div></body></html>"""


def krxa_v324_shell(title, role, body):
    return f"""<!doctype html><html><head><meta charset='utf-8'><title>{title}</title>
<style>
body{{margin:0;background:#07101f;color:#eaf2ff;font-family:Arial,'Noto Sans KR',sans-serif}}
header{{position:sticky;top:0;z-index:10;background:#0d1729;border-bottom:1px solid #314563;padding:10px;display:flex;gap:8px;align-items:center;flex-wrap:wrap}}
b{{font-size:17px}} .pill{{border:1px solid #4a77bd;border-radius:999px;padding:4px 10px;color:#8dffbd;font-size:12px}}
a.btn,button{{background:#2f6df6;color:white;border:0;border-radius:8px;padding:8px 12px;text-decoration:none;cursor:pointer}}
.gray{{background:#43536d!important}} .red{{background:#d93636!important}} .green{{background:#059669!important}}
.grid{{display:grid;grid-template-columns:360px 1fr 390px;gap:12px;padding:12px}} .grid2{{display:grid;grid-template-columns:1fr 1fr;gap:12px;padding:12px}}
.panel{{background:#121d31;border:1px solid #314563;border-radius:14px;overflow:hidden}} .panel h3{{margin:0;background:#101a2d;border-bottom:1px solid #314563;padding:10px}} .body{{padding:10px}}
textarea,input,select,pre{{background:#040a14;color:#eaf2ff;border:1px solid #30445f;border-radius:8px;padding:8px;width:100%;box-sizing:border-box}} textarea{{min-height:90px}} pre{{white-space:pre-wrap;min-height:150px;max-height:55vh;overflow:auto}}
table{{width:100%;border-collapse:collapse;font-size:12px}}td,th{{border-bottom:1px solid #2d405c;padding:6px;text-align:left}} a{{color:#8db7ff}} .flowbar{{display:flex;align-items:center;gap:8px;padding:10px 12px;background:#081426;border-bottom:1px solid #314563;position:sticky;top:58px;z-index:4}}.step{{border:1px solid #43638d;border-radius:999px;padding:6px 12px;background:#13233c;color:#9eb2cf;font-size:12px;font-weight:bold}}.step.active{{background:#059669;color:white;border-color:#4ade80}}.step.running{{background:#b45309;color:white;border-color:#fbbf24}}.step.error{{background:#dc2626;color:white;border-color:#fca5a5}}.arrow{{color:#8db7ff}}.flowmeta{{margin-left:auto;color:#8dffbd;font-size:12px}}
.llmbar{{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:88px;z-index:8;flex-wrap:wrap}}.llmbar button{{padding:7px 12px}}.llmbar .mode{{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{{color:#fca5a5}}.llmbar .good{{color:#86efac}}</style>
<script>
async function api(path,opt={{}}){{const r=await fetch(path,{{headers:{{'Content-Type':'application/json'}},...opt}});const t=await r.text();try{{return JSON.parse(t)}}catch(e){{return {{ok:false,status:r.status,raw:t}}}}}}
function out(id,x){{let e=document.getElementById(id); if(e)e.textContent=typeof x==='string'?x:JSON.stringify(x,null,2)}}
function openWin(p){{window.open(p,'_blank','width=1180,height=780,menubar=no,toolbar=no,location=yes,resizable=yes,scrollbars=yes')}}
async function runChat(){{let q=document.getElementById('chatInput')?.value||'';out('mainOut',await api('/api/chat/run',{{method:'POST',body:JSON.stringify({{user_id:'default',message:q}})}}));refreshAll()}}
async function makeReport(){{let q=document.getElementById('reportQuery')?.value||'보고서 작성';out('fileOut',await api('/api/krxa/web/learning/run',{{method:'POST',body:JSON.stringify({{user_id:'default',query:q}})}}));loadFiles()}}
async function loadFiles(){{let d=await api('/api/files/reports');let files=((d.manifest||{{}}).files||[]);let e=document.getElementById('fileList'); if(e)e.innerHTML='<table><tr><th>파일</th><th>다운로드</th><th>관리</th></tr>'+files.map(f=>{{let n=f.name||f.stored_name||f.original_name;return `<tr><td>${{n}}</td><td><a target="_blank" href="${{f.download_url}}">다운로드</a></td><td><button class="red" onclick="deleteFile('${{n}}')">삭제</button></td></tr>`}}).join('')+'</table>';out('fileOut',d)}}
async function deleteFile(n){{if(confirm('삭제 '+n+'?')){{out('fileOut',await api('/api/files/delete',{{method:'POST',body:JSON.stringify({{filename:n}})}}));loadFiles()}}}}
async function loadTasks(){{let d=await api('/api/tasks/auto/status');let items=((d.queue||{{}}).items||[]);let e=document.getElementById('taskList');if(e)e.innerHTML='<table><tr><th>상태</th><th>작업</th><th>관리</th></tr>'+items.map(t=>`<tr><td>${{t.status}}</td><td>${{t.title}}</td><td><button onclick="tick()">실행</button><button class="red" onclick="deleteTask('${{t.id}}')">삭제</button></td></tr>`).join('')+'</table>';out('taskOut',d)}}
async function addTask(){{out('taskOut',await api('/api/tasks/add',{{method:'POST',body:JSON.stringify({{user_id:'default',title:document.getElementById('taskTitle')?.value||'새 작업',action:document.getElementById('taskAction')?.value||'manual_action'}})}}));loadTasks()}}
async function deleteTask(id){{out('taskOut',await api('/api/tasks/delete',{{method:'POST',body:JSON.stringify({{id}})}}));loadTasks()}}
async function tick(){{out('taskOut',await api('/api/autonomous/tick',{{method:'POST',body:JSON.stringify({{user_id:'default'}})}}));refreshAll()}}
async function loadMemory(){{let d=await api('/api/memory/user?user_id=default');let items=((d.memory||{{}}).items||[]);let e=document.getElementById('memoryList');if(e)e.innerHTML='<table><tr><th>ID</th><th>내용</th><th>관리</th></tr>'+items.map(m=>`<tr><td>${{m.id||''}}</td><td>${{((m.final||m.input||'')+'').slice(0,90)}}</td><td><button class="red" onclick="deleteMemory('${{m.id}}')">삭제</button></td></tr>`).join('')+'</table>';out('memoryOut',d)}}
async function addMemory(){{out('memoryOut',await api('/api/memory/add',{{method:'POST',body:JSON.stringify({{user_id:'default',text:document.getElementById('memoryInput')?.value||'새 기억'}})}}));loadMemory()}}
async function deleteMemory(id){{out('memoryOut',await api('/api/memory/delete',{{method:'POST',body:JSON.stringify({{user_id:'default',id}})}}));loadMemory()}}
async function loadDB(){{let db=document.getElementById('dbSelect')?.value||'krxai_db';let d=await api('/api/db/'+db);let ed=document.getElementById('dbEditor');if(ed)ed.value=JSON.stringify(d.data||d,null,2);out('dbOut',d)}}
async function saveDB(){{let data;try{{data=JSON.parse(document.getElementById('dbEditor').value)}}catch(e){{alert('JSON 오류 '+e);return}}out('dbOut',await api('/api/db/save',{{method:'POST',body:JSON.stringify({{db:document.getElementById('dbSelect').value,data:data.data||data}})}}))}}
async function loadLog(){{out('logOut',await api('/api/apilog'));out('rtOut',await api('/api/realtime/status'))}}

async function loadFlow(){{
  let d=await api('/api/flow/status');
  let f=d.flow||{{}};
  ['chat','task','engine','memory','file'].forEach(k=>{{
    let e=document.getElementById('flow_'+k);
    if(e){{e.className='step '+((f[k]=='active'||f[k]=='done'||f[k]=='ready')?'active':(f[k]=='pending'||f[k]=='created'?'running':''));}}
  }});
  let m=document.getElementById('flow_meta');
  if(m)m.textContent='TASK '+(d.counts?.tasks||0)+' · MEMORY '+(d.counts?.memory||0)+' · FILE '+(d.counts?.files||0);
}}

async function refreshAll(){{loadFiles();loadTasks();loadMemory();loadLog();loadFlow();try{{loadDB()}}catch(e){{}}}}
window.onload=refreshAll
</script></head><body>
<header><b>{title}</b><span class='pill'>{role}</span>
<a class='btn' href='/user' target='_blank'>사용자 UI</a>
<a class='btn' href='/admin' target='_blank'>관리자 UI</a>
<a class='btn' href='/dev' target='_blank'>개발자 UI</a>
<a class='btn gray' href='/control' target='_blank'>통합 CONTROL</a>
<button class='green' onclick='tick()'>AUTO TICK</button>
</header>{krxa_v325_flowbar_html()}{body}</body></html>"""

def krxa_v324_user_ui_html():
    body = """<div class='grid'>
<section class='panel'><h3>사용자 UI 연결바</h3><div class='body'><button onclick="openWin('/user/chat')">CHAT 새창</button> <button onclick="openWin('/user/status')">진행상태 새창</button> <button onclick="openWin('/user/files')">FILES 새창</button><textarea id='chatInput'>삼성전자 전략 보고서 작성</textarea><button onclick='runChat()'>요청 실행</button><pre id='mainOut'>대기</pre></div></section>
<section class='panel'><h3>결과 보고</h3><div class='body'><textarea id='reportQuery'>삼성전자 전략 보고서</textarea><button onclick='makeReport()'>보고서 생성</button><pre id='fileOut'>대기</pre></div></section>
<section class='panel'><h3>다운로드</h3><div class='body'><div id='fileList'></div></div></section>
</div>"""
    return krxa_v324_shell('KRXA USER UI V33.1','사용자용: 요청·결과·다운로드',body)

def krxa_v324_admin_ui_html():
    body = """<div class='grid'>
<section class='panel'><h3>관리자 요청 바</h3><div class='body'><button onclick="openWin('/admin/tasks')">TASK 관리 새창</button> <button onclick="openWin('/admin/approval')">승인/거부 새창</button> <button onclick="openWin('/admin/reports')">결과 보고 새창</button> <button onclick="openWin('/admin/result')">결과 관리 새창</button> <button onclick="openWin('/admin/memory')">MEMORY 관리 새창</button><textarea id='chatInput'>관리 지시 입력</textarea><button onclick='runChat()'>관리 지시 실행</button><pre id='mainOut'>대기</pre></div></section>
<section class='panel'><h3>TASK 관리</h3><div class='body'><input id='taskTitle' value='새 관리 작업'><input id='taskAction' value='manual_action'><button onclick='addTask()'>작업 추가</button><button onclick='tick()'>대기 실행</button><div id='taskList'></div><pre id='taskOut'></pre></div></section>
<section class='panel'><h3>결과 보고 / FILE</h3><div class='body'><textarea id='reportQuery'>관리 결과 보고서</textarea><button onclick='makeReport()'>보고서 생성</button><div id='fileList'></div><pre id='fileOut'></pre></div></section>
</div><div class='grid2'><section class='panel'><h3>MEMORY 관리</h3><div class='body'><textarea id='memoryInput'>관리 메모</textarea><button onclick='addMemory()'>추가</button><div id='memoryList'></div><pre id='memoryOut'></pre></div></section><section class='panel'><h3>REALTIME / API LOG</h3><div class='body'><button onclick='loadLog()'>로그 새로고침</button><pre id='rtOut'></pre><pre id='logOut'></pre></div></section></div>"""
    return krxa_v324_shell('KRXA ADMIN UI V33.1','관리자용: 관리지시·실행·결과보고',body)

def krxa_v324_dev_ui_html():
    body = """<div class='grid'>
<section class='panel'><h3>개발자 기능 바</h3><div class='body'><button onclick="openWin('/dev/new')">신규</button> <button onclick="openWin('/dev/edit?path=app.py')">수정</button> <button onclick="openWin('/dev/patch')">변경/PATCH</button> <button onclick="openWin('/dev/db')">DB</button> <button onclick="openWin('/dev/api')">API</button> <button onclick="openWin('/dev/log')">LOG</button> <button onclick="openWin('/dev/debug')">DEBUG</button><hr><textarea id='taskTitle'>개발 작업</textarea><input id='taskAction' value='dev_patch'><button onclick='addTask()'>개발 작업 추가</button><pre id='taskOut'></pre></div></section>
<section class='panel'><h3>DB / CONFIG EDITOR</h3><div class='body'><select id='dbSelect'><option>krxai_db</option><option>language_db</option><option>krxai_root</option><option>chat_history</option><option>role_based_operating_ui_policy</option><option>full_operating_control_ui_policy</option></select><button onclick='loadDB()'>조회</button><button onclick='saveDB()'>저장</button><textarea id='dbEditor' style='min-height:360px'></textarea><pre id='dbOut'></pre></div></section>
<section class='panel'><h3>API / DEBUG LOG</h3><div class='body'><button onclick='loadLog()'>API LOG</button><button onclick='loadTasks()'>TASK</button><button onclick='loadFiles()'>FILES</button><button onclick='loadMemory()'>MEMORY</button><pre id='logOut'></pre><pre id='rtOut'></pre></div></section>
</div>"""
    return krxa_v324_shell('KRXA DEVELOPER UI V33.1','개발자용: 신규·수정·변경·DB·API·DEBUG',body)

def krxa_v324_child_ui_html(kind):
    body = f"""<div class='grid2'><section class='panel'><h3>{kind.upper()} 새창 실행</h3><div class='body'><textarea id='chatInput'>{kind} 작업 요청</textarea><button onclick='runChat()'>요청 실행</button><button onclick='addTask()'>작업 추가</button><button onclick='tick()'>AUTO TICK</button><pre id='mainOut'>대기</pre></div></section><section class='panel'><h3>상태 / 결과</h3><div class='body'><div id='taskList'></div><div id='fileList'></div><pre id='taskOut'></pre><pre id='fileOut'></pre></div></section></div>"""
    return krxa_v324_shell('KRXA ROLE WINDOW V33.1 - '+kind.upper(),'새창 실행 화면',body)


def krxa_v325_flow_status():
    queue = krxa_v323_queue_load() if 'krxa_v323_queue_load' in globals() else load_json(CORE_DIR / 'krxai_task_queue.json', {'items': []})
    tasks = queue.get('items', [])
    files = krxa_v323_manifest().get('files', []) if 'krxa_v323_manifest' in globals() else []
    memory = krxa_user_memory_load('default') if 'krxa_user_memory_load' in globals() else {'items': []}
    recent_task = tasks[-1] if tasks else {}
    return {
        'ok': True,
        'schema': 'KRXA_FLOW_STATUS_V32_9',
        'flow': {
            'chat': 'ready',
            'task': recent_task.get('status', 'ready') if recent_task else 'ready',
            'engine': 'ready',
            'memory': 'active' if memory.get('items') else 'ready',
            'file': 'active' if files else 'ready'
        },
        'counts': {
            'tasks': len(tasks),
            'memory': len(memory.get('items', [])),
            'files': len(files)
        },
        'recent_task': recent_task,
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

def krxa_v325_flowbar_html():
    return """
<div class='flowbar'>
  <div class='step' id='flow_chat'>CHAT</div>
  <div class='arrow'>→</div>
  <div class='step' id='flow_task'>TASK</div>
  <div class='arrow'>→</div>
  <div class='step' id='flow_engine'>ENGINE</div>
  <div class='arrow'>→</div>
  <div class='step' id='flow_memory'>MEMORY</div>
  <div class='arrow'>→</div>
  <div class='step' id='flow_file'>FILE</div>
  <div class='flowmeta' id='flow_meta'>대기</div>
</div>
"""

def krxa_v325_child_ui_html(kind):
    k = (kind or 'work').replace('/', '_')
    body = f"""<div class='grid2'>
<section class='panel'><h3>{k.upper()} 작업창</h3><div class='body'>
<textarea id='chatInput'>{k} 작업 요청</textarea>
<button onclick='runChat()'>요청 실행</button>
<button onclick='addTask()'>작업 추가</button>
<button onclick='tick()'>AUTO TICK</button>
<button onclick='refreshAll()'>새로고침</button>
<pre id='mainOut'>대기</pre>
</div></section>
<section class='panel'><h3>연결 상태</h3><div class='body'>
<div id='taskList'></div>
<div id='fileList'></div>
<div id='memoryList'></div>
<pre id='taskOut'></pre><pre id='fileOut'></pre><pre id='memoryOut'></pre>
</div></section></div>"""
    return krxa_v324_shell('KRXA WORK WINDOW V33.1 - '+k.upper(), '클릭 새 작업창', body) if 'krxa_v324_shell' in globals() else '<html><body>work window</body></html>'


def krxa_v326_dashboard_state():
    try:
        flow = krxa_v325_flow_status() if 'krxa_v325_flow_status' in globals() else {'flow': {}, 'counts': {}}
    except Exception:
        flow = {'flow': {}, 'counts': {}}
    try:
        queue = krxa_v323_queue_load() if 'krxa_v323_queue_load' in globals() else load_json(CORE_DIR / 'krxai_task_queue.json', {'items': []})
    except Exception:
        queue = {'items': []}
    try:
        manifest = krxa_v323_manifest() if 'krxa_v323_manifest' in globals() else {'files': []}
    except Exception:
        manifest = {'files': []}
    try:
        mem = krxa_user_memory_load('default') if 'krxa_user_memory_load' in globals() else {'items': []}
    except Exception:
        mem = {'items': []}

    tasks = queue.get('items', [])
    files = manifest.get('files', [])
    memories = mem.get('items', [])
    latest_task = tasks[-1] if tasks else {}
    latest_file = files[0] if files else {}
    status = latest_task.get('status', 'ready') if latest_task else 'ready'
    if status in ['done', 'executed', 'complete', 'completed']:
        overall = 'done'
    elif status in ['pending', 'created']:
        overall = 'pending'
    elif status in ['running', 'processing']:
        overall = 'running'
    elif status in ['error', 'failed']:
        overall = 'error'
    else:
        overall = 'ready'

    return {
        'ok': True,
        'schema': 'KRXA_DASHBOARD_STATE_V32_9',
        'overall': overall,
        'summary': {
            'tasks': len(tasks),
            'pending': len([t for t in tasks if t.get('status') in ['pending', 'created']]),
            'done': len([t for t in tasks if t.get('status') in ['done', 'executed', 'complete', 'completed']]),
            'files': len(files),
            'memory': len(memories)
        },
        'latest': {
            'task': latest_task,
            'file': latest_file,
            'memory': memories[-1] if memories else {}
        },
        'flow': flow.get('flow', {}),
        'counts': flow.get('counts', {}),
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

def krxa_v326_real_control_center_html(role='user'):
    role = (role or 'user').lower()
    title_map = {'user':'KRXA USER CONTROL CENTER V33.1','admin':'KRXA ADMIN CONTROL CENTER V33.1','dev':'KRXA DEV CONTROL CENTER V33.1'}
    role_label = {'user':'사용자 관제: 요청·진행·결과','admin':'관리자 관제: 지시·승인·실행·보고','dev':'개발자 관제: 수정·DB·API·DEBUG'}
    extra = ''
    if role == 'admin':
        extra = """
<section class='card wide'><h3>관리자 실행 센터</h3>
<div class='toolbar'><button onclick="openWin('/admin/approval')">승인/거부</button><button onclick="openWin('/admin/tasks')">TASK 관리</button><button onclick="openWin('/admin/result')">결과 관리</button><button onclick='tick()'>대기 작업 실행</button></div>
<div id='taskCards' class='cards'></div></section>
<section class='card'><h3>MEMORY 관리</h3><textarea id='memoryInput'>관리 메모</textarea><button onclick='addMemory()'>기억 추가</button><div id='memoryCards'></div></section>
"""
    elif role == 'dev':
        extra = """
<section class='card wide'><h3>개발자 실행 센터</h3>
<div class='toolbar'><button onclick="openWin('/dev/new')">신규</button><button onclick="openWin('/dev/edit?path=app.py')">수정</button><button onclick="openWin('/dev/patch')">PATCH</button><button onclick="openWin('/dev/db')">DB</button><button onclick="openWin('/dev/api')">API</button><button onclick="openWin('/dev/debug')">DEBUG</button></div>
<select id='dbSelect'><option>krxai_db</option><option>language_db</option><option>krxai_root</option><option>chat_history</option><option>real_control_center_ui_policy</option></select>
<div class='toolbar'><button onclick='loadDB()'>DB 조회</button><button onclick='saveDB()'>DB 저장</button></div>
<textarea id='dbEditor' class='codebox'></textarea></section>
<section class='card'><h3>API / DEBUG</h3><button onclick='loadLog()'>로그 조회</button><pre id='logOut' class='detail'></pre></section>
"""
    else:
        extra = """
<section class='card wide'><h3>사용자 진행 상황</h3><div class='progressWrap'><div id='progressBar'></div></div><div id='resultCards' class='cards'></div></section>
<section class='card'><h3>빠른 작업</h3><button onclick="openWin('/user/status')">진행상태 새창</button><button onclick="openWin('/user/files')">다운로드 새창</button><button onclick='tick()'>상태 갱신</button></section>
"""
    return f"""<!doctype html><html><head><meta charset='utf-8'><title>{title_map.get(role,title_map['user'])}</title>
<style>
:root{{--bg:#07101f;--panel:#121d31;--head:#0d1729;--line:#314563;--btn:#2f6df6;--txt:#eaf2ff;--muted:#9eb2cf;--green:#10b981;--yellow:#f59e0b;--red:#ef4444;--blue:#3b82f6}}
*{{box-sizing:border-box}}body{{margin:0;background:linear-gradient(135deg,#07101f,#10234a);color:var(--txt);font-family:Arial,'Noto Sans KR',sans-serif;min-height:100vh}}
header{{position:sticky;top:0;z-index:10;background:#0d1729;border-bottom:1px solid var(--line);padding:10px 14px;display:flex;align-items:center;gap:10px;flex-wrap:wrap}}
header b{{font-size:18px}}.pill{{border:1px solid #4a77bd;border-radius:999px;padding:5px 12px;color:#8dffbd;font-size:12px}}
button,a.btn{{background:var(--btn);color:white;border:0;border-radius:10px;padding:9px 14px;margin:2px;cursor:pointer;text-decoration:none;font-weight:700}}.gray{{background:#43536d!important}}.green{{background:#059669!important}}.red{{background:#d93636!important}}
.flowbar{{display:flex;align-items:center;gap:8px;padding:12px 14px;background:#081426;border-bottom:1px solid var(--line);position:sticky;top:58px;z-index:9}}.step{{border:1px solid #43638d;border-radius:999px;padding:8px 14px;background:#13233c;color:#9eb2cf;font-size:12px;font-weight:bold}}.step.active{{background:#059669;color:white;border-color:#4ade80}}.step.running{{background:#b45309;color:white;border-color:#fbbf24}}.step.error{{background:#dc2626;color:white;border-color:#fca5a5}}.arrow{{color:#8db7ff}}.flowmeta{{margin-left:auto;color:#8dffbd;font-size:12px}}
.layout{{display:grid;grid-template-columns:360px 1fr 390px;gap:14px;padding:14px}}.card{{background:rgba(18,29,49,.94);border:1px solid var(--line);border-radius:18px;overflow:hidden;box-shadow:0 12px 30px rgba(0,0,0,.25)}}.card h3{{margin:0;background:#101a2d;border-bottom:1px solid var(--line);padding:12px 14px}}.card>div,.card>textarea,.card>pre,.card>select{{margin:12px}}.wide{{grid-column:span 2}}
textarea,input,select,pre{{background:#040a14;color:var(--txt);border:1px solid #30445f;border-radius:10px;padding:10px;width:calc(100% - 24px)}}textarea{{min-height:92px}}.codebox{{min-height:330px;font-family:monospace}}.detail{{white-space:pre-wrap;max-height:240px;overflow:auto;display:none}}.detail.open{{display:block}}
.kpi{{display:grid;grid-template-columns:repeat(4,1fr);gap:10px}}.kpiBox{{background:#07101f;border:1px solid #28415f;border-radius:14px;padding:14px}}.kpiBox strong{{display:block;font-size:28px;margin-top:6px}}.statusDot{{display:inline-block;width:12px;height:12px;border-radius:50%;margin-right:6px;background:var(--blue)}}.statusDot.done{{background:var(--green)}}.statusDot.pending{{background:var(--yellow)}}.statusDot.error{{background:var(--red)}}.statusDot.running{{background:var(--blue);animation:pulse 1s infinite}}@keyframes pulse{{50%{{opacity:.35}}}}
.cards{{display:grid;gap:10px}}.miniCard{{background:#07101f;border:1px solid #28415f;border-radius:14px;padding:12px}}.miniCard .title{{font-weight:800;margin-bottom:6px}}.miniCard .sub{{font-size:12px;color:var(--muted)}}.toolbar{{display:flex;gap:8px;flex-wrap:wrap}}.progressWrap{{height:16px;background:#07101f;border:1px solid #28415f;border-radius:999px;overflow:hidden}}#progressBar{{height:100%;width:0%;background:linear-gradient(90deg,#2f6df6,#10b981);transition:.4s}}
table{{width:100%;border-collapse:collapse;font-size:12px}}td,th{{border-bottom:1px solid #2d405c;padding:7px;text-align:left}}a{{color:#8db7ff}}
.llmbar{{display:flex;align-items:center;gap:8px;padding:10px 14px;background:#07101f;border-bottom:1px solid #314563;position:sticky;top:88px;z-index:8;flex-wrap:wrap}}.llmbar button{{padding:7px 12px}}.llmbar .mode{{border:1px solid #4a77bd;border-radius:999px;padding:6px 12px;color:#8dffbd;font-size:12px}}.llmbar .bad{{color:#fca5a5}}.llmbar .good{{color:#86efac}}</style>
<script>
async function api(path,opt={{}}){{const r=await fetch(path,{{headers:{{'Content-Type':'application/json'}},...opt}});const t=await r.text();try{{return JSON.parse(t)}}catch(e){{return {{ok:false,status:r.status,raw:t}}}}}}
function openWin(p){{window.open(p,'_blank','width=1180,height=780,menubar=no,toolbar=no,location=yes,resizable=yes,scrollbars=yes')}}
function setText(id,x){{let e=document.getElementById(id); if(e)e.textContent=typeof x==='string'?x:JSON.stringify(x,null,2)}}
function setStep(k,state){{let e=document.getElementById('flow_'+k); if(!e)return; e.className='step '+(state==='error'?'error':(state==='pending'||state==='created'?'running':'active'));}}
async function refreshState(){{
 const s=await api('/api/dashboard/state'); const sum=s.summary||{{}};
 document.getElementById('kpiTasks').textContent=sum.tasks||0; document.getElementById('kpiPending').textContent=sum.pending||0; document.getElementById('kpiFiles').textContent=sum.files||0; document.getElementById('kpiMemory').textContent=sum.memory||0;
 ['chat','task','engine','memory','file'].forEach(k=>setStep(k,(s.flow||{{}})[k]||'ready'));
 document.getElementById('flow_meta').textContent='TASK '+(sum.tasks||0)+' · MEMORY '+(sum.memory||0)+' · FILE '+(sum.files||0);
 let p=20; if(sum.tasks)p=45; if(sum.memory)p=70; if(sum.files)p=100; let pb=document.getElementById('progressBar'); if(pb)pb.style.width=p+'%';
 let dot=document.getElementById('overallDot'); if(dot)dot.className='statusDot '+(s.overall||'ready'); let label=document.getElementById('overallLabel'); if(label)label.textContent=s.overall||'ready';
 renderTaskCards(); renderFileCards(); renderMemoryCards();
}}

function krxaResultText(d){{
  if(!d)return '';
  return d.final || d.result || d.text || d.output || (d.llm && (d.llm.final||d.llm.output)) || (d.task && d.task.title) || JSON.stringify(d,null,2);
}}
function renderUserResult(d,q){{
  let e=document.getElementById('resultCards');
  if(!e)return;
  let txt=krxaResultText(d);
  let used=(d && (d.llm_used || d.last_llm_used || (d.judgement&&d.judgement.llm_used)))?'YES':'NO';
  e.innerHTML=`<div class='miniCard'><div class='title'>요청 결과</div><div class='sub'>${{q||''}} · LLM_USED ${{used}}</div><pre class='detail open'>${{txt}}</pre></div>` + e.innerHTML;
}}
async function runChat(){{let q=document.getElementById('chatInput')?.value||'';let d=await api('/api/chat/run',{{method:'POST',body:JSON.stringify({{user_id:'default',message:q}})}});setText('mainOut',d);renderUserResult(d,q);document.getElementById('detailMain')?.classList.add('open');refreshState();}}

async function makeReport(){{let q=document.getElementById('reportQuery')?.value||'보고서 작성';let d=await api('/api/krxa/web/learning/run',{{method:'POST',body:JSON.stringify({{user_id:'default',query:q}})}});setText('fileOut',d);renderUserResult(d,q);refreshState();}}
async function tick(){{let d=await api('/api/autonomous/tick',{{method:'POST',body:JSON.stringify({{user_id:'default'}})}});setText('taskOut',d);refreshState();}}
async function renderFileCards(){{let d=await api('/api/files/reports');let files=((d.manifest||{{}}).files||[]);let e=document.getElementById('fileCards'); if(!e)return; e.innerHTML=files.slice(0,8).map(f=>{{let n=f.name||f.stored_name||f.original_name;return `<div class='miniCard'><div class='title'>${{n}}</div><div class='sub'>${{f.size||0}} bytes</div><a class='btn' target='_blank' href='${{f.download_url}}'>다운로드</a><button class='red' onclick="deleteFile('${{n}}')">삭제</button></div>`}}).join(''); setText('fileOut',d);}}
async function deleteFile(n){{if(confirm('삭제 '+n+'?')){{setText('fileOut',await api('/api/files/delete',{{method:'POST',body:JSON.stringify({{filename:n}})}}));refreshState();}}}}
async function renderTaskCards(){{let d=await api('/api/tasks/auto/status');let items=((d.queue||{{}}).items||[]);let e=document.getElementById('taskCards'); if(!e)return; e.innerHTML=items.slice(-8).reverse().map(t=>`<div class='miniCard'><div class='title'><span class='statusDot ${{t.status}}'></span>${{t.title}}</div><div class='sub'>${{t.id}} · ${{t.status}}</div><button onclick='tick()'>실행</button><button class='red' onclick="deleteTask('${{t.id}}')">삭제</button></div>`).join(''); setText('taskOut',d);}}
async function addTask(){{setText('taskOut',await api('/api/tasks/add',{{method:'POST',body:JSON.stringify({{user_id:'default',title:document.getElementById('taskTitle')?.value||'새 작업',action:document.getElementById('taskAction')?.value||'manual_action'}})}}));refreshState();}}
async function deleteTask(id){{setText('taskOut',await api('/api/tasks/delete',{{method:'POST',body:JSON.stringify({{id}})}}));refreshState();}}
async function renderMemoryCards(){{let d=await api('/api/memory/user?user_id=default');let items=((d.memory||{{}}).items||[]);let e=document.getElementById('memoryCards'); if(!e)return; e.innerHTML=items.slice(-6).reverse().map(m=>`<div class='miniCard'><div class='title'>${{m.id||'memory'}}</div><div class='sub'>${{((m.final||m.input||'')+'').slice(0,120)}}</div><button class='red' onclick="deleteMemory('${{m.id}}')">삭제</button></div>`).join(''); setText('memoryOut',d);}}
async function addMemory(){{setText('memoryOut',await api('/api/memory/add',{{method:'POST',body:JSON.stringify({{user_id:'default',text:document.getElementById('memoryInput')?.value||'새 기억'}})}}));refreshState();}}
async function deleteMemory(id){{setText('memoryOut',await api('/api/memory/delete',{{method:'POST',body:JSON.stringify({{user_id:'default',id}})}}));refreshState();}}
async function loadDB(){{let db=document.getElementById('dbSelect')?.value||'krxai_db';let d=await api('/api/db/'+db);let ed=document.getElementById('dbEditor');if(ed)ed.value=JSON.stringify(d.data||d,null,2);setText('dbOut',d)}}
async function saveDB(){{let data;try{{data=JSON.parse(document.getElementById('dbEditor').value)}}catch(e){{alert('JSON 오류 '+e);return}}setText('dbOut',await api('/api/db/save',{{method:'POST',body:JSON.stringify({{db:document.getElementById('dbSelect').value,data:data.data||data}})}}))}}
async function loadLog(){{setText('logOut',await api('/api/apilog'));setText('rtOut',await api('/api/realtime/status'));document.getElementById('logOut')?.classList.add('open')}}
function toggle(id){{document.getElementById(id)?.classList.toggle('open')}}


async function llmControlStatus(){{
  try {{
    let d=await api('/api/llm/control/status');
    let e=document.getElementById('llmControlStatus');
    if(e){{
      let key=d.api_key_state||'UNKNOWN';
      let cls=key==='SET'?'good':'bad';
      e.innerHTML='MODE <b>'+(d.llm_mode||'-')+'</b> · KEY <span class="'+cls+'">'+key+'</span> · MODEL '+(d.model||'-')+' · LAST_USED '+(d.last_llm_used?'YES':'NO');
    }}
    setText('llmOut',d);
    return d;
  }} catch(err) {{
    let e=document.getElementById('llmControlStatus');
    if(e) e.innerHTML='LLM STATUS ERROR';
    return {{ok:false,error:String(err)}};
  }}
}}
async function llmControlMode(m){{
  try {{
    let d=await api('/api/llm/control/mode',{{method:'POST',body:JSON.stringify({{mode:m}})}});
    setText('llmOut',d);
    await llmControlStatus();
    return d;
  }} catch(err) {{
    setText('llmOut',{{ok:false,error:String(err)}});
    return {{ok:false,error:String(err)}};
  }}
}}


window.onload=()=>{{refreshState();llmControlStatus();setInterval(refreshState,3000);setInterval(llmControlStatus,3000)}}
</script></head><body>
<header><b>{title_map.get(role,title_map['user'])}</b><span class='pill'>{role_label.get(role,role)}</span>
<a class='btn' href='/user' target='_blank'>사용자 UI</a><a class='btn' href='/admin' target='_blank'>관리자 UI</a><a class='btn' href='/dev' target='_blank'>개발자 UI</a><a class='btn gray' href='/control' target='_blank'>통합 CONTROL</a><button class='green' onclick='tick()'>AUTO TICK</button></header>
<div class='flowbar'><div class='step' id='flow_chat'>CHAT</div><div class='arrow'>→</div><div class='step' id='flow_task'>TASK</div><div class='arrow'>→</div><div class='step' id='flow_engine'>ENGINE</div><div class='arrow'>→</div><div class='step' id='flow_memory'>MEMORY</div><div class='arrow'>→</div><div class='step' id='flow_file'>FILE</div><div class='flowmeta' id='flow_meta'>대기</div></div>
<div class='llmbar'><span class='mode'>KRXAI ↔ LLM CONTROL</span><button class='red' onclick="llmControlMode('off')">OFF</button><button class='green' onclick="llmControlMode('auto')">AUTO</button><button onclick="llmControlMode('force')">FORCE</button><button class='gray' onclick="llmControlMode('user')">USER</button><span id='llmControlStatus' class='mode'>확인 중</span></div><div class='layout'>
<section class='card'><h3>현재 상태</h3><div class='kpi'><div class='kpiBox'>TASK<strong id='kpiTasks'>0</strong></div><div class='kpiBox'>대기<strong id='kpiPending'>0</strong></div><div class='kpiBox'>FILE<strong id='kpiFiles'>0</strong></div><div class='kpiBox'>MEMORY<strong id='kpiMemory'>0</strong></div></div><div style='margin:12px'><span id='overallDot' class='statusDot'></span><b id='overallLabel'>ready</b></div></section>
<section class='card wide'><h3>요청 / 실행</h3><textarea id='chatInput'>삼성전자 전략 보고서 작성</textarea><div class='toolbar'><button onclick='runChat()'>요청 실행</button><button onclick='tick()'>대기 실행</button><button onclick="toggle('detailMain')">상세 JSON</button></div><pre id='mainOut' class='detail'></pre></section>
<section class='card'><h3>보고서 / 파일</h3><textarea id='reportQuery'>삼성전자 전략 보고서</textarea><div class='toolbar'><button onclick='makeReport()'>보고서 생성</button><button onclick='renderFileCards()'>파일 새로고침</button></div><div id='fileCards' class='cards'></div><pre id='fileOut' class='detail'></pre></section>
{extra}
<section class='card'><h3>최근 TASK</h3><div id='taskCards' class='cards'></div><pre id='taskOut' class='detail'></pre></section>
<section class='card'><h3>최근 MEMORY</h3><div id='memoryCards' class='cards'></div><pre id='memoryOut' class='detail'></pre></section>
</div></body></html>"""


def krxa_v327_llm_config_path():
    return CORE_DIR / 'krxa_llm_connection_config.json'

def krxa_v327_llm_config():
    default = {
        'schema': 'KRXA_LLM_CONNECTION_CONFIG_V32_9',
        'llm_enabled': False,
        'mode': 'internal_first',
        'provider': 'openai',
        'model_env': 'OPENAI_MODEL',
        'api_key_env': 'OPENAI_API_KEY',
        'updated_at': datetime.now(timezone.utc).isoformat()
    }
    return load_json(krxa_v327_llm_config_path(), default)

def krxa_v327_llm_save(cfg):
    cfg['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(krxa_v327_llm_config_path(), cfg)
    try:
        krxa_event_emit('LLM_CONFIG_UPDATE', {'llm_enabled': cfg.get('llm_enabled'), 'mode': cfg.get('mode')}, source='v32_7_llm_control')
    except Exception:
        pass
    return cfg

def krxa_v327_llm_toggle(enabled=None):
    cfg = krxa_v327_llm_config()
    if enabled is None:
        cfg['llm_enabled'] = not bool(cfg.get('llm_enabled'))
    else:
        cfg['llm_enabled'] = bool(enabled)
    return {'ok': True, 'config': krxa_v327_llm_save(cfg)}

def krxa_v327_llm_status():
    cfg = krxa_v327_llm_config()
    api_key_state = 'SET' if os.environ.get(cfg.get('api_key_env', 'OPENAI_API_KEY')) else 'MISSING'
    model = os.environ.get(cfg.get('model_env', 'OPENAI_MODEL'), os.environ.get('OPENAI_MODEL', 'gpt-4.1'))
    return {
        'ok': True,
        'schema': 'KRXA_LLM_STATUS_V32_9',
        'llm_enabled': bool(cfg.get('llm_enabled')),
        'mode': cfg.get('mode', 'internal_first'),
        'provider': cfg.get('provider', 'openai'),
        'api_key_state': api_key_state,
        'model': model,
        'rule': 'OFF이면 외부 ChatGPT API 호출 금지, ON이면 필요한 경우만 호출',
        'config': cfg,
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

def krxa_v327_should_use_llm(message='', intent=''):
    cfg = krxa_v327_llm_config()
    if not cfg.get('llm_enabled'):
        return False, 'LLM OFF: 내부 KRXAI_DB + 언어DB + MEMORY 판단'
    text = (message or '').lower()
    # 필요 시만 외부 LLM 사용: 보고서/분석/검색/복잡한 자연어 요청 중심
    if intent in ['report_or_research', 'report', 'research', 'analysis']:
        return True, 'LLM ON: 보고서/분석 요청으로 보조 판단 허용'
    if any(k in text for k in ['보고서', '리포트', '분석', '전략', '복잡', '설명', '정리']):
        return True, 'LLM ON: 복합 자연어 처리 보조 허용'
    return False, 'LLM ON이지만 내부 판단으로 충분'

def krxa_v327_optional_llm_call(message, intent=''):
    use, reason = krxa_v327_should_use_llm(message, intent)
    if not use:
        return {'llm_enabled': krxa_v327_llm_config().get('llm_enabled', False), 'llm_used': False, 'llm_reason': reason, 'llm_output': None}
    try:
        # OpenAI SDK가 설치되어 있고 API KEY가 있을 때만 실제 호출.
        from openai import OpenAI
        client = OpenAI()
        model = os.environ.get('OPENAI_MODEL', 'gpt-4.1')
        res = client.chat.completions.create(
            model=model,
            messages=[
                {'role': 'system', 'content': 'You are KRXAI external LLM bridge. Return concise Korean operational judgement.'},
                {'role': 'user', 'content': message}
            ],
            temperature=0.2
        )
        output = res.choices[0].message.content
        return {'llm_enabled': True, 'llm_used': True, 'llm_reason': reason, 'llm_output': output, 'model': model}
    except Exception as e:
        return {'llm_enabled': True, 'llm_used': False, 'llm_reason': 'LLM 호출 실패, 내부 판단으로 fallback', 'llm_error': str(e), 'llm_output': None}

# 기존 krxa_chat_run을 감싸서 llm_enabled/llm_used/llm_reason을 결과에 반드시 붙인다.
if 'krxa_chat_run' in globals() and 'krxa_chat_run_v327_original' not in globals():
    krxa_chat_run_v327_original = krxa_chat_run
    def krxa_chat_run(user_id='default', message=''):
        base = krxa_chat_run_v327_original(user_id, message)
        intent = base.get('intent') or base.get('task', {}).get('type') or ''
        llm = krxa_v327_optional_llm_call(message, intent)
        base['llm_mode'] = llm.get('llm_mode')
        base['force_llm'] = llm.get('force_llm')
        base['llm_enabled'] = llm.get('llm_enabled')
        base['llm_used'] = llm.get('llm_used')
        base['llm_reason'] = llm.get('llm_reason')
        if llm.get('llm_output'):
            base['llm_output'] = llm.get('llm_output')
            # final은 내부 결과를 유지하되 보조 판단을 별도 필드로 둔다.
        try:
            krxa_user_memory_store(user_id, {
                'role': 'llm_connection_control',
                'input': message,
                'intent': intent,
                'final': base.get('final'),
                'llm_enabled': base.get('llm_enabled'),
                'llm_used': base.get('llm_used'),
                'llm_reason': base.get('llm_reason'),
                'engine': 'KRXA_LLM_CONNECTION_ONOFF_CONTROL_V32_9'
            })
        except Exception:
            pass
        return base

def krxa_v327_llm_control_panel_html():
    return """<section class='card'><h3>LLM 연결 제어</h3>
<div class='body'>
<div class='toolbar'>
<button class='green' onclick='llmSet(true)'>LLM ON</button>
<button class='red' onclick='llmSet(false)'>LLM OFF</button>
<button onclick='llmStatus()'>상태 확인</button>
</div>
<div class='miniCard'><div class='title'>현재 LLM 상태</div><div id='llmBadge' class='sub'>확인 중</div></div>
<pre id='llmOut' class='detail open'></pre>
</div></section>"""

def krxa_v327_llm_script():
    return """
async function llmStatus(){
  let d=await api('/api/llm/config');
  let b=document.getElementById('llmBadge');
  if(b)b.textContent='MODE '+(d.llm_mode|| (d.llm_enabled?'auto':'off'))+' · '+(d.llm_enabled?'ON':'OFF')+' · FORCE '+(d.force_llm?'YES':'NO')+' · KEY '+d.api_key_state+' · '+d.model;
  setText('llmOut',d);
}
async function llmSet(v){
  let d=await api('/api/llm/toggle',{method:'POST',body:JSON.stringify({enabled:v})});
  setText('llmOut',d); llmStatus();
}
"""


def krxa_v328_set_llm_mode(mode='auto'):
    mode = (mode or 'auto').lower()
    cfg = krxa_v327_llm_config() if 'krxa_v327_llm_config' in globals() else load_json(CORE_DIR / 'krxa_llm_connection_config.json', {})
    if mode == 'off':
        cfg['llm_enabled'] = False
        cfg['force_llm'] = False
        cfg['llm_mode'] = 'off'
    elif mode == 'force':
        cfg['llm_enabled'] = True
        cfg['force_llm'] = True
        cfg['llm_mode'] = 'force'
    else:
        cfg['llm_enabled'] = True
        cfg['force_llm'] = False
        cfg['llm_mode'] = 'auto'
        mode = 'auto'
    cfg['mode'] = 'off_auto_force'
    cfg['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(CORE_DIR / 'krxa_llm_connection_config.json', cfg)
    try:
        krxa_event_emit('LLM_MODE_SET', {'llm_mode': cfg.get('llm_mode'), 'llm_enabled': cfg.get('llm_enabled'), 'force_llm': cfg.get('force_llm')}, source='v32_8_llm_mode')
    except Exception:
        pass
    return {'ok': True, 'schema': 'KRXA_LLM_MODE_SET_V32_9', 'mode': mode, 'config': cfg}

def krxa_v328_llm_status():
    base = krxa_v327_llm_status() if 'krxa_v327_llm_status' in globals() else {'ok': True, 'config': load_json(CORE_DIR / 'krxa_llm_connection_config.json', {})}
    cfg = base.get('config', {})
    if not cfg.get('llm_enabled'):
        llm_mode = 'off'
    elif cfg.get('force_llm'):
        llm_mode = 'force'
    else:
        llm_mode = 'auto'
    base['schema'] = 'KRXA_LLM_STATUS_V32_9'
    base['llm_mode'] = llm_mode
    base['force_llm'] = bool(cfg.get('force_llm'))
    base['mode_description'] = {
        'off': '외부 ChatGPT API 호출 금지',
        'auto': '내부 판단 우선, 필요 시만 ChatGPT 호출',
        'force': '요청마다 ChatGPT 강제 호출'
    }.get(llm_mode)
    return base

# v32.7 판단 함수를 FORCE/AUTO/OFF로 재정의
def krxa_v327_should_use_llm(message='', intent=''):
    cfg = krxa_v327_llm_config()
    if not cfg.get('llm_enabled'):
        return False, 'LLM OFF: 외부 호출 금지, 내부 KRXAI_DB + 언어DB + MEMORY 판단'
    if cfg.get('force_llm'):
        return True, 'LLM FORCE: 사용자가 강제 호출 모드를 선택함'
    text = (message or '').lower()
    if intent in ['report_or_research', 'report', 'research', 'analysis']:
        return True, 'LLM AUTO: 보고서/분석 요청으로 보조 판단 허용'
    if any(k in text for k in ['보고서', '리포트', '분석', '전략', '복잡', '설명', '정리', '비교']):
        return True, 'LLM AUTO: 복합 자연어 처리 보조 허용'
    return False, 'LLM AUTO: 내부 판단으로 충분'

# v32.7 optional call도 상태 필드 보강
if 'krxa_v327_optional_llm_call' in globals():
    krxa_v328_optional_llm_call_original = krxa_v327_optional_llm_call
    def krxa_v327_optional_llm_call(message, intent=''):
        cfg = krxa_v327_llm_config()
        if not cfg.get('llm_enabled'):
            llm_mode = 'off'
        elif cfg.get('force_llm'):
            llm_mode = 'force'
        else:
            llm_mode = 'auto'
        result = krxa_v328_optional_llm_call_original(message, intent)
        result['llm_mode'] = llm_mode
        result['force_llm'] = bool(cfg.get('force_llm'))
        return result


def krxa_v329_llm_control_status():
    st = krxa_v328_llm_status() if 'krxa_v328_llm_status' in globals() else krxa_v327_llm_status()
    cfg = st.get('config', {})
    mode = st.get('llm_mode') or ('off' if not cfg.get('llm_enabled') else ('force' if cfg.get('force_llm') else 'auto'))
    last = load_json(CORE_DIR / 'krxa_llm_last_call.json', {})
    return {
        'ok': True,
        'schema': 'KRXA_LLM_CONTROL_BAR_STATUS_V32_9',
        'llm_mode': mode,
        'llm_enabled': bool(cfg.get('llm_enabled')),
        'force_llm': bool(cfg.get('force_llm')),
        'user_select_mode': bool(cfg.get('user_select_mode')),
        'api_key_state': st.get('api_key_state') or ('SET' if os.environ.get('OPENAI_API_KEY') else 'MISSING'),
        'model': st.get('model') or os.environ.get('OPENAI_MODEL', 'gpt-4.1'),
        'last_llm_used': last.get('llm_used', False),
        'last_llm_reason': last.get('llm_reason', ''),
        'last_updated_at': last.get('updated_at'),
        'config': cfg,
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

def krxa_v329_set_llm_control_mode(mode='auto'):
    mode = (mode or 'auto').lower()
    if mode == 'user':
        result = krxa_v328_set_llm_mode('auto') if 'krxa_v328_set_llm_mode' in globals() else krxa_v327_llm_toggle(True)
        cfg = result.get('config', krxa_v327_llm_config())
        cfg['llm_enabled'] = True
        cfg['force_llm'] = False
        cfg['user_select_mode'] = True
        cfg['llm_mode'] = 'user'
        cfg['updated_at'] = datetime.now(timezone.utc).isoformat()
        save_json(CORE_DIR / 'krxa_llm_connection_config.json', cfg)
        return {'ok': True, 'schema': 'KRXA_LLM_CONTROL_BAR_MODE_SET_V32_9', 'mode': 'user', 'config': cfg}
    result = krxa_v328_set_llm_mode(mode) if 'krxa_v328_set_llm_mode' in globals() else krxa_v327_llm_toggle(mode != 'off')
    cfg = result.get('config', krxa_v327_llm_config())
    cfg['user_select_mode'] = False
    cfg['llm_mode'] = mode if mode in ['off','auto','force'] else 'auto'
    cfg['updated_at'] = datetime.now(timezone.utc).isoformat()
    save_json(CORE_DIR / 'krxa_llm_connection_config.json', cfg)
    result['schema'] = 'KRXA_LLM_CONTROL_BAR_MODE_SET_V32_9'
    result['config'] = cfg
    return result

def krxa_v327_should_use_llm(message='', intent=''):
    cfg = krxa_v327_llm_config()
    msg = (message or '').lower()
    if not cfg.get('llm_enabled'):
        return False, 'LLM OFF: 외부 호출 금지'
    if cfg.get('llm_mode') == 'user' or cfg.get('user_select_mode'):
        if any(k in msg for k in ['ai 사용', 'chatgpt', 'gpt', 'llm', '외부 ai', '외부 판단']):
            return True, 'LLM USER: 사용자가 AI/ChatGPT 호출을 명시함'
        return False, 'LLM USER: 사용자가 외부 AI 호출을 선택하지 않음'
    if cfg.get('force_llm'):
        return True, 'LLM FORCE: 강제 호출 모드'
    if intent in ['report_or_research', 'report', 'research', 'analysis']:
        return True, 'LLM AUTO: 보고서/분석 요청'
    if any(k in msg for k in ['보고서', '리포트', '분석', '전략', '복잡', '설명', '정리', '비교']):
        return True, 'LLM AUTO: 복합 자연어 처리 보조'
    return False, 'LLM AUTO: 내부 판단으로 충분'

if 'krxa_v327_optional_llm_call' in globals() and 'krxa_v329_optional_original' not in globals():
    krxa_v329_optional_original = krxa_v327_optional_llm_call
    def krxa_v327_optional_llm_call(message, intent=''):
        result = krxa_v329_optional_original(message, intent)
        try:
            save_json(CORE_DIR / 'krxa_llm_last_call.json', {
                'message': message,
                'intent': intent,
                'llm_mode': result.get('llm_mode'),
                'force_llm': result.get('force_llm'),
                'llm_enabled': result.get('llm_enabled'),
                'llm_used': result.get('llm_used'),
                'llm_reason': result.get('llm_reason'),
                'model': result.get('model'),
                'llm_error': result.get('llm_error'),
                'updated_at': datetime.now(timezone.utc).isoformat()
            })
        except Exception:
            pass
        return result


# ---------------- V33.1 DEV SOURCE TREE + USER RESULT RENDER ----------------
DEV_STORAGE_DIR = ROOT / 'storage'
DEV_EDITABLE_ROOTS = [
    BASE_DIR / 'app.py',
    BASE_DIR / 'README.md',
    BASE_DIR / 'index.html',
    CORE_DIR,
    DEV_STORAGE_DIR,
    BASE_DIR / 'scripts',
    BASE_DIR / 'samples',
]

def krxa_v330_safe_rel_path(path_value='app.py'):
    raw = (path_value or 'app.py').replace('\\', '/').lstrip('/')
    target = (BASE_DIR / raw).resolve()
    base = BASE_DIR.resolve()
    if not str(target).startswith(str(base)):
        raise ValueError('path outside project is not allowed')
    if target.is_dir():
        raise ValueError('directory cannot be edited directly')
    allowed = False
    for root in DEV_EDITABLE_ROOTS:
        rr = root.resolve()
        if target == rr or str(target).startswith(str(rr) + os.sep):
            allowed = True
            break
    if not allowed:
        raise ValueError('path is not in editable whitelist')
    return raw, target

def krxa_v330_dev_source_tree():
    allowed_ext = {'.py','.json','.html','.css','.js','.md','.txt','.bat'}
    items = []
    roots = [('root', BASE_DIR), ('core', CORE_DIR), ('storage', DEV_STORAGE_DIR), ('scripts', BASE_DIR / 'scripts'), ('samples', BASE_DIR / 'samples')]
    seen = set()
    for label, root in roots:
        if not root.exists():
            continue
        files = [root] if root.is_file() else [p for p in root.rglob('*') if p.is_file()]
        for p in files:
            if p.name.startswith('.') and p.name not in ['.nojekyll']:
                continue
            if p.suffix.lower() not in allowed_ext:
                continue
            rel = p.relative_to(BASE_DIR).as_posix()
            if rel in seen:
                continue
            seen.add(rel)
            try:
                size = p.stat().st_size
            except Exception:
                size = 0
            items.append({'name': p.name, 'path': rel, 'group': rel.split('/')[0] if '/' in rel else 'root', 'size': size, 'ext': p.suffix.lower() or 'file'})
    items.sort(key=lambda x: (x['group'], x['path']))
    return {'ok': True, 'schema': 'KRXA_DEV_SOURCE_TREE_V33_0', 'items': items, 'count': len(items), 'updated_at': datetime.now(timezone.utc).isoformat()}

def krxa_v330_dev_source_read(path_value='app.py'):
    try:
        rel, target = krxa_v330_safe_rel_path(path_value)
        if not target.exists():
            return {'ok': False, 'error': 'file not found', 'path': rel}
        text = target.read_text(encoding='utf-8', errors='replace')
        return {'ok': True, 'schema': 'KRXA_DEV_SOURCE_READ_V33_0', 'path': rel, 'size': len(text), 'content': text, 'updated_at': datetime.now(timezone.utc).isoformat()}
    except Exception as e:
        return {'ok': False, 'schema': 'KRXA_DEV_SOURCE_READ_V33_0', 'error': str(e), 'path': path_value}

def krxa_v330_dev_source_save(path_value='app.py', content=''):
    try:
        rel, target = krxa_v330_safe_rel_path(path_value)
        backup_dir = DEV_STORAGE_DIR / 'dev_backups'
        backup_dir.mkdir(parents=True, exist_ok=True)
        if target.exists():
            stamp = datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')
            backup_name = rel.replace('/','__') + '.' + stamp + '.bak'
            (backup_dir / backup_name).write_text(target.read_text(encoding='utf-8', errors='replace'), encoding='utf-8')
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(content or '', encoding='utf-8')
        compile_ok = None
        compile_error = ''
        if rel == 'app.py' or rel.endswith('.py'):
            try:
                py_compile.compile(str(target), doraise=True)
                compile_ok = True
            except Exception as e:
                compile_ok = False
                compile_error = str(e)
        try:
            krxa_event_emit('DEV_SOURCE_SAVED', {'path': rel, 'compile_ok': compile_ok}, source='v33_dev_source_tree')
        except Exception:
            pass
        return {'ok': True, 'schema': 'KRXA_DEV_SOURCE_SAVE_V33_0', 'path': rel, 'size': len(content or ''), 'compile_ok': compile_ok, 'compile_error': compile_error, 'note': 'app.py 같은 서버 소스 수정은 저장 후 Render 재배포/재시작이 필요합니다. JSON/DB성 파일은 즉시 반영됩니다.', 'updated_at': datetime.now(timezone.utc).isoformat()}
    except Exception as e:
        return {'ok': False, 'schema': 'KRXA_DEV_SOURCE_SAVE_V33_0', 'error': str(e), 'path': path_value}

def krxa_v330_dev_source_editor_html(path_value='app.py'):
    safe_path = (path_value or 'app.py').replace("'", '').replace('"','')
    return f"""<!doctype html><html><head><meta charset='utf-8'><title>KRXA DEV SOURCE EDITOR V33.1</title>
<style>
:root{{--bg:#07101f;--panel:#121d31;--head:#0d1729;--line:#314563;--btn:#2f6df6;--txt:#eaf2ff;--muted:#9eb2cf;--green:#10b981;--red:#ef4444}}
*{{box-sizing:border-box}}body{{margin:0;background:linear-gradient(135deg,#07101f,#10234a);color:var(--txt);font-family:Arial,'Noto Sans KR',sans-serif;height:100vh;overflow:hidden}}
header{{height:58px;background:#0d1729;border-bottom:1px solid var(--line);display:flex;gap:10px;align-items:center;padding:10px 14px}}
button,.btn{{background:var(--btn);color:white;border:0;border-radius:10px;padding:10px 14px;font-weight:bold;cursor:pointer}}button.red{{background:var(--red)}}button.green{{background:var(--green)}}
.wrap{{display:grid;grid-template-columns:310px 1fr;height:calc(100vh - 58px)}}
.tree{{border-right:1px solid var(--line);background:#0b1628;overflow:auto;padding:12px}}
.editor{{display:grid;grid-template-rows:auto 1fr auto;height:100%;padding:12px;gap:10px}}
.file{{padding:8px;border-bottom:1px solid #20344f;cursor:pointer;font-size:13px;color:#dbeafe}}
.file:hover{{background:#13233c}}.group{{margin-top:12px;color:#8dffbd;font-weight:bold}}
textarea{{width:100%;height:100%;background:#020812;color:#eaf2ff;border:1px solid var(--line);border-radius:10px;padding:12px;font-family:Consolas,monospace;font-size:13px;line-height:1.35}}
pre{{margin:0;background:#020812;color:#dbeafe;border:1px solid var(--line);border-radius:10px;padding:10px;max-height:130px;overflow:auto}}
input{{background:#020812;color:#eaf2ff;border:1px solid var(--line);border-radius:10px;padding:10px;width:520px}}
</style></head><body>
<header><b>KRXA DEV SOURCE TREE EDITOR V33.1</b><span>사용자UI/관제UI 프로그램 수정</span><button onclick="location.href='/dev'">개발자 UI</button><button onclick="location.href='/user'">사용자 UI</button></header>
<div class='wrap'>
  <aside class='tree'><h3>프로그램 트리</h3><div id='treeList'>불러오는 중</div></aside>
  <main class='editor'>
    <div><input id='path' value='{safe_path}'><button onclick='loadFile()'>열기</button><button class='green' onclick='saveFile()'>완료/저장</button><button class='red' onclick='deployGit()'>Git 자동배포</button><button onclick='deployConfig()'>배포상태</button><button onclick='openUser()'>사용자UI 확인</button></div>
    <textarea id='code'>대기</textarea>
    <pre id='out'>대기</pre>
  </main>
</div>
<script>
async function api(path,opt={{}}){{const r=await fetch(path,{{headers:{{'Content-Type':'application/json'}},...opt}});const t=await r.text();try{{return JSON.parse(t)}}catch(e){{return {{ok:false,status:r.status,raw:t}}}}}}
function esc(s){{return (s||'').replace(/[&<>]/g,m=>({{'&':'&amp;','<':'&lt;','>':'&gt;'}}[m]))}}
async function loadTree(){{
  let d=await api('/api/dev/source/tree');
  let items=d.items||[];
  let by={{}};
  items.forEach(x=>{{(by[x.group]||(by[x.group]=[])).push(x)}});
  let html='';
  Object.keys(by).sort().forEach(g=>{{html+=`<div class='group'>${{g}}</div>`;by[g].forEach(f=>{{html+=`<div class='file' onclick="selectFile('${{f.path}}')">📄 ${{esc(f.path)}} <span style='color:#8aa'>${{f.size}}b</span></div>`}})}});
  document.getElementById('treeList').innerHTML=html||'파일 없음';
}}
function selectFile(p){{document.getElementById('path').value=p;loadFile();}}
async function loadFile(){{
  let p=document.getElementById('path').value;
  let d=await api('/api/dev/source/read?path='+encodeURIComponent(p));
  document.getElementById('out').textContent=JSON.stringify(d,null,2);
  if(d.ok)document.getElementById('code').value=d.content;
}}
async function saveFile(){{
  let p=document.getElementById('path').value;
  let content=document.getElementById('code').value;
  let d=await api('/api/dev/source/save',{{method:'POST',body:JSON.stringify({{path:p,content}})}});
  document.getElementById('out').textContent=JSON.stringify(d,null,2);
  if(d.compile_ok===false) alert('저장은 됐지만 Python 컴파일 오류가 있습니다. 백업이 생성되었습니다.');
}}
async function deployConfig(){{
  let d=await api('/api/dev/deploy/config');
  document.getElementById('out').textContent=JSON.stringify(d,null,2);
}}
async function deployGit(){{
  let p=document.getElementById('path').value||'app.py';
  if(!confirm('현재 파일을 GitHub에 자동 반영하고 Render Auto Deploy를 실행하시겠습니까?')) return;
  let d=await api('/api/dev/deploy/github',{{method:'POST',body:JSON.stringify({{paths:[p],message:'KRXA auto deploy from dev UI: '+p}})}});
  document.getElementById('out').textContent=JSON.stringify(d,null,2);
  if(d.ok) alert('GitHub 반영 완료. Render Auto Deploy가 켜져 있으면 곧 재배포됩니다.');
  else alert('자동배포 실패: '+(d.error||JSON.stringify(d.errors||d)));
}}
function openUser(){{window.open('/user','_blank')}}
window.onload=()=>{{loadTree();loadFile();}}
</script></body></html>"""


# ---------------- V33.1 RENDER + GITHUB AUTO DEPLOY ----------------
def krxa_v331_deploy_config():
    return {
        'ok': True,
        'schema': 'KRXA_RENDER_GIT_AUTO_DEPLOY_CONFIG_V33_1',
        'github_repo': os.environ.get('GITHUB_REPO', 'infokrxa-lgtm/KRXA'),
        'github_branch': os.environ.get('GITHUB_BRANCH', 'main'),
        'has_github_token': bool(os.environ.get('GITHUB_TOKEN') or os.environ.get('GH_TOKEN')),
        'render_service': os.environ.get('RENDER_SERVICE_NAME', 'krxa'),
        'mode': 'github_contents_api',
        'note': 'GITHUB_TOKEN 환경변수가 있어야 개발자 UI에서 GitHub 자동 반영이 가능합니다.'
    }

def krxa_v331_github_api(method, url, token, body=None):
    data = None
    if body is not None:
        data = json.dumps(body).encode('utf-8')
    req = urllib.request.Request(url, data=data, method=method)
    req.add_header('Authorization', 'Bearer ' + token)
    req.add_header('Accept', 'application/vnd.github+json')
    req.add_header('X-GitHub-Api-Version', '2022-11-28')
    req.add_header('User-Agent', 'KRXA-V33.1-AUTO-DEPLOY')
    if data is not None:
        req.add_header('Content-Type', 'application/json')
    with urllib.request.urlopen(req, timeout=30) as res:
        raw = res.read().decode('utf-8', errors='replace')
        try:
            return json.loads(raw) if raw else {}
        except Exception:
            return {'raw': raw}

def krxa_v331_get_github_sha(repo, branch, path, token):
    url = f'https://api.github.com/repos/{repo}/contents/{urllib.parse.quote(path)}?ref={urllib.parse.quote(branch)}'
    try:
        r = krxa_v331_github_api('GET', url, token)
        return r.get('sha')
    except Exception:
        return None

def krxa_v331_deploy_files_to_github(paths=None, message=None):
    token = os.environ.get('GITHUB_TOKEN') or os.environ.get('GH_TOKEN')
    repo = os.environ.get('GITHUB_REPO', 'infokrxa-lgtm/KRXA')
    branch = os.environ.get('GITHUB_BRANCH', 'main')
    if not token:
        return {
            'ok': False,
            'schema': 'KRXA_RENDER_GIT_AUTO_DEPLOY_V33_1',
            'error': 'GITHUB_TOKEN 환경변수가 없습니다. Render Environment에 GITHUB_TOKEN을 등록해야 합니다.',
            'required_env': ['GITHUB_TOKEN', 'GITHUB_REPO=infokrxa-lgtm/KRXA', 'GITHUB_BRANCH=main']
        }
    if not paths:
        paths = ['app.py']
    if isinstance(paths, str):
        paths = [paths]
    deployed = []
    errors = []
    for p in paths:
        try:
            rel, target = krxa_v330_safe_rel_path(p) if 'krxa_v330_safe_rel_path' in globals() else (p, (BASE_DIR / p).resolve())
            if not target.exists():
                errors.append({'path': rel, 'error': 'file not found'})
                continue
            content_bytes = target.read_bytes()
            import base64
            encoded = base64.b64encode(content_bytes).decode('ascii')
            sha = krxa_v331_get_github_sha(repo, branch, rel, token)
            body = {
                'message': message or f'KRXA auto deploy: {rel}',
                'content': encoded,
                'branch': branch
            }
            if sha:
                body['sha'] = sha
            url = f'https://api.github.com/repos/{repo}/contents/{urllib.parse.quote(rel)}'
            r = krxa_v331_github_api('PUT', url, token, body)
            deployed.append({'path': rel, 'sha': (r.get('content') or {}).get('sha'), 'html_url': (r.get('content') or {}).get('html_url')})
        except Exception as e:
            errors.append({'path': p, 'error': str(e)})
    ok = len(deployed) > 0 and not errors
    try:
        krxa_event_emit('GITHUB_AUTO_DEPLOY', {'deployed': deployed, 'errors': errors}, source='v33_1_render_git_auto_deploy')
    except Exception:
        pass
    return {
        'ok': ok,
        'schema': 'KRXA_RENDER_GIT_AUTO_DEPLOY_V33_1',
        'repo': repo,
        'branch': branch,
        'deployed': deployed,
        'errors': errors,
        'note': 'GitHub 반영 후 Render Auto Deploy가 켜져 있으면 자동으로 재배포됩니다.',
        'updated_at': datetime.now(timezone.utc).isoformat()
    }

class KRXAHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def _read_json(self):
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            if content_length <= 0:
                return {}
            body = self.rfile.read(content_length)
            if not body:
                return {}
            return json.loads(body.decode('utf-8'))
        except Exception as e:
            return {'_json_error': str(e)}

    def _send_html(self, html_text, status=200):
        body = html_text.encode('utf-8')
        self.send_response(status)
        self.send_header('Content-Type', 'text/html; charset=utf-8')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _send_json(self, data, status=200):
        # KRXA_SKIP_API_LOG_PATCH
        try:
            krxa_api_log_add(getattr(self, 'command', 'GET'), getattr(self, 'path', ''), status=str(status))
        except Exception:
            pass
        body = json.dumps(data, ensure_ascii=False, indent=2).encode('utf-8')
        self.send_response(status)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.end_headers()

    def do_GET(self):
        clean_path = self.path.split('?')[0]
        if clean_path.startswith('/api/dev/deploy/config') or clean_path.startswith('/api/deploy/status'):
            return self._send_json(krxa_v331_deploy_config())
        if clean_path.startswith('/api/dev/source/tree'):
            return self._send_json(krxa_v330_dev_source_tree())
        if clean_path.startswith('/api/dev/source/read'):
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            return self._send_json(krxa_v330_dev_source_read((qs.get('path') or ['app.py'])[0]))
        if clean_path.startswith('/dev/source') or clean_path.startswith('/dev/edit'):
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            return self._send_html(krxa_v330_dev_source_editor_html((qs.get('path') or ['app.py'])[0]))
        if clean_path.startswith('/api/llm/control/status'):
            return self._send_json(krxa_v329_llm_control_status())
        if clean_path.startswith('/api/llm/mode'):
            return self._send_json(krxa_v328_llm_status())
        if clean_path.startswith('/api/dashboard/state'):
            return self._send_json(krxa_v326_dashboard_state())
        if clean_path == '/user' or clean_path == '/user2':
            return self._send_html(krxa_v326_real_control_center_html('user'))
        if clean_path == '/admin':
            return self._send_html(krxa_v326_real_control_center_html('admin'))
        if clean_path == '/dev':
            return self._send_html(krxa_v326_real_control_center_html('dev'))
        if clean_path == '/user' or clean_path.startswith('/user/'):
            return self._send_html(krxa_v325_child_ui_html(clean_path.split('/user/',1)[1]) if clean_path.startswith('/user/') else krxa_v324_user_ui_html())
        if clean_path == '/admin':
            return self._send_html(krxa_v324_admin_ui_html())
        if clean_path.startswith('/admin/'):
            return self._send_html(krxa_v325_child_ui_html(clean_path.split('/admin/',1)[1]))
        if clean_path == '/dev':
            return self._send_html(krxa_v324_dev_ui_html())
        if clean_path.startswith('/dev/'):
            return self._send_html(krxa_v325_child_ui_html(clean_path.split('/dev/',1)[1]))
        if clean_path in ['/control', '/admin', '/operate', '/v32/control', '/v32.3']:
            return self._send_html(krxa_v323_full_control_html())
        if clean_path.startswith('/api/apilog'):
            return self._send_json(krxa_v323_apilog())
        if clean_path.startswith('/api/realtime/status'):
            return self._send_json(krxa_v323_realtime())
        if clean_path.startswith('/api/files/reports'):
            return self._send_json({'ok': True, 'manifest': krxa_v323_manifest()})
        if clean_path in ['/chat', '/integrated', '/v32'] or clean_path.startswith('/chat/'):
            return self._send_html(krxa_chat_integrated_ui_html())
        if clean_path.startswith('/api/chat/history'):
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            uid = (qs.get('user_id') or ['default'])[0]
            return self._send_json({'ok': True, 'history': krxa_chat_history_load(uid)})
        if clean_path.lower().startswith('/hts') or clean_path.startswith('/desktop'):
            return self._send_html(krxa_realtime_desktop_html() if 'krxa_realtime_desktop_html' in globals() else krxa_hts_desktop_html())
        clean_path = self.path.split('?')[0]
        if clean_path.startswith('/api/files/reports'):
            return self._send_json(krxa_safe_files_reports_response())
        if clean_path.startswith('/HTS') or clean_path.startswith('/hts') or clean_path.startswith('/desktop'):
            return self._send_html(krxa_realtime_desktop_html() if 'krxa_realtime_desktop_html' in globals() else krxa_hts_desktop_html())
        
        clean_path = self.path.split('?', 1)[0]
        if clean_path == '/':
            self.path = '/index.html'
            return super().do_GET()
        if clean_path == '/user':
            self.path = '/samples/m2m/user/user_ui.html'
            return super().do_GET()
        if clean_path == '/multi':
            self.path = '/samples/m2m/user/multi_user_ui.html'
            return super().do_GET()
        if clean_path == '/control':
            self.path = '/admin/control_ui.html'
            return super().do_GET()
        if clean_path == '/krxai':
            self.path = '/admin/krxai_console.html'
            return super().do_GET()
        if clean_path in ('/files', '/file_exchange'):
            self.path = '/admin/file_exchange.html'
            return super().do_GET()
        if clean_path in ('/admin', '/admin_ui'):
            self.path = '/admin/admin_ui.html'
            return super().do_GET()
        if clean_path == '/dev':
            self.path = '/admin/dev_ui.html'
            return super().do_GET()
        if clean_path == '/room':
            self.path = '/samples/m2m/user/room_ui.html'
            return super().do_GET()
        if clean_path == '/memory':
            self.path = '/admin/memory_ui.html'
            return super().do_GET()
        if clean_path == '/tasks':
            self.path = '/admin/task_queue_ui.html'
            return super().do_GET()
        if clean_path == '/trust':
            self.path = '/admin/trust_guard_ui.html'
            return super().do_GET()
        if clean_path == '/db':
            self.path = '/admin/db_manager_ui.html'
            return super().do_GET()

        if clean_path.startswith('/api/krxai/discussion/status') or clean_path.startswith('/api/krxai/discuss/status'):
            return self._send_json(krxai_discussion_status())
        if clean_path.startswith('/api/krxai/autonomous/status'):
            return self._send_json(krxai_autonomous_status())
        if clean_path.startswith('/api/krxai/autonomous/run'):
            return self._send_json(krxai_autonomous_run('GET'))
        if clean_path.startswith('/api/control/status'):
            return self._send_json(control_status_payload())
        if clean_path.startswith('/api/control/principles'):
            return self._send_json({'ok': True, 'principles': {k: load_json(v, {}) for k, v in PRINCIPLE_FILES.items()}})
        if clean_path.startswith('/api/control/flow'):
            return self._send_json({'ok': True, 'flow': load_json(ROOMS_FILE, initial_state())})
        if clean_path.startswith('/api/control/db'):
            return self._send_json({'ok': True, 'db': {name: load_json(path, {}) for name, path in DB_FILES.items()}})
        if clean_path.startswith('/api/control/files'):
            return self._send_json({'ok': True, 'files': get_file_manifest()})
        if clean_path.startswith('/api/control/trust'):
            return self._send_json({'ok': True, 'trust': load_json(CORE_DIR / 'trust_guard.json', {})})
        if clean_path.startswith('/api/control/exec'):
            return self._send_json({'ok': True, 'exec': load_json(CORE_DIR / 'autostart_policy.json', {}), 'start_command': 'python app.py'})
        if clean_path.startswith('/api/control/krxai'):
            return self._send_json({'ok': True, 'krxai_last': load_json(USER_DIR / 'krxai_console_last.json', {}), 'llm': get_llm_bridge_config()})
        if clean_path.startswith('/api/control/ai'):
            cfg = get_llm_bridge_config()
            return self._send_json({
                'ok': True,
                'schema': 'KRXA_CONTROL_AI_STATUS_V1',
                'console': '/krxai',
                'ask_api': '/api/krxai/thinking',
                'llm_config_api': '/api/llm/config',
                'policy': cfg.get('call_policy', {}),
                'enabled': bool(cfg.get('enabled')),
                'has_api_key': bool(os.environ.get('OPENAI_API_KEY')),
                'mode': cfg.get('mode', 'smart_call'),
                'rule': 'DB first, call LLM only when policy says needed.'
            })
        if clean_path.startswith('/api/state'):
            return self._send_json(load_json(ROOMS_FILE, initial_state()))
        if clean_path.startswith('/api/health'):
            return self._send_json({'ok': True, 'service': 'KRXA', 'state': 'READY', 'routes': ['/', '/user', '/multi', '/control', '/krxai', '/files', '/api/control/status', '/api/krxai/discussion/status', '/api/krxai/autonomous/status', '/api/krxai/autonomous/run']})
        if clean_path.startswith('/api/platform/status'):
            return self._send_json(platform_status_payload())
        if clean_path.startswith('/files') or clean_path.startswith('/reports'):
            return self._send_html(krxa_file_manager_html())
        if clean_path.startswith('/desktop') or clean_path.startswith('/hts'):
            return self._send_html(krxa_realtime_desktop_html())
        if clean_path.startswith('/window/'):
            name = clean_path.split('/window/', 1)[1].strip('/') or 'krxai'
            return self._send_html(krxa_workspace_html(name))
        if clean_path.startswith('/api/memory/user'):
            qs = urllib.parse.parse_qs(urllib.parse.urlparse(self.path).query)
            uid = (qs.get('user_id') or ['default'])[0]
            return self._send_json({'ok': True, 'memory': krxa_user_memory_load(uid)})
        if clean_path.startswith('/api/tasks/auto/status'):
            return self._send_json({'ok': True, 'engine': load_json(CORE_DIR / 'krxa_auto_task_engine.json', {}), 'queue': load_json(CORE_DIR / 'krxai_task_queue.json', {})})
        if clean_path.startswith('/api/files/reports'):
            manifest = krxa_manifest_load_sorted()
            return self._send_json({'ok': True, 'manifest': manifest})
        if clean_path.startswith('/api/files/download/'):
            name = clean_path.split('/api/files/download/', 1)[1]
            safe = Path(name).name
            target = ROOT / 'files' / 'reports' / safe
            if not target.exists():
                alt = FILE_EXCHANGE_DIR / safe
                if alt.exists():
                    target = alt
                else:
                    return self._send_json({'ok': False, 'error': 'file_not_found', 'filename': safe}, status=404)
            data = target.read_bytes()
            ctype = mimetypes.guess_type(str(target))[0] or 'application/octet-stream'
            self.send_response(200)
            self.send_header('Content-Type', ctype)
            self.send_header('Content-Disposition', f'attachment; filename="{safe}"')
            self.send_header('Content-Length', str(len(data)))
            self.end_headers()
            self.wfile.write(data)
            return
        if clean_path.startswith('/api/krxa/web/learning/status'):
            return self._send_json({
                'ok': True,
                'schema': 'KRXA_AUTO_WEB_LEARNING_LOOP_STATUS_V33.1.3_2',
                'policy': load_json(CORE_DIR / 'krxa_auto_web_learning_policy.json', {}),
                'state': load_json(CORE_DIR / 'krxa_auto_web_learning_state.json', {}),
                'external_fetch_enabled': True,
                'scope': 'user_request_only',
                'updated_at': datetime.now(timezone.utc).isoformat()
            })
        if clean_path.startswith('/api/krxa/web/status'):
            return self._send_json({
                'ok': True,
                'schema': 'KRXA_WEB_RESEARCH_STATUS_V33.1.3',
                'policy': load_json(CORE_DIR / 'krxa_web_research_policy.json', {}),
                'state': load_json(CORE_DIR / 'krxa_web_research_state.json', {}),
                'trusted_mode': True,
                'external_fetch_enabled': False,
                'updated_at': datetime.now(timezone.utc).isoformat()
            })
        if clean_path.startswith('/api/krxai/thinking/status'):
            engine = load_json(CORE_DIR / 'krxai_thinking_engine.json', {})
            return self._send_json({
                'ok': True,
                'schema': 'KRXAI_THINKING_ENGINE_STATUS_V33.1.3',
                'root_loaded': bool(load_json(CORE_DIR / 'krxai_root.json', {})),
                'engine_loaded': bool(engine),
                'memory_index': load_json(CORE_DIR / 'krxai_memory_index.json', {'items': []}),
                'sequence': engine.get('sequence', []),
                'response_schema': engine.get('response_schema', {}),
                'llm_required': False,
                'personality_loaded': bool(load_json(CORE_DIR / 'krxai_personality_profile.json', {})),
                'auto_learning_loaded': bool(load_json(CORE_DIR / 'krxai_auto_learning_policy.json', {})),
                'seed_v2_loaded': bool(load_json(CORE_DIR / 'krxai_core_memory_seed_v2.json', {})),
                'updated_at': datetime.now(timezone.utc).isoformat()
            })
        if clean_path.startswith('/api/role-ui/registry'):
            return self._send_json({'ok': True, 'registry': load_json(CORE_DIR / 'role_ui_registry.json', {})})
        if clean_path.startswith('/api/memory'):
            return self._send_json({'ok': True, 'memory': load_json(CORE_DIR / 'krxai_autonomous_memory.json', {})})
        if clean_path.startswith('/api/tasks'):
            return self._send_json({'ok': True, 'tasks': load_json(CORE_DIR / 'krxai_task_queue.json', {})})
        if clean_path.startswith('/api/files/download/'):
            stored = safe_filename(clean_path.rsplit('/', 1)[-1])
            path = FILE_EXCHANGE_DIR / stored
            if not path.exists() or not path.is_file():
                return self._send_json({'ok': False, 'error': 'file not found'}, 404)
            self.send_response(200)
            self.send_header('Content-Type', 'application/octet-stream')
            self.send_header('Content-Disposition', f'attachment; filename="{stored}"')
            self.send_header('Content-Length', str(path.stat().st_size))
            self.end_headers()
            with path.open('rb') as f:
                shutil.copyfileobj(f, self.wfile)
            return
        if clean_path.startswith('/api/files'):
            return self._send_json({'ok': True, 'manifest': get_file_manifest()})
        if clean_path.startswith('/api/files/upload'):
            try:
                ensure_file_exchange()
                upload = parse_multipart_upload(self)
                original = upload['filename']
                stored = safe_filename(original)
                target = FILE_EXCHANGE_DIR / stored
                if target.exists():
                    stem = target.stem
                    suffix = target.suffix
                    stored = f"{stem}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}{suffix}"
                    target = FILE_EXCHANGE_DIR / stored
                target.write_bytes(upload['content'])
                manifest = add_file_to_manifest(original, stored, target.stat().st_size, upload.get('note', ''))
                return self._send_json({'ok': True, 'file': stored, 'manifest': manifest})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 400)
        if clean_path.startswith('/api/files/delete'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                stored = safe_filename(data.get('stored_name', ''))
                path = FILE_EXCHANGE_DIR / stored
                if path.exists() and path.is_file():
                    path.unlink()
                manifest = get_file_manifest()
                manifest['files'] = [f for f in manifest.get('files', []) if f.get('stored_name') != stored]
                save_json(FILE_MANIFEST, manifest)
                return self._send_json({'ok': True, 'deleted': stored, 'manifest': manifest})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 400)
        if clean_path.startswith('/api/db/'):
            name = clean_path.rsplit('/', 1)[-1]
            path = DB_FILES.get(name)
            if not path:
                return self._send_json({'ok': False, 'error': 'unknown db'}, 404)
            return self._send_json({'ok': True, 'db': name, 'file': str(path.relative_to(ROOT)), 'data': load_json(path, {})})
        if clean_path.startswith('/api/llm/diagnostic'):
            cfg = get_llm_bridge_config()
            return self._send_json({
                'ok': True,
                'schema': 'KRXA_LLM_DIAGNOSTIC_V21',
                'env': {
                    'has_openai_api_key': bool(os.environ.get('OPENAI_API_KEY')),
                    'model': os.environ.get('OPENAI_MODEL', cfg.get('model'))
                },
                'config_enabled': bool(cfg.get('enabled')),
                'provider': cfg.get('provider'),
                'routes': ['/api/llm/config', '/api/llm/log', '/api/llm/diagnostic', '/api/krxai/thinking'],
                'updated_at': datetime.now(timezone.utc).isoformat()
            })
        if clean_path.startswith('/api/llm/config'):
            cfg = get_llm_bridge_config()
            cfg = dict(cfg)
            cfg['has_api_key'] = bool(os.environ.get('OPENAI_API_KEY'))
            return self._send_json({'ok': True, 'config': cfg})
        if clean_path.startswith('/api/llm/log'):
            return self._send_json({'ok': True, 'log': load_json(LLM_LOG_FILE, {'schema': 'KRXA_LLM_CALL_LOG_V1', 'items': []})})
        return super().do_GET()

    def do_POST(self):
        clean_path = self.path.split('?')[0]
        if clean_path.startswith('/api/dev/deploy/github') or clean_path.startswith('/api/deploy/github'):
            data = self._read_json()
            return self._send_json(krxa_v331_deploy_files_to_github(data.get('paths') or ['app.py'], data.get('message') or 'KRXA V33.1 auto deploy from dev UI'))
        if clean_path.startswith('/api/dev/source/save'):
            data = self._read_json()
            return self._send_json(krxa_v330_dev_source_save(data.get('path') or 'app.py', data.get('content') or ''))
        if clean_path.startswith('/api/llm/control/mode'):
            data = self._read_json()
            return self._send_json(krxa_v329_set_llm_control_mode(data.get('mode') or 'auto'))
        if clean_path.startswith('/api/llm/mode') or clean_path.startswith('/api/llm/set_mode'):
            data = self._read_json()
            return self._send_json(krxa_v329_set_llm_control_mode(data.get('mode') or data.get('llm_mode') or 'auto'))
        if clean_path.startswith('/api/tasks/add'):
            data = self._read_json()
            return self._send_json(krxa_v323_task_add(data.get('title','새 작업'), data.get('action','manual_action'), data.get('user_id','default'), data.get('type','manual')))
        if clean_path.startswith('/api/tasks/delete'):
            data = self._read_json()
            return self._send_json(krxa_v323_task_delete(data.get('id') or data.get('task_id')))
        if clean_path.startswith('/api/autonomous/tick') or clean_path.startswith('/api/tasks/auto/tick'):
            data = self._read_json()
            return self._send_json(krxa_v323_task_tick(data.get('user_id','default')))
        if clean_path.startswith('/api/memory/add'):
            data = self._read_json()
            return self._send_json(krxa_v323_memory_add(data.get('user_id','default'), data.get('text') or data.get('input') or (data.get('entry') or {}).get('final','')))
        if clean_path.startswith('/api/memory/delete'):
            data = self._read_json()
            return self._send_json(krxa_v323_memory_delete(data.get('user_id','default'), data.get('id') or data.get('item_id')))
        if clean_path.startswith('/api/files/delete'):
            data = self._read_json()
            return self._send_json(krxa_v323_file_delete(data.get('filename') or data.get('name')))
        if clean_path.startswith('/api/db/save'):
            data = self._read_json()
            return self._send_json(krxa_v323_db_save(data.get('db') or data.get('name'), data.get('data', {})))
        if clean_path.startswith('/api/chat/run'):
            data = self._read_json()
            return self._send_json(krxa_chat_run(data.get('user_id', 'default'), data.get('message') or data.get('text') or data.get('input') or ''))
        if clean_path.startswith('/api/autonomous/tick') or clean_path.startswith('/api/tasks/auto/tick'):
            data = self._read_json()
            return self._send_json(krxa_autonomous_safe_tick(data.get('user_id', 'default')) if 'krxa_autonomous_safe_tick' in globals() else krxa_auto_task_tick(data.get('user_id', 'default')))
        
        clean_path = self.path.split('?')[0]
        if clean_path.startswith('/api/report/from-query'):
            data = self._read_json()
            query = data.get('query') or data.get('message') or data.get('text') or data.get('title') or 'KRXA Report'
            result = krxa_auto_web_learning_run(query if krxa_is_report_request(query) else query + ' 보고서', data.get('user_id', 'default'))
            return self._send_json(result)
        if clean_path.startswith('/api/tasks/auto/tick'):
            data = self._read_json()
            return self._send_json(krxa_auto_task_tick(data.get('user_id', 'default')))
        if clean_path.startswith('/api/report/export'):
            data = self._read_json()
            report = data.get('report') or {'title': data.get('title', 'KRXA Report'), 'content': data.get('content', '')}
            return self._send_json({'ok': True, 'files': krxa_report_export(report, data.get('user_id', 'default'), data.get('title', 'KRXA Report'))})
        if clean_path.startswith('/api/krxa/web/learning/run'):
            data = self._read_json()
            query = data.get('query') or data.get('message') or data.get('text') or ''
            return self._send_json(krxa_auto_web_learning_run(query, data.get('user_id', 'default')))
        if clean_path.startswith('/api/krxa/web/research'):
            data = self._read_json()
            query = data.get('query') or data.get('message') or data.get('text') or ''
            return self._send_json(krxa_web_research_record(query))
        
        clean_path = self.path.split('?', 1)[0]
        if clean_path.startswith('/api/files/upload'):
            try:
                ensure_file_exchange()
                upload = parse_multipart_upload(self)
                original = upload['filename']
                stored = safe_filename(original)
                target = FILE_EXCHANGE_DIR / stored
                if target.exists():
                    stem = target.stem
                    suffix = target.suffix
                    stored = f"{stem}_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}{suffix}"
                    target = FILE_EXCHANGE_DIR / stored
                target.write_bytes(upload['content'])
                manifest = add_file_to_manifest(original, stored, target.stat().st_size, upload.get('note', ''))
                return self._send_json({'ok': True, 'file': stored, 'manifest': manifest})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 400)
        if clean_path.startswith('/api/files/delete'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                stored = safe_filename(data.get('stored_name', ''))
                path = FILE_EXCHANGE_DIR / stored
                if path.exists() and path.is_file():
                    path.unlink()
                manifest = get_file_manifest()
                manifest['files'] = [f for f in manifest.get('files', []) if f.get('stored_name') != stored]
                save_json(FILE_MANIFEST, manifest)
                return self._send_json({'ok': True, 'deleted': stored, 'manifest': manifest})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 400)
        if clean_path.startswith('/api/db/'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                name = clean_path.rsplit('/', 1)[-1]
                path = DB_FILES.get(name)
                if not path:
                    return self._send_json({'ok': False, 'error': 'unknown db'}, 404)
                payload = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                data = payload.get('data')
                if not isinstance(data, dict):
                    return self._send_json({'ok': False, 'error': 'data must be a JSON object'}, 400)
                save_json(path, data)
                return self._send_json({'ok': True, 'db': name, 'file': str(path.relative_to(ROOT)), 'updated_at': datetime.now(timezone.utc).isoformat()})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        if clean_path.startswith('/api/krxai/thinking/ask'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                return self._send_json(krxai_internal_think(data.get('text', ''), data.get('role', 'krxai')))
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        if clean_path.startswith('/api/tasks/add'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                queue = load_json(CORE_DIR / 'krxai_task_queue.json', {'items': []})
                item = {'id': 'task_' + datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S_%f'), 'title': data.get('title','KRXA task'), 'action': data.get('action','internal'), 'status': 'pending', 'requires_llm': False, 'created_at': datetime.now(timezone.utc).isoformat()}
                queue.setdefault('items', []).insert(0, item)
                queue['updated_at'] = datetime.now(timezone.utc).isoformat()
                save_json(CORE_DIR / 'krxai_task_queue.json', queue)
                return self._send_json({'ok': True, 'task': item, 'queue': queue})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        if clean_path.startswith('/api/krxai/discussion') or clean_path.startswith('/api/krxai/discuss'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                return self._send_json(krxai_discussion_core(data.get('text', ''), data.get('source', 'krxai')))
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        if clean_path.startswith('/api/krxai/autonomous/run'):
            return self._send_json(krxai_autonomous_run('POST'))
        if clean_path.startswith('/api/krxai/thinking') or clean_path.startswith('/api/control/ai/ask'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                return self._send_json(krxai_think_structured(data.get('message') or data.get('text') or data.get('input') or '', 'krxai'))
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        if clean_path.startswith('/api/llm/config'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                cfg = get_llm_bridge_config()
                for key in ['enabled', 'provider', 'model', 'mode', 'call_policy']:
                    if key in data:
                        cfg[key] = data[key]
                cfg['updated_at'] = datetime.now(timezone.utc).isoformat()
                save_json(LLM_BRIDGE_FILE, cfg)
                cfg['has_api_key'] = bool(os.environ.get('OPENAI_API_KEY'))
                return self._send_json({'ok': True, 'config': cfg})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        if clean_path.startswith('/api/send'):
            length = int(self.headers.get('Content-Length', '0') or 0)
            try:
                data = json.loads(self.rfile.read(length).decode('utf-8') or '{}')
                msg, state = add_message(
                    data.get('room_id', 'room_default'),
                    data.get('from', 'user_A'),
                    data.get('to', 'user_B'),
                    data.get('text', '')
                )
                return self._send_json({'ok': True, 'message': msg, 'state': state})
            except Exception as e:
                return self._send_json({'ok': False, 'error': str(e)}, 500)
        return self._send_json({'ok': False, 'error': 'not found'}, 404)




if __name__ == '__main__':
    save_json(ROOMS_FILE, load_json(ROOMS_FILE, initial_state()))
    ensure_file_exchange()
    get_llm_bridge_config()
    print('KRXA THINKING ENGINE V33.1.3.2 HTTP ROUTE FIX: READY')
    print('LLM ENV: OPENAI_API_KEY=' + ('SET' if os.environ.get('OPENAI_API_KEY') else 'NOT_SET') + ', OPENAI_MODEL=' + os.environ.get('OPENAI_MODEL', 'gpt-4.1'))
    print(f'BIND: {HOST}:{PORT}')
    print('ROUTES: / /user /multi /control /krxai /files /api/control/status /api/control/ai /api/state /api/send /api/db/<name> /api/krxai/thinking /api/krxai/discussion /api/krxai/discussion/status /api/krxai/autonomous/status /api/krxai/autonomous/run /api/llm/config /api/llm/log /api/files /api/health')
    ThreadingHTTPServer((HOST, PORT), KRXAHandler).serve_forever()
